import base64
import io
import json
import os
import re
import uuid
import smtplib
from email.message import EmailMessage
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone, date, time
from typing import List, Dict, Any, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image
import streamlit as st

# --- Optional OpenAI (kept for PDF parsing flow) ---
try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # type: ignore

from graph_client import GraphClient, GraphConfig, GraphAPIError, GraphAuthError
from audit_log import AuditLog, InterviewStatus, LogLevel, has_processed_message, log_structured, mark_message_processed
from ics_utils import ICSInvite, stable_uid, ICSValidationError, create_ics_from_interview, generate_cancellation_ics
from timezone_utils import to_utc, from_utc, iso_utc, is_valid_timezone, safe_zoneinfo
from export_utils import (
    export_interviews_csv,
    export_audit_log_csv,
    format_audit_entry_human,
    filter_interviews_for_export,
    filter_audit_entries,
    AUDIT_ACTION_DESCRIPTIONS,
)
from calendar_parser import (
    CalendarParser,
    CalendarFormat,
    ParserConfig,
    ParseResult,
    pdf_to_images_enhanced,
)


# ----------------------------
# Input Validation
# ----------------------------
import re as _re
from typing import Tuple as _Tuple

# Email regex (RFC 5322 simplified)
_EMAIL_REGEX = _re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')

# Date/time patterns from OpenAI output
_DATE_REGEX = _re.compile(r'^\d{4}-\d{2}-\d{2}$')
_TIME_REGEX = _re.compile(r'^\d{2}:\d{2}$')


class ValidationError(ValueError):
    """Raised when input validation fails."""
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(f"{field}: {message}")


# Maximum number of candidates allowed in a single batch
MAX_CANDIDATES = 20


@dataclass
class CandidateValidationResult:
    """Result of validating a single candidate entry."""
    original: str           # Original input string
    email: Optional[str]    # Validated email (None if invalid)
    name: str               # Parsed name (empty string if not provided)
    is_valid: bool          # True if email is valid
    error: Optional[str]    # Error message if invalid


@dataclass
class SchedulingResult:
    """Result of scheduling an interview for one candidate."""
    candidate_email: str
    candidate_name: str
    success: bool
    event_id: Optional[str]
    teams_url: Optional[str]
    error: Optional[str]
    warnings: Optional[List[str]] = None
    recipients: Optional[List[str]] = None


def _ensure_candidate_name(name: Optional[str], email: str) -> str:
    """Derive display name from email if name is empty.

    Examples:
        - john.doe@example.com -> "John Doe"
        - jane_smith@corp.com -> "Jane Smith"
        - "" with email -> parsed name or "Candidate"
    """
    if name and name.strip():
        return name.strip()
    if email and "@" in email:
        prefix = email.split("@")[0]
        parts = prefix.replace(".", " ").replace("_", " ").replace("-", " ").split()
        if parts and not prefix.isdigit():
            return " ".join(word.capitalize() for word in parts)
    return "Candidate"


@dataclass
class CompanyConfig:
    """Company branding configuration for emails."""
    name: str
    logo_url: Optional[str]
    primary_color: str
    website: Optional[str]
    sender_email: str

    @property
    def signature_name(self) -> str:
        """Return the company signature name for emails."""
        return f"{self.name} Talent Acquisition Team"


@dataclass
class LayoutConfig:
    """Layout and display configuration for UI branding."""
    show_sidebar: bool
    show_footer: bool
    show_powered_by: bool
    header_style: str  # "full", "compact", "minimal"


def validate_email(email: str, field_name: str = "email") -> str:
    """Validate email format. Returns cleaned email or raises ValidationError."""
    if not email:
        raise ValidationError(field_name, "Email is required")
    email = email.strip().lower()
    if not _EMAIL_REGEX.match(email):
        raise ValidationError(field_name, f"Invalid email format: {email}")
    if len(email) > 254:  # RFC 5321 limit
        raise ValidationError(field_name, "Email too long (max 254 characters)")
    return email


def validate_email_optional(email: Optional[str], field_name: str = "email") -> Optional[str]:
    """Validate email if provided, return None if empty."""
    if not email or not email.strip():
        return None
    return validate_email(email, field_name)


# Pattern for "Name <email>" format
_NAME_EMAIL_PATTERN = _re.compile(r'^(.+?)\s*<([^>]+)>$')


def _parse_single_candidate(entry: str) -> CandidateValidationResult:
    """
    Parse a single candidate entry.
    Supports formats: 'email@example.com' or 'Name <email@example.com>'
    """
    entry = entry.strip()
    if not entry:
        return CandidateValidationResult(
            original=entry,
            email=None,
            name="",
            is_valid=False,
            error="Empty entry"
        )

    # Try to match "Name <email>" format
    match = _NAME_EMAIL_PATTERN.match(entry)
    if match:
        name = match.group(1).strip()
        email_raw = match.group(2).strip()
    else:
        name = ""
        email_raw = entry

    try:
        validated_email = validate_email(email_raw, "candidate email")
        return CandidateValidationResult(
            original=entry,
            email=validated_email,
            name=name,
            is_valid=True,
            error=None
        )
    except ValidationError as e:
        return CandidateValidationResult(
            original=entry,
            email=None,
            name=name,
            is_valid=False,
            error=e.message
        )


def parse_candidate_emails(raw_input: str) -> List[CandidateValidationResult]:
    """
    Parse semicolon-separated candidate emails with optional names.

    Formats supported:
    - "email@example.com"
    - "Name <email@example.com>"
    - "email1@example.com; email2@example.com"

    Returns list of CandidateValidationResult objects with validation status.
    Enforces MAX_CANDIDATES limit and detects duplicates.
    """
    results: List[CandidateValidationResult] = []
    if not raw_input or not raw_input.strip():
        return results

    entries = [e.strip() for e in raw_input.split(';') if e.strip()]

    # Check for exceeding limit
    if len(entries) > MAX_CANDIDATES:
        # Parse first MAX_CANDIDATES normally, mark rest as exceeding limit
        for i, entry in enumerate(entries):
            if i < MAX_CANDIDATES:
                results.append(_parse_single_candidate(entry))
            else:
                results.append(CandidateValidationResult(
                    original=entry,
                    email=None,
                    name="",
                    is_valid=False,
                    error=f"Exceeds maximum of {MAX_CANDIDATES} candidates"
                ))
        return results

    # Track seen emails for duplicate detection
    seen_emails: set = set()

    for entry in entries:
        result = _parse_single_candidate(entry)

        # Check for duplicates among valid entries
        if result.is_valid and result.email:
            if result.email in seen_emails:
                results.append(CandidateValidationResult(
                    original=result.original,
                    email=None,
                    name=result.name,
                    is_valid=False,
                    error="Duplicate email address"
                ))
            else:
                seen_emails.add(result.email)
                results.append(result)
        else:
            results.append(result)

    return results


def validate_slot(slot: dict) -> _Tuple[str, str, str]:
    """Validate slot dict from OpenAI parsing. Returns (date, start, end) tuple."""
    if not isinstance(slot, dict):
        raise ValidationError("slot", "Slot must be a dictionary")

    date = slot.get("date", "")
    start = slot.get("start", "")
    end = slot.get("end", "")

    if not _DATE_REGEX.match(date):
        raise ValidationError("slot.date", f"Invalid date format: {date}. Expected YYYY-MM-DD")
    if not _TIME_REGEX.match(start):
        raise ValidationError("slot.start", f"Invalid start time format: {start}. Expected HH:MM")
    if end and not _TIME_REGEX.match(end):
        raise ValidationError("slot.end", f"Invalid end time format: {end}. Expected HH:MM")

    return date, start, end


# ----------------------------
# Configuration helpers
# ----------------------------
def get_secret(key: str, default: Any = None) -> Any:
    # st.secrets behaves like a dict on Streamlit Cloud; local dev can use env vars too.
    try:
        if key in st.secrets:
            return st.secrets.get(key, default)
    except Exception:
        pass
    return os.getenv(key.upper(), default)


def get_default_timezone() -> str:
    return get_secret("default_timezone", "UTC")


def get_audit_log_path() -> str:
    return get_secret("audit_log_path", "audit_log.db")


def _get_slots_path() -> str:
    """Get path for persistent parsed slots file."""
    return get_secret("slots_storage_path", "parsed_slots.json")


def _load_persisted_slots() -> Dict[str, Any]:
    """Load persisted slots from file. Returns dict with slots, intersections, and panel data."""
    path = _get_slots_path()
    try:
        if os.path.exists(path):
            with open(path, 'r') as f:
                data = json.load(f)
                if isinstance(data, dict) and "slots" in data:
                    return data
    except Exception:
        pass
    return {"slots": [], "computed_intersections": [], "panel_interviewers": [], "next_interviewer_id": 1}


def _save_persisted_slots() -> None:
    """Save current slots and panel interviewer data to persistent storage."""
    path = _get_slots_path()
    try:
        # Serialize panel_interviewers without the non-serializable 'file' field
        panel = st.session_state.get("panel_interviewers", [])
        serializable_panel = [
            {k: v for k, v in interviewer.items() if k != "file"}
            for interviewer in panel
        ]
        data = {
            "slots": st.session_state.get("slots", []),
            "computed_intersections": st.session_state.get("computed_intersections", []),
            "panel_interviewers": serializable_panel,
            "next_interviewer_id": st.session_state.get("next_interviewer_id", 1),
        }
        with open(path, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception:
        pass


def _get_branding_settings_path() -> str:
    """Get path for persistent branding settings file."""
    return get_secret("branding_settings_path", "branding_settings.json")


def _load_branding_settings() -> Dict[str, Any]:
    """Load branding settings from persistent storage."""
    path = _get_branding_settings_path()
    try:
        if os.path.exists(path):
            with open(path, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_branding_settings(settings: Dict[str, Any]) -> None:
    """Save branding settings to persistent storage."""
    path = _get_branding_settings_path()
    try:
        with open(path, 'w') as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        st.warning(f"Could not save branding settings: {e}")


def _get_email_templates_path() -> str:
    """Get path for persistent email templates file."""
    return get_secret("email_templates_path", "email_templates.json")


def _load_email_templates() -> Dict[str, Any]:
    """Load saved email templates from persistent storage."""
    path = _get_email_templates_path()
    try:
        if os.path.exists(path):
            with open(path, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_email_template(name: str, template: Dict[str, Any]) -> bool:
    """Save an email template to persistent storage."""
    path = _get_email_templates_path()
    try:
        templates = _load_email_templates()
        templates[name] = template
        with open(path, 'w') as f:
            json.dump(templates, f, indent=2)
        return True
    except Exception as e:
        st.error(f"Could not save template: {e}")
        return False


def _delete_email_template(name: str) -> bool:
    """Delete an email template from persistent storage."""
    path = _get_email_templates_path()
    try:
        templates = _load_email_templates()
        if name in templates:
            del templates[name]
            with open(path, 'w') as f:
                json.dump(templates, f, indent=2)
            return True
    except Exception as e:
        st.error(f"Could not delete template: {e}")
    return False


def get_graph_config() -> Optional[GraphConfig]:
    tenant_id = get_secret("graph_tenant_id")
    client_id = get_secret("graph_client_id")
    client_secret = get_secret("graph_client_secret")
    scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")

    if not (tenant_id and client_id and client_secret and scheduler_mailbox):
        return None
    return GraphConfig(
        tenant_id=str(tenant_id),
        client_id=str(client_id),
        client_secret=str(client_secret),
        scheduler_mailbox=str(scheduler_mailbox),
    )


def get_company_config() -> CompanyConfig:
    """
    Load company branding configuration.
    Checks session state for user overrides first, then falls back to secrets.
    """
    # Check session state for user customizations (if initialized)
    custom_name = st.session_state.get("custom_company_name") if hasattr(st, "session_state") else None
    custom_logo = st.session_state.get("custom_logo_data") if hasattr(st, "session_state") else None
    email_logo = st.session_state.get("email_logo_url") if hasattr(st, "session_state") else None
    custom_color = st.session_state.get("custom_primary_color") if hasattr(st, "session_state") else None

    # Priority: email_logo_url > custom_logo_data > secret
    logo = email_logo or custom_logo or get_secret("company_logo_url")

    return CompanyConfig(
        name=custom_name or get_secret("company_name", "PowerDash HR"),
        logo_url=logo,
        primary_color=custom_color or get_secret("company_primary_color", "#0066CC"),
        website=get_secret("company_website"),
        sender_email=get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"),
    )


def get_layout_config() -> LayoutConfig:
    """Load layout configuration from secrets for UI branding."""
    return LayoutConfig(
        show_sidebar=get_secret("show_sidebar", False),
        show_footer=get_secret("show_footer", True),
        show_powered_by=get_secret("show_powered_by", True),
        header_style=get_secret("header_style", "full"),
    )


def _lighten_color(hex_color: str, factor: float) -> str:
    """Lighten a hex color by mixing with white."""
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f"#{r:02x}{g:02x}{b:02x}"


def _darken_color(hex_color: str, factor: float) -> str:
    """Darken a hex color."""
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r * (1 - factor))
    g = int(g * (1 - factor))
    b = int(b * (1 - factor))
    return f"#{r:02x}{g:02x}{b:02x}"


def _get_logo_src(path_or_url: Optional[str]) -> Optional[str]:
    """Convert a local file path or URL to a src attribute value for img tags.

    For local files, returns a base64 data URL.
    For URLs (http/https), returns the URL as-is.
    Returns None if path is None or file doesn't exist.
    """
    if not path_or_url:
        return None

    # If it's a URL, return as-is
    if path_or_url.startswith(('http://', 'https://')):
        return path_or_url

    # For local files, convert to base64 data URL
    try:
        # Handle relative paths from app directory
        if not os.path.isabs(path_or_url):
            app_dir = os.path.dirname(os.path.abspath(__file__))
            path_or_url = os.path.join(app_dir, path_or_url)

        if not os.path.exists(path_or_url):
            return None

        with open(path_or_url, 'rb') as f:
            data = f.read()

        # Determine MIME type from extension
        ext = os.path.splitext(path_or_url)[1].lower()
        mime_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.svg': 'image/svg+xml',
            '.ico': 'image/x-icon',
        }
        mime_type = mime_types.get(ext, 'image/png')

        b64 = base64.b64encode(data).decode('utf-8')
        return f"data:{mime_type};base64,{b64}"
    except Exception:
        return None


def graph_enabled() -> bool:
    return get_graph_config() is not None


def get_openai_client() -> Optional[Any]:
    """
    Create an OpenAI client using Streamlit secrets or environment variable.
    Keeps backward compatibility with older secret key names.
    """
    api_key = get_secret("openai_api_key") or get_secret("OPENAI_API_KEY")
    if not api_key:
        st.warning("OpenAI API key not found. PDF availability parsing will be limited.")
        return None
    if OpenAI is None:
        st.warning("OpenAI SDK not available (openai). PDF availability parsing will be limited.")
        return None
    return OpenAI(api_key=api_key)


# ----------------------------
# PDF / image parsing helpers (existing behavior)
# ----------------------------
def image_to_base64(image: Image.Image) -> str:
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def parse_slots_from_image(image: Image.Image, interviewer_timezone: Optional[str] = None, display_timezone: Optional[str] = None) -> List[Dict[str, str]]:
    """
    Use OpenAI vision to parse free/busy calendar images into slots.

    This function uses CalendarParser with format detection and confidence scoring.
    Debug info is stored in st.session_state["parser_debug_info"] when debug mode is enabled.

    Expected return format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30", "confidence": 0.95},
      ...
    ]
    """
    client = get_openai_client()
    if not client:
        return []

    # Build parser config from secrets
    debug_mode = get_secret("parser_debug_mode", "false").lower() == "true"
    config = ParserConfig(
        debug_mode=debug_mode,
        pdf_dpi=int(get_secret("parser_pdf_dpi", "300")),
    )

    # Create parser and set model
    parser = CalendarParser(client, config)
    parser.set_model(get_secret("openai_model", "gpt-5.2"))

    try:
        # Parse with format detection
        result = parser.parse_image(
            image,
            interviewer_timezone=interviewer_timezone,
            display_timezone=display_timezone
        )

        # Store debug info in session state if enabled
        if debug_mode:
            debug_info = {
                "detected_format": result.detected_format.value,
                "format_confidence": result.format_confidence,
                "preprocessing_applied": result.preprocessing_applied,
                "raw_response": result.raw_response[:2000] if result.raw_response else None,
                "slot_count": len(result.slots),
            }
            # Append to existing debug info list or create new
            if "parser_debug_info" not in st.session_state:
                st.session_state["parser_debug_info"] = []
            st.session_state["parser_debug_info"].append(debug_info)

        # Convert to legacy format (backward compatible)
        return result.to_legacy_format()

    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Calendar parser error: {e}",
            action="parse_slots_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def pdf_to_images(pdf_bytes: bytes, max_pages: int = 3) -> List[Image.Image]:
    """Convert PDF to images. Returns empty list on error instead of crashing.

    Uses configurable DPI (default 300) for better parsing accuracy.
    Set parser_pdf_dpi secret to adjust.
    """
    dpi = int(get_secret("parser_pdf_dpi", "300"))
    images: List[Image.Image] = []
    doc = None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i in range(min(len(doc), max_pages)):
            try:
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=dpi)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                images.append(img)
            except Exception as e:
                log_structured(
                    LogLevel.WARNING,
                    f"Failed to process PDF page {i}: {e}",
                    action="pdf_page_process",
                    error_type="pdf_error",
                )
    except Exception as e:
        st.error(f"Failed to open PDF: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to open PDF: {e}",
            action="pdf_open",
            error_type="pdf_error",
            exc_info=True,
        )
    finally:
        if doc:
            doc.close()
    return images


def docx_to_text(docx_bytes: bytes) -> str:
    """
    Extract text from a Word document including paragraphs and tables.
    Returns empty string on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        st.warning("python-docx not installed. Word document parsing unavailable.")
        log_structured(
            LogLevel.ERROR,
            "python-docx not installed",
            action="docx_import",
            error_type="import_error",
        )
        return ""

    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        text_parts: List[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)

        # Extract tables (important for calendar/availability data)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        st.error(f"Failed to read Word document: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to read Word document: {e}",
            action="docx_read",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return ""


def docx_extract_images(docx_bytes: bytes, max_images: int = 5) -> List[Image.Image]:
    """
    Extract embedded images from a Word document.
    Returns empty list on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    images: List[Image.Image] = []
    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))

        # Access the document's related parts to find images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data)).convert("RGB")
                    images.append(img)
                    if len(images) >= max_images:
                        break
                except Exception as e:
                    log_structured(
                        LogLevel.WARNING,
                        f"Failed to extract image from docx: {e}",
                        action="docx_image_extract",
                        error_type="image_error",
                    )
                    continue

        return images
    except Exception as e:
        log_structured(
            LogLevel.WARNING,
            f"Failed to extract images from Word document: {e}",
            action="docx_image_extract",
            error_type=type(e).__name__,
        )
        return []


def split_slot_by_duration(slot: Dict[str, str], duration_minutes: int) -> List[Dict[str, str]]:
    """
    Split a large availability window into discrete meeting slots.

    Args:
        slot: {"date": "2025-01-15", "start": "09:00", "end": "14:00", ...}
        duration_minutes: Meeting duration (e.g., 60)

    Returns:
        List of slots, each exactly duration_minutes long.
        Preserves all other slot metadata (inferred_tz, source, etc.)
    """
    from datetime import datetime, timedelta

    # Guard against invalid duration
    if duration_minutes <= 0:
        return []

    try:
        start_dt = datetime.strptime(f"{slot['date']}T{slot['start']}", "%Y-%m-%dT%H:%M")
        end_dt = datetime.strptime(f"{slot['date']}T{slot['end']}", "%Y-%m-%dT%H:%M")
    except (ValueError, KeyError):
        return []

    # Get total available minutes
    total_minutes = (end_dt - start_dt).total_seconds() / 60

    # If window is smaller than duration, return empty
    if total_minutes < duration_minutes:
        return []

    # If window exactly equals duration, return the slot as-is
    if total_minutes == duration_minutes:
        return [slot.copy()]

    # Split into discrete slots
    result = []
    current_start = start_dt
    duration_delta = timedelta(minutes=duration_minutes)

    while current_start + duration_delta <= end_dt:
        current_end = current_start + duration_delta
        new_slot = slot.copy()
        new_slot["start"] = current_start.strftime("%H:%M")
        new_slot["end"] = current_end.strftime("%H:%M")
        result.append(new_slot)
        current_start = current_end

    return result


def filter_out_scheduled_slots(
    slots: List[Dict[str, str]],
    scheduled_interviews: List[Dict[str, Any]],
    source_timezone: str = "UTC",
) -> List[Dict[str, str]]:
    """
    Filter out slots that overlap with already-scheduled interviews.

    Args:
        slots: List of available slots ({"date": "YYYY-MM-DD", "start": "HH:MM", "end": "HH:MM"})
        scheduled_interviews: List of interview dicts with start_utc and end_utc fields
        source_timezone: Timezone of the slots (for converting to UTC for comparison)

    Returns:
        Filtered list of slots that don't conflict with scheduled interviews
    """
    from datetime import datetime
    import pytz

    if not slots or not scheduled_interviews:
        return slots

    try:
        source_tz = pytz.timezone(source_timezone)
    except Exception:
        source_tz = pytz.UTC

    # Parse scheduled interview times (stored in UTC)
    scheduled_ranges = []
    for interview in scheduled_interviews:
        try:
            start_utc = interview.get("start_utc", "")
            end_utc = interview.get("end_utc", "")
            if start_utc and end_utc:
                # Parse UTC times - handle both formats
                for fmt in ["%Y-%m-%dT%H:%M:%SZ", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S"]:
                    try:
                        start_dt = datetime.strptime(start_utc, fmt).replace(tzinfo=pytz.UTC)
                        end_dt = datetime.strptime(end_utc, fmt).replace(tzinfo=pytz.UTC)
                        scheduled_ranges.append((start_dt, end_dt))
                        break
                    except ValueError:
                        continue
        except Exception:
            continue

    if not scheduled_ranges:
        return slots

    # Filter slots
    filtered_slots = []
    for slot in slots:
        try:
            slot_date = slot.get("date", "")
            slot_start = slot.get("start", "")
            slot_end = slot.get("end", "")

            if not (slot_date and slot_start and slot_end):
                continue

            # Parse slot times in source timezone
            slot_start_dt = source_tz.localize(
                datetime.strptime(f"{slot_date}T{slot_start}", "%Y-%m-%dT%H:%M")
            ).astimezone(pytz.UTC)
            slot_end_dt = source_tz.localize(
                datetime.strptime(f"{slot_date}T{slot_end}", "%Y-%m-%dT%H:%M")
            ).astimezone(pytz.UTC)

            # Check for overlaps with any scheduled interview
            has_conflict = False
            for sched_start, sched_end in scheduled_ranges:
                # Overlap exists if: slot_start < sched_end AND slot_end > sched_start
                if slot_start_dt < sched_end and slot_end_dt > sched_start:
                    has_conflict = True
                    break

            if not has_conflict:
                filtered_slots.append(slot)
        except Exception:
            # If parsing fails, include the slot to be safe
            filtered_slots.append(slot)

    return filtered_slots


def parse_slots_from_text(text: str) -> List[Dict[str, str]]:
    """
    Use OpenAI to parse free/busy text into slots.
    Expected JSON format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30"},
      ...
    ]
    """
    if not text or not text.strip():
        return []

    client = get_openai_client()
    if not client:
        return []

    # Get current year for inference
    current_year = datetime.now().year

    prompt = f"""You are extracting FREE/AVAILABLE time slots from text describing someone's availability.

IMPORTANT RULES:
1. Only extract slots explicitly marked as FREE, AVAILABLE, or OPEN
2. Do NOT include busy/blocked/unavailable times
3. Convert all dates to YYYY-MM-DD format
4. Convert all times to 24-hour HH:MM format
5. If year is not specified, assume {current_year}
6. If end time is not specified, assume 1 hour duration
7. Only include slots that are at least 30 minutes

DATE FORMAT EXAMPLES:
- "Monday Dec 3" -> "{current_year}-12-03"
- "12/03/2025" -> "2025-12-03"
- "3rd December" -> "{current_year}-12-03"
- "Dec 3, 2025" -> "2025-12-03"

TIME FORMAT EXAMPLES:
- "9am-10am" -> start: "09:00", end: "10:00"
- "09:00-10:00" -> start: "09:00", end: "10:00"
- "9:00 AM to 10:00 AM" -> start: "09:00", end: "10:00"
- "2pm-3:30pm" -> start: "14:00", end: "15:30"

Return ONLY valid JSON as a list of objects with keys: date, start, end.
If no free slots found, return an empty list [].

TEXT TO PARSE:
{text}"""

    try:
        resp = client.chat.completions.create(
            model=get_secret("openai_model", "gpt-5.2"),
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns strict JSON. Never include markdown formatting."},
                {"role": "user", "content": prompt},
            ],
        )
        content = resp.choices[0].message.content.strip() if resp.choices else ""

        # Strip code fences if present (same pattern as parse_slots_from_image)
        if content.startswith("```"):
            content = content.strip("`")
            if "\n" in content:
                content = content.split("\n", 1)[1].strip()

        slots = json.loads(content) if content else []
        valid_slots = []
        for s in slots:
            if isinstance(s, dict) and all(k in s for k in ("date", "start", "end")):
                valid_slots.append({
                    "date": str(s["date"]),
                    "start": str(s["start"]),
                    "end": str(s["end"])
                })
        return valid_slots
    except json.JSONDecodeError as e:
        st.error(f"OpenAI returned invalid JSON: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI JSON parse error: {e}",
            action="parse_slots_text_openai",
            error_type="json_decode_error",
        )
        return []
    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI text API error: {e}",
            action="parse_slots_text_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def ensure_session_state() -> None:
    defaults = {
        "slots": [],
        "last_graph_event_id": "",
        "last_teams_join_url": "",
        "last_invite_uid": "",
        "last_invite_ics_bytes": b"",
        "selected_timezone": get_default_timezone(),
        "candidate_timezone": get_default_timezone(),
        "duration_minutes": 30,
        # Panel interview support
        "panel_interviewers": [],  # List of {id, name, email, file, slots, timezone}
        "next_interviewer_id": 1,  # Auto-increment for unique widget keys
        "slot_filter_mode": "all_available",  # "all_available" | "any_n" | "show_all"
        "slot_filter_min_n": 1,  # Minimum N for "any_n" mode
        "computed_intersections": [],  # Intersection slots with availability metadata
        "editing_slot_index": None,  # Track which slot is being edited: (interviewer_idx, slot_idx) or None
        "editing_parsed_slot_index": None,  # int index into session_state["slots"] or None
        "adding_parsed_slot": False,  # Whether the add-slot form is visible
        # Interview management UI state
        "cancelling_interview_id": None,  # ID of interview being cancelled (for confirmation dialog)
        "rescheduling_interview_id": None,  # ID of interview being rescheduled (for confirmation dialog)
        "viewing_interview_history": None,  # Event ID for viewing history
        "interview_status_filter": "All",  # Status filter for interviews list
        # Branding customization (overrides secrets) - loaded from persistent storage
        "custom_company_name": None,  # Override company name from secrets
        "custom_logo_data": None,  # Base64 encoded logo data
        "email_logo_url": None,  # Logo URL for email templates
        "custom_primary_color": None,  # Override primary brand color
        "custom_background_color": None,  # Override background color
        "_branding_loaded": False,  # Track if branding was loaded from file
        # Audit log view state
        "audit_view_mode": "Table",  # "Timeline" | "Table" | "Raw"
        "audit_entry_limit": 300,  # Entry limit selector
        "audit_action_filter": "All",  # Action type filter
        "audit_status_filter": "All",  # Status filter (All/Success/Failed)
        "audit_search": "",  # Search term
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Load persistent branding settings on first run
    if not st.session_state.get("_branding_loaded"):
        saved = _load_branding_settings()
        if saved:
            st.session_state["custom_company_name"] = saved.get("company_name")
            st.session_state["custom_logo_data"] = saved.get("logo_data")
            st.session_state["custom_primary_color"] = saved.get("primary_color")
            st.session_state["custom_background_color"] = saved.get("background_color")
        st.session_state["_branding_loaded"] = True

    # Load persisted slots on first run
    if not st.session_state.get("_slots_loaded"):
        saved_slots = _load_persisted_slots()
        if saved_slots.get("slots"):
            st.session_state["slots"] = saved_slots["slots"]
            st.session_state["computed_intersections"] = saved_slots.get("computed_intersections", [])
        if saved_slots.get("panel_interviewers"):
            # Restore panel interviewers (add file=None since uploads don't persist)
            restored_panel = [
                {**interviewer, "file": None}
                for interviewer in saved_slots["panel_interviewers"]
            ]
            st.session_state["panel_interviewers"] = restored_panel
            st.session_state["next_interviewer_id"] = saved_slots.get("next_interviewer_id", 1)
        st.session_state["_slots_loaded"] = True


def format_slot_label(slot: Dict[str, str]) -> str:
    """Basic slot label for internal use, with confidence indicator if available."""
    base_label = f"{slot['date']} {slot['start']}–{slot['end']}"

    # Add confidence indicator if present
    confidence = slot.get("confidence")
    if confidence is not None:
        try:
            conf_pct = float(confidence) * 100
            if conf_pct < 70:
                # Low confidence - warning indicator
                return f"⚠️ {base_label} ({conf_pct:.0f}%)"
            elif conf_pct >= 90:
                # High confidence - no indicator needed
                return base_label
            else:
                # Medium confidence - subtle indicator
                return f"{base_label} ({conf_pct:.0f}%)"
        except (ValueError, TypeError):
            pass

    return base_label


def format_slot_for_email(
    slot: Dict[str, str],
    source_timezone: Optional[str] = None,
    target_timezone: Optional[str] = None
) -> str:
    """
    Format slot in a professional, human-readable way for emails.
    Example: "Monday, January 26, 2026 • 9:00 AM – 9:30 AM (PST)"

    Args:
        slot: Slot dict with 'date', 'start', 'end' keys
        source_timezone: The timezone the slot times are in (e.g., interviewer's timezone)
        target_timezone: The timezone to display times in (e.g., candidate's timezone)
    """
    from timezone_utils import safe_zoneinfo, to_utc, from_utc

    try:
        date_obj = datetime.strptime(slot['date'], "%Y-%m-%d")
        start_time = datetime.strptime(slot['start'], "%H:%M")
        end_time = datetime.strptime(slot['end'], "%H:%M")

        # Create full datetime objects
        start_dt = datetime.combine(date_obj.date(), start_time.time())
        end_dt = datetime.combine(date_obj.date(), end_time.time())

        # If timezones provided, convert from source to target
        if source_timezone and target_timezone and source_timezone != target_timezone:
            source_zi, _ = safe_zoneinfo(source_timezone, fallback="UTC")

            # Make timezone-aware in source timezone
            start_dt = start_dt.replace(tzinfo=source_zi)
            end_dt = end_dt.replace(tzinfo=source_zi)

            # Convert to target timezone
            start_dt = from_utc(to_utc(start_dt), target_timezone)
            end_dt = from_utc(to_utc(end_dt), target_timezone)

            # Get the date in target timezone (might be different day)
            date_obj = start_dt

        # Format: "Monday, January 26, 2026"
        date_str = date_obj.strftime("%A, %B %d, %Y")

        # Format time in 12-hour with AM/PM: "9:00 AM"
        start_str = start_dt.strftime("%I:%M %p").lstrip("0")
        end_str = end_dt.strftime("%I:%M %p").lstrip("0")

        # Add timezone abbreviation if we did a conversion
        if target_timezone and source_timezone and source_timezone != target_timezone:
            tz_abbrev = start_dt.strftime("%Z") if hasattr(start_dt, 'tzinfo') and start_dt.tzinfo else ""
            if tz_abbrev:
                return f"{date_str} • {start_str} – {end_str} ({tz_abbrev})"

        return f"{date_str} • {start_str} – {end_str}"
    except (ValueError, KeyError):
        # Fallback to basic format
        return f"{slot.get('date', '')} {slot.get('start', '')}–{slot.get('end', '')}"


def _merge_slots(manual_slots: List[Dict], uploaded_slots: List[Dict]) -> List[Dict]:
    """Merge slots, preferring manual over uploaded for duplicates."""
    seen = {}
    for s in manual_slots:
        key = (s["date"], s["start"], s["end"])
        seen[key] = s
    for s in uploaded_slots:
        key = (s["date"], s["start"], s["end"])
        if key not in seen:
            seen[key] = s
    return list(seen.values())


def _add_manual_slot(interviewer_idx: int, slot_date, start_time, end_time) -> bool:
    """Add a manually entered slot with validation. Returns True if successful."""
    from datetime import date as date_type, time as time_type

    errors = []

    # Validate end time is after start time
    if end_time <= start_time:
        errors.append("End time must be after start time")

    # Validate minimum duration (30 minutes)
    start_dt = datetime.combine(date.today(), start_time)
    end_dt = datetime.combine(date.today(), end_time)
    duration_minutes = (end_dt - start_dt).seconds // 60
    if duration_minutes < 30:
        errors.append("Slot must be at least 30 minutes")

    # Validate not in the past
    if slot_date < date.today():
        errors.append("Cannot add slots in the past")

    if errors:
        for err in errors:
            st.error(err)
        return False

    # Create slot in standard format
    new_slot = {
        "date": slot_date.strftime("%Y-%m-%d"),
        "start": start_time.strftime("%H:%M"),
        "end": end_time.strftime("%H:%M"),
        "source": "manual",
    }

    # Get existing slots for this interviewer
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.error("Invalid interviewer index")
        return False

    existing_slots = interviewers[interviewer_idx].get("slots", [])

    # Check for duplicates
    slot_key = (new_slot["date"], new_slot["start"], new_slot["end"])
    for s in existing_slots:
        if (s["date"], s["start"], s["end"]) == slot_key:
            st.warning("This slot already exists")
            return False

    existing_slots.append(new_slot)
    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = existing_slots
    _save_persisted_slots()
    st.success(f"Added slot: {format_slot_label(new_slot)}")
    return True


def _delete_interviewer_slot(interviewer_idx: int, slot_idx: int) -> None:
    """Delete a slot by index from an interviewer's slots."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if 0 <= slot_idx < len(slots):
        deleted = slots.pop(slot_idx)
        st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
        _save_persisted_slots()
        st.toast(f"Deleted: {format_slot_label(deleted)}")
        st.rerun()


def _render_interviewer_slots(interviewer_idx: int, interviewer_id: int) -> None:
    """Render editable list of current slots for an interviewer."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])

    if not slots:
        st.info("No slots added yet. Use the form above or upload a calendar.")
        return

    st.markdown(f"**Current Slots ({len(slots)}):**")

    for idx, slot in enumerate(slots):
        col_label, col_edit, col_delete = st.columns([4, 1, 1])

        with col_label:
            source_badge = " manual" if slot.get("source") == "manual" else " uploaded"
            st.text(f"{source_badge} {format_slot_label(slot)}")

        with col_edit:
            if st.button("Edit", key=f"edit_slot_{interviewer_id}_{idx}"):
                st.session_state["editing_slot_index"] = (interviewer_idx, idx)
                st.rerun()

        with col_delete:
            if st.button("Del", key=f"del_slot_{interviewer_id}_{idx}"):
                _delete_interviewer_slot(interviewer_idx, idx)

    # Clear all button
    if len(slots) > 1:
        if st.button("Clear All Slots", key=f"clear_all_{interviewer_id}", type="secondary"):
            st.session_state["panel_interviewers"][interviewer_idx]["slots"] = []
            _save_persisted_slots()
            st.rerun()


def _render_manual_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render the form to add a new manual slot."""
    st.caption("Add availability slots manually instead of uploading a calendar")

    col_date, col_start, col_end, col_btn = st.columns([2, 1.5, 1.5, 1])

    with col_date:
        slot_date = st.date_input(
            "Date",
            value=date.today(),
            key=f"manual_slot_date_{interviewer_id}",
            min_value=date.today(),
        )
    with col_start:
        slot_start = st.time_input(
            "Start",
            value=time(9, 0),
            key=f"manual_slot_start_{interviewer_id}",
        )
    with col_end:
        slot_end = st.time_input(
            "End",
            value=time(10, 0),
            key=f"manual_slot_end_{interviewer_id}",
        )
    with col_btn:
        st.write("")  # Vertical spacing
        if st.button("+ Add", key=f"add_manual_slot_{interviewer_id}", type="primary"):
            if _add_manual_slot(interviewer_idx, slot_date, slot_start, slot_end):
                st.rerun()


def _render_edit_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render edit form when a slot is being edited."""
    edit_info = st.session_state.get("editing_slot_index")
    if edit_info is None:
        return

    edit_interviewer_idx, edit_slot_idx = edit_info

    # Only render if this is the interviewer being edited
    if edit_interviewer_idx != interviewer_idx:
        return

    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.session_state["editing_slot_index"] = None
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if edit_slot_idx >= len(slots):
        st.session_state["editing_slot_index"] = None
        return

    slot = slots[edit_slot_idx]

    st.markdown("---")
    st.markdown(f"**Editing:** {format_slot_label(slot)}")

    col_date, col_start, col_end = st.columns(3)
    with col_date:
        new_date = st.date_input(
            "Date",
            value=datetime.strptime(slot["date"], "%Y-%m-%d").date(),
            key=f"edit_slot_date_{interviewer_id}",
        )
    with col_start:
        new_start = st.time_input(
            "Start",
            value=datetime.strptime(slot["start"], "%H:%M").time(),
            key=f"edit_slot_start_{interviewer_id}",
        )
    with col_end:
        new_end = st.time_input(
            "End",
            value=datetime.strptime(slot["end"], "%H:%M").time(),
            key=f"edit_slot_end_{interviewer_id}",
        )

    col_save, col_cancel = st.columns(2)
    with col_save:
        if st.button("Save Changes", type="primary", key=f"save_edit_{interviewer_id}"):
            # Validate
            if new_end <= new_start:
                st.error("End time must be after start time")
            elif new_date < date.today():
                st.error("Cannot set date in the past")
            else:
                duration = (datetime.combine(date.today(), new_end) - datetime.combine(date.today(), new_start)).seconds // 60
                if duration < 30:
                    st.error("Slot must be at least 30 minutes")
                else:
                    # Update the slot
                    slots[edit_slot_idx] = {
                        "date": new_date.strftime("%Y-%m-%d"),
                        "start": new_start.strftime("%H:%M"),
                        "end": new_end.strftime("%H:%M"),
                        "source": slot.get("source", "manual"),
                    }
                    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
                    st.session_state["editing_slot_index"] = None
                    _save_persisted_slots()
                    st.success("Slot updated!")
                    st.rerun()

    with col_cancel:
        if st.button("Cancel", key=f"cancel_edit_{interviewer_id}"):
            st.session_state["editing_slot_index"] = None
            st.rerun()

    st.markdown("---")


def _delete_parsed_slot(slot_idx: int) -> None:
    """Remove a slot from session_state['slots'] and computed_intersections."""
    slots = st.session_state.get("slots", [])
    if 0 <= slot_idx < len(slots):
        removed = slots.pop(slot_idx)
        st.session_state["slots"] = slots
        # Also remove from computed_intersections if present
        intersections = st.session_state.get("computed_intersections", [])
        key = (removed["date"], removed["start"], removed["end"])
        st.session_state["computed_intersections"] = [
            s for s in intersections if (s["date"], s["start"], s["end"]) != key
        ]
        # Reset editing index if it was pointing at or beyond removed slot
        if st.session_state.get("editing_parsed_slot_index") is not None:
            if st.session_state["editing_parsed_slot_index"] >= len(st.session_state["slots"]):
                st.session_state["editing_parsed_slot_index"] = None
        _save_persisted_slots()
        st.toast(f"Deleted slot: {format_slot_label(removed)}")
        st.rerun()


def _render_parsed_slots_list(filtered_slots: List[Dict[str, str]]) -> None:
    """Render editable list of parsed slots with Edit/Delete buttons."""
    if not filtered_slots:
        return

    st.markdown("**Parsed Slots:**")

    # Determine if multiple interviewers are involved
    panel_interviewers = st.session_state.get("panel_interviewers", [])
    interviewer_count = len([i for i in panel_interviewers if i.get("slots")])

    # Interviewer filter for multi-interviewer scenarios
    display_slots = filtered_slots
    if interviewer_count > 1:
        # Use actual panel interviewer names for the filter
        interviewer_names = [
            i.get("name") or i.get("email") or f"Interviewer {i['id']}"
            for i in panel_interviewers if i.get("slots")
        ]
        if interviewer_names:
            options = ["All Interviewers"] + interviewer_names
            selected_interviewer = st.selectbox(
                "View slots for",
                options=options,
                key="parsed_slots_interviewer_filter",
            )
            if selected_interviewer != "All Interviewers":
                display_slots = [
                    s for s in filtered_slots
                    if selected_interviewer in s.get("available_names", [])
                ]

    all_slots = st.session_state.get("slots", [])

    # Scrollable container for the slot list (max 300px height)
    with st.container(height=300 if len(display_slots) > 5 else None):
        for slot in display_slots:
            # Find the absolute index in session_state["slots"]
            slot_key = (slot["date"], slot["start"], slot["end"])
            abs_idx = None
            for j, s in enumerate(all_slots):
                if (s["date"], s["start"], s["end"]) == slot_key:
                    abs_idx = j
                    break
            if abs_idx is None:
                continue

            col_label, col_actions = st.columns([5, 2])
            with col_label:
                st.markdown(f"{format_slot_label(slot)}")
            with col_actions:
                btn_edit, btn_del = st.columns(2)
                with btn_edit:
                    if st.button("✏️", key=f"parsed_edit_{abs_idx}", help="Edit slot"):
                        st.session_state["editing_parsed_slot_index"] = abs_idx
                        st.session_state["adding_parsed_slot"] = False
                        st.rerun()
                with btn_del:
                    if st.button("🗑️", key=f"parsed_del_{abs_idx}", help="Delete slot"):
                        _delete_parsed_slot(abs_idx)

    # Clear All button (outside scrollable container)
    if len(filtered_slots) > 1:
        if st.button("Clear All Slots", key="clear_all_parsed_slots"):
            st.session_state["slots"] = []
            st.session_state["computed_intersections"] = []
            st.session_state["editing_parsed_slot_index"] = None
            st.session_state["adding_parsed_slot"] = False
            _save_persisted_slots()
            st.toast("All slots cleared.")
            st.rerun()


def _render_parsed_slot_edit_form() -> None:
    """Render inline edit form for the currently-editing parsed slot."""
    edit_idx = st.session_state.get("editing_parsed_slot_index")
    if edit_idx is None:
        return

    slots = st.session_state.get("slots", [])
    if edit_idx >= len(slots):
        st.session_state["editing_parsed_slot_index"] = None
        return

    slot = slots[edit_idx]

    st.markdown("---")
    st.markdown(f"**Editing:** {format_slot_label(slot)}")

    col_date, col_start, col_end = st.columns(3)
    with col_date:
        new_date = st.date_input(
            "Date",
            value=datetime.strptime(slot["date"], "%Y-%m-%d").date(),
            key="parsed_edit_date",
        )
    with col_start:
        new_start = st.time_input(
            "Start",
            value=datetime.strptime(slot["start"], "%H:%M").time(),
            key="parsed_edit_start",
        )
    with col_end:
        new_end = st.time_input(
            "End",
            value=datetime.strptime(slot["end"], "%H:%M").time(),
            key="parsed_edit_end",
        )

    col_save, col_cancel = st.columns(2)
    with col_save:
        if st.button("Save Changes", type="primary", key="parsed_edit_save"):
            if new_end <= new_start:
                st.error("End time must be after start time")
            elif new_date < date.today():
                st.error("Cannot set date in the past")
            else:
                duration = (datetime.combine(date.today(), new_end) - datetime.combine(date.today(), new_start)).seconds // 60
                if duration < 30:
                    st.error("Slot must be at least 30 minutes")
                else:
                    old_key = (slot["date"], slot["start"], slot["end"])
                    updated_slot = {
                        "date": new_date.strftime("%Y-%m-%d"),
                        "start": new_start.strftime("%H:%M"),
                        "end": new_end.strftime("%H:%M"),
                        "source": slot.get("source", "parsed"),
                    }
                    # Preserve any extra metadata
                    for k in slot:
                        if k not in updated_slot:
                            updated_slot[k] = slot[k]
                    st.session_state["slots"][edit_idx] = updated_slot
                    # Update in computed_intersections too
                    intersections = st.session_state.get("computed_intersections", [])
                    for idx, s in enumerate(intersections):
                        if (s["date"], s["start"], s["end"]) == old_key:
                            for k in updated_slot:
                                intersections[idx][k] = updated_slot[k]
                            break
                    st.session_state["computed_intersections"] = intersections
                    st.session_state["editing_parsed_slot_index"] = None
                    _save_persisted_slots()
                    st.success("Slot updated!")
                    st.rerun()

    with col_cancel:
        if st.button("Cancel", key="parsed_edit_cancel"):
            st.session_state["editing_parsed_slot_index"] = None
            st.rerun()

    st.markdown("---")


def _render_add_parsed_slot_form() -> None:
    """Render form to add a new parsed slot."""
    if not st.session_state.get("adding_parsed_slot"):
        if st.button("+ Add Slot", key="add_parsed_slot_btn"):
            st.session_state["adding_parsed_slot"] = True
            st.session_state["editing_parsed_slot_index"] = None
            st.rerun()
        return

    st.markdown("---")
    st.markdown("**Add New Slot:**")

    col_date, col_start, col_end = st.columns(3)
    with col_date:
        new_date = st.date_input("Date", value=date.today(), key="add_parsed_date")
    with col_start:
        new_start = st.time_input("Start", value=datetime.strptime("09:00", "%H:%M").time(), key="add_parsed_start")
    with col_end:
        new_end = st.time_input("End", value=datetime.strptime("10:00", "%H:%M").time(), key="add_parsed_end")

    col_save, col_cancel = st.columns(2)
    with col_save:
        if st.button("Add", type="primary", key="add_parsed_save"):
            if new_end <= new_start:
                st.error("End time must be after start time")
            elif new_date < date.today():
                st.error("Cannot set date in the past")
            else:
                duration = (datetime.combine(date.today(), new_end) - datetime.combine(date.today(), new_start)).seconds // 60
                if duration < 30:
                    st.error("Slot must be at least 30 minutes")
                else:
                    new_slot = {
                        "date": new_date.strftime("%Y-%m-%d"),
                        "start": new_start.strftime("%H:%M"),
                        "end": new_end.strftime("%H:%M"),
                        "source": "manual",
                    }
                    slot_key = (new_slot["date"], new_slot["start"], new_slot["end"])
                    # Check for duplicates
                    existing_keys = [(s["date"], s["start"], s["end"]) for s in st.session_state.get("slots", [])]
                    if slot_key in existing_keys:
                        st.error("This slot already exists.")
                    else:
                        st.session_state["slots"].append(new_slot)
                        st.session_state["computed_intersections"].append(new_slot.copy())
                        st.session_state["adding_parsed_slot"] = False
                        _save_persisted_slots()
                        st.toast(f"Added slot: {format_slot_label(new_slot)}")
                        st.rerun()

    with col_cancel:
        if st.button("Cancel", key="add_parsed_cancel"):
            st.session_state["adding_parsed_slot"] = False
            st.rerun()

    st.markdown("---")


def _render_parser_debug_panel() -> None:
    """Render collapsible debug panel showing parser diagnostics."""
    debug_mode = get_secret("parser_debug_mode", "false").lower() == "true"
    if not debug_mode:
        return

    debug_info = st.session_state.get("parser_debug_info", [])
    if not debug_info:
        return

    with st.expander("🔍 Parser Debug Info", expanded=False):
        for i, info in enumerate(debug_info):
            st.markdown(f"**Parse #{i + 1}:**")

            col1, col2 = st.columns(2)
            with col1:
                detected_format = info.get("detected_format", "unknown")
                format_conf = info.get("format_confidence", 0) * 100
                st.markdown(f"- **Format:** {detected_format}")
                st.markdown(f"- **Format confidence:** {format_conf:.0f}%")

            with col2:
                preprocessing = info.get("preprocessing_applied", [])
                slot_count = info.get("slot_count", 0)
                st.markdown(f"- **Preprocessing:** {', '.join(preprocessing) if preprocessing else 'None'}")
                st.markdown(f"- **Slots found:** {slot_count}")

            # Show raw response (truncated) if available
            raw_response = info.get("raw_response")
            if raw_response:
                st.text_area(
                    f"Raw response (parse #{i + 1})",
                    value=raw_response[:1500] + ("..." if len(raw_response) > 1500 else ""),
                    height=100,
                    key=f"debug_raw_response_{i}",
                    disabled=True
                )

            if i < len(debug_info) - 1:
                st.markdown("---")

        # Clear debug info button
        if st.button("Clear Debug Info", key="clear_parser_debug"):
            st.session_state["parser_debug_info"] = []
            st.rerun()


def extract_common_timezone(slots: List[Dict[str, str]]) -> Optional[str]:
    """
    Extract the most common inferred timezone from parsed slots.

    Returns IANA timezone or None if no timezone was inferred.
    """
    from collections import Counter
    from timezone_utils import infer_timezone_from_abbreviation

    tz_abbrevs = [s.get("inferred_tz") for s in slots if s.get("inferred_tz")]
    if not tz_abbrevs:
        return None

    # Get most common abbreviation
    most_common = Counter(tz_abbrevs).most_common(1)[0][0]

    # Convert to IANA timezone name
    iana_tz, matched, _ = infer_timezone_from_abbreviation(most_common)
    return iana_tz if matched else None


# ----------------------------
# Email helpers (existing, with updated secret key names)
# ----------------------------
def build_scheduling_email(role_title: str, recruiter_name: str, slots: List[Dict[str, str]]) -> str:
    slot_lines = "\n".join([f"- {format_slot_label(s)}" for s in slots]) if slots else "- (No slots extracted)"
    return f"""Hi there,

Thanks for your time. Please choose one of the following interview times for the role: {role_title}

Available slots:
{slot_lines}

Reply with your preferred option and we will confirm the invite.

Best regards,
{recruiter_name}
Talent Acquisition
"""


def _build_logo_html(company: CompanyConfig) -> str:
    """Build logo HTML section, or empty string if no logo URL configured.

    Supports both URLs and local file paths. Local files are converted to base64 data URLs.
    Note: Some email clients may block base64 images; hosted URLs are more reliable.
    """
    if not company.logo_url:
        return ""

    # Convert local files to base64 data URL, or use URL as-is
    logo_src = _get_logo_src(company.logo_url)
    if not logo_src:
        return ""

    return f'''
    <tr>
        <td align="center" style="padding: 20px 0 10px 0;">
            <img src="{logo_src}" alt="{company.name}"
                 style="max-height: 60px; max-width: 200px; height: auto; display: block;" />
        </td>
    </tr>
    '''


# System font stack for cross-platform email rendering
_EMAIL_FONT_STACK = "-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif"


def build_branded_email_html(
    candidate_name: str,
    role_title: str,
    slots: List[Dict[str, str]],
    company: CompanyConfig,
    custom_message: Optional[str] = None,
    source_timezone: Optional[str] = None,
    target_timezone: Optional[str] = None,
) -> str:
    """
    Build professional HTML email with company branding.

    Uses inline CSS only for email client compatibility.
    Max width 600px, table-based layout.

    Args:
        source_timezone: The timezone the slot times are stored in
        target_timezone: The timezone to display times in (candidate's timezone)
    """
    # Logo section or company name header
    logo_html = _build_logo_html(company)
    if not logo_html:
        # Show company name as header if no logo
        logo_html = f'''
        <tr>
            <td align="center" style="padding: 30px 40px 20px 40px; background: linear-gradient(135deg, {company.primary_color} 0%, {company.primary_color}dd 100%);">
                <h1 style="margin: 0; color: #ffffff; font-family: {_EMAIL_FONT_STACK}; font-size: 24px; font-weight: 600; letter-spacing: -0.5px;">
                    {company.name}
                </h1>
            </td>
        </tr>
        '''

    # Greeting with candidate name
    greeting = f"Dear {candidate_name}," if candidate_name and candidate_name.strip() else "Hello,"

    # Optional custom message
    custom_section = f'<p style="margin: 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">{custom_message}</p>' if custom_message else ""

    # Build slot list HTML with numbered options - using better formatting
    slot_items = ""
    for idx, slot in enumerate(slots, start=1):
        formatted_slot = format_slot_for_email(slot, source_timezone, target_timezone)
        slot_items += f'''
        <tr>
            <td style="padding: 14px 16px; border-left: 4px solid {company.primary_color};
                       background-color: #f8f9fa; border-radius: 0 4px 4px 0; font-family: {_EMAIL_FONT_STACK};">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%">
                    <tr>
                        <td width="32" valign="top">
                            <span style="display: inline-block; width: 26px; height: 26px; line-height: 26px; text-align: center;
                                         background-color: {company.primary_color}; color: #ffffff; border-radius: 50%;
                                         font-size: 13px; font-weight: 600;">{idx}</span>
                        </td>
                        <td style="padding-left: 12px; font-size: 15px; color: #333333;">
                            {formatted_slot}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
        <tr><td style="height: 10px;"></td></tr>
        '''
    if not slots:
        slot_items = f'<tr><td style="padding: 14px 16px; color: #666; font-family: {_EMAIL_FONT_STACK}; background-color: #f8f9fa; border-radius: 4px;">(No slots available)</td></tr>'

    # Website link for footer
    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); overflow: hidden;">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 20px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 16px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <p style="margin: 0 0 20px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.7;">
                                Thank you for your interest in the <strong style="color: {company.primary_color};">{role_title}</strong> position at <strong>{company.name}</strong>. We were impressed with your background and would like to invite you for an interview.
                            </p>
                            {custom_section}
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; font-weight: 600;">
                                Please select one of the following available times:
                            </p>
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="margin: 16px 0 24px 0;">
                                {slot_items}
                            </table>
                            <div style="background-color: #e8f4fd; border-radius: 6px; padding: 16px; margin-top: 20px;">
                                <p style="margin: 0; color: #1a5a96; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; line-height: 1.6;">
                                    <strong>How to respond:</strong> Simply reply to this email with the <strong>number</strong> of your preferred time slot (e.g., "1" or "2"), and we'll send you a calendar invitation with all the details.
                                </p>
                            </div>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 24px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; line-height: 1.6;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
                <!-- Footer -->
                <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="max-width: 600px; width: 100%;">
                    <tr>
                        <td style="padding: 16px 0; text-align: center;">
                            <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 12px; color: #999999;">
                                This email was sent from {company.sender_email}
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_confirmation_email_html(
    candidate_name: str,
    role_title: str,
    interview_time: str,
    teams_url: Optional[str],
    interviewer_names: List[str],
    company: CompanyConfig,
) -> str:
    """
    Build confirmation email after candidate selects a slot.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        interview_time: Formatted interview time in candidate's timezone
        teams_url: Optional Microsoft Teams meeting URL
        interviewer_names: List of interviewer names
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    # Teams meeting section
    meeting_section = ""
    if teams_url:
        meeting_section = f'''
        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
               style="margin: 16px 0; background-color: #f0f7ff; border-left: 4px solid {company.primary_color}; border-radius: 4px;">
            <tr>
                <td style="padding: 16px;">
                    <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                        Microsoft Teams Meeting
                    </p>
                    <a href="{teams_url}" style="color: {company.primary_color}; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; word-break: break-all;">
                        Join Meeting
                    </a>
                </td>
            </tr>
        </table>
        '''

    # Interviewers list
    interviewers_html = ""
    if interviewer_names:
        names_list = ", ".join(interviewer_names)
        interviewers_html = f'<p style="margin: 8px 0 0 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #555555;"><strong>Interviewer(s):</strong> {names_list}</p>'

    # Website link
    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <p style="margin: 0 0 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                Your interview for the <strong style="color: #333333;">{role_title}</strong> position at {company.name} has been confirmed.
                            </p>
                            <!-- Interview details box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #333333;">
                                            <strong>Date & Time:</strong> {interview_time}
                                        </p>
                                        {interviewers_html}
                                    </td>
                                </tr>
                            </table>
                            {meeting_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                A calendar invitation has been sent to your email. If you need to reschedule, please reply to this email.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_branded_email_plain(
    candidate_name: str,
    role_title: str,
    slots: List[Dict[str, str]],
    company: CompanyConfig,
    source_timezone: Optional[str] = None,
    target_timezone: Optional[str] = None,
) -> str:
    """Plain text version of branded email for fallback."""
    greeting = f"Dear {candidate_name}," if candidate_name and candidate_name.strip() else "Hello,"
    slot_lines = "\n".join([f"  {idx}. {format_slot_for_email(s, source_timezone, target_timezone)}" for idx, s in enumerate(slots, start=1)]) if slots else "  (No slots available)"

    footer_parts = [company.signature_name]
    if company.website:
        footer_parts.append(company.website)

    return f"""{greeting}

Thank you for your interest in the {role_title} position at {company.name}. We were impressed with your background and would like to invite you for an interview.

Please select one of the following available times:

{slot_lines}

HOW TO RESPOND: Simply reply to this email with the NUMBER of your preferred time slot (e.g., "1" or "2"), and we will send you a calendar invitation with all the details.

Best regards,
{chr(10).join(footer_parts)}

---
Sent from {company.sender_email}
"""


def build_cancellation_email_html(
    candidate_name: str,
    role_title: str,
    interview_time: str,
    reason: str,
    custom_message: Optional[str],
    company: CompanyConfig,
) -> str:
    """
    Build HTML email for interview cancellation notification.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        interview_time: Formatted interview time
        reason: Cancellation reason
        custom_message: Optional additional message
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    custom_section = ""
    if custom_message:
        custom_section = f'''
        <p style="margin: 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
            {custom_message}
        </p>
        '''

    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <!-- Cancellation notice box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #fff3cd; border-left: 4px solid #ffc107; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 16px; font-weight: 600; color: #856404;">
                                            Interview Cancelled
                                        </p>
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #856404;">
                                            We regret to inform you that your interview for the <strong>{role_title}</strong> position
                                            scheduled for <strong>{interview_time}</strong> has been cancelled.
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            <!-- Reason -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #333333;">
                                            <strong>Reason:</strong> {reason}
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            {custom_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                We apologize for any inconvenience this may cause. If you have any questions, please reply to this email.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_reschedule_email_html(
    candidate_name: str,
    role_title: str,
    old_time: str,
    new_time: str,
    teams_url: Optional[str],
    company: CompanyConfig,
) -> str:
    """
    Build HTML email for interview reschedule notification.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        old_time: Previous interview time (formatted)
        new_time: New interview time (formatted)
        teams_url: Optional Microsoft Teams meeting URL
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    meeting_section = ""
    if teams_url:
        meeting_section = f'''
        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
               style="margin: 16px 0; background-color: #f0f7ff; border-left: 4px solid {company.primary_color}; border-radius: 4px;">
            <tr>
                <td style="padding: 16px;">
                    <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                        Microsoft Teams Meeting
                    </p>
                    <a href="{teams_url}" style="color: {company.primary_color}; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; word-break: break-all;">
                        Join Meeting
                    </a>
                </td>
            </tr>
        </table>
        '''

    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <!-- Reschedule notice box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #d4edda; border-left: 4px solid #28a745; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 16px; font-weight: 600; color: #155724;">
                                            Interview Rescheduled
                                        </p>
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #155724;">
                                            Your interview for the <strong>{role_title}</strong> position has been rescheduled.
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            <!-- Time comparison -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%">
                                            <tr>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #666;">
                                                    <strong>Previous Time:</strong>
                                                </td>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #999; text-decoration: line-through;">
                                                    {old_time}
                                                </td>
                                            </tr>
                                            <tr>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #155724;">
                                                    <strong>New Time:</strong>
                                                </td>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; font-weight: 600; color: #155724;">
                                                    {new_time}
                                                </td>
                                            </tr>
                                        </table>
                                    </td>
                                </tr>
                            </table>
                            {meeting_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                An updated calendar invitation has been sent. Please update your calendar accordingly.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def _smtp_cfg() -> Optional[Dict[str, Any]]:
    # New keys (preferred)
    host = get_secret("smtp_host")
    port = get_secret("smtp_port")
    username = get_secret("smtp_username")
    password = get_secret("smtp_password")
    smtp_from = get_secret("smtp_from")

    # Backward-compat keys (older app)
    if not host:
        host = get_secret("smtp_server")

    if not (host and username and password):
        return None

    return {
        "host": str(host),
        "port": int(port or 587),
        "username": str(username),
        "password": str(password),
        "from": str(smtp_from or username),
    }


def send_email_smtp(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
    content_type: str = "Text",
    plain_text_body: Optional[str] = None,
) -> bool:
    """
    Send email using SMTP (Gmail or other SMTP server).

    Args:
        content_type: "Text" for plain text, "HTML" for HTML emails
        plain_text_body: Optional plain text fallback for HTML emails (multipart/alternative)
    """
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText

    cfg = _smtp_cfg()
    if not cfg:
        return False  # SMTP not configured, caller can try alternative

    # For HTML emails, use multipart/alternative for better client compatibility
    if content_type.upper() == "HTML":
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"] = cfg["from"]
        msg["To"] = ", ".join([e for e in to_emails if e])
        if cc_emails:
            msg["Cc"] = ", ".join([e for e in cc_emails if e])
        msg["MIME-Version"] = "1.0"

        # Add plain text version first (fallback)
        if plain_text_body:
            plain_part = MIMEText(plain_text_body, "plain", "utf-8")
        else:
            # Generate basic plain text from HTML by stripping tags
            import re
            plain_fallback = re.sub(r'<[^>]+>', '', body)
            plain_fallback = re.sub(r'\s+', ' ', plain_fallback).strip()
            plain_part = MIMEText(plain_fallback, "plain", "utf-8")
        msg.attach(plain_part)

        # Add HTML version (preferred) - must come after plain text
        html_part = MIMEText(body, "html", "utf-8")
        msg.attach(html_part)
    else:
        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = cfg["from"]
        msg["To"] = ", ".join([e for e in to_emails if e])
        if cc_emails:
            msg["Cc"] = ", ".join([e for e in cc_emails if e])
        msg.set_content(body)

    if attachment:
        from email.mime.base import MIMEBase
        from email import encoders

        if isinstance(msg, MIMEMultipart):
            # For MIMEMultipart, manually create and attach
            att_part = MIMEBase(
                attachment.get("maintype", "application"),
                attachment.get("subtype", "octet-stream"),
            )
            att_part.set_payload(attachment["data"])
            encoders.encode_base64(att_part)
            att_part.add_header(
                "Content-Disposition",
                "attachment",
                filename=attachment.get("filename", "attachment.bin"),
            )
            msg.attach(att_part)
        else:
            # For EmailMessage, use add_attachment
            msg.add_attachment(
                attachment["data"],
                maintype=attachment.get("maintype", "application"),
                subtype=attachment.get("subtype", "octet-stream"),
                filename=attachment.get("filename", "attachment.bin"),
            )

    try:
        with smtplib.SMTP(cfg["host"], cfg["port"]) as server:
            server.starttls()
            server.login(cfg["username"], cfg["password"])
            server.send_message(msg)
        # Log the email type for debugging
        log_structured(
            LogLevel.INFO,
            f"Email sent successfully via SMTP",
            action="smtp_send",
            details={
                "content_type": content_type,
                "is_multipart": isinstance(msg, MIMEMultipart),
                "mime_type": msg.get_content_type() if hasattr(msg, 'get_content_type') else "unknown",
            },
        )
        return True
    except smtplib.SMTPAuthenticationError as e:
        st.error(f"SMTP authentication failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP authentication failed: {e}",
            action="smtp_send",
            error_type="smtp_auth_error",
        )
        return False
    except smtplib.SMTPException as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type="smtp_error",
        )
        return False
    except Exception as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def send_email_graph(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
    content_type: str = "Text",
    plain_text_body: Optional[str] = None,
) -> bool:
    """
    Send email - tries SMTP (Gmail) first, falls back to Microsoft Graph API.

    Args:
        plain_text_body: For HTML emails, include plain text version for multipart/alternative

    Args:
        subject: Email subject line
        body: Email body (plain text or HTML)
        to_emails: List of recipient email addresses
        cc_emails: Optional list of CC recipients
        attachment: Optional attachment dict with filename, data, maintype, subtype
        content_type: "Text" for plain text, "HTML" for HTML emails
    """
    # Validate recipients before attempting to send
    valid_recipients = [e for e in to_emails if e and e.strip()]
    if not valid_recipients:
        st.error("At least one recipient email is required to send an email.")
        return False

    # Try SMTP first (Gmail)
    smtp_cfg = _smtp_cfg()
    if smtp_cfg:
        try:
            result = send_email_smtp(
                subject=subject,
                body=body,
                to_emails=valid_recipients,
                cc_emails=cc_emails,
                attachment=attachment,
                content_type=content_type,
                plain_text_body=plain_text_body,
            )
            if result:
                format_info = "(HTML with multipart/alternative)" if content_type.upper() == "HTML" else "(plain text)"
                st.success(f"Email sent via Gmail SMTP {format_info}")
                return True
        except Exception as e:
            log_structured(
                LogLevel.WARNING,
                f"SMTP send failed, will try Graph API: {e}",
                action="smtp_send_fallback",
            )

    # Fall back to Graph API
    cfg = get_graph_config()
    if not cfg:
        st.warning("Neither SMTP nor Graph API is configured for sending emails.")
        return False

    try:
        client = GraphClient(cfg)
        graph_attachment = None
        if attachment:
            graph_attachment = {
                "name": attachment.get("filename", "attachment.bin"),
                "contentBytes": attachment.get("data"),
                "contentType": f"{attachment.get('maintype', 'application')}/{attachment.get('subtype', 'octet-stream')}",
            }
        client.send_mail(
            subject=subject,
            body=body,
            to_recipients=valid_recipients,
            cc_recipients=[e for e in (cc_emails or []) if e and e.strip()] or None,
            content_type=content_type,
            attachment=graph_attachment,
        )
        return True
    except Exception as e:
        st.error(f"Graph email send failed: {e}")
        return False


def fetch_emails_imap(include_read: bool = False, limit: int = 20) -> Tuple[List[Dict[str, Any]], Optional[str], bool]:
    """
    Fetch emails from scheduler mailbox via IMAP (Gmail).
    Returns (emails, error_message, is_configured) tuple.
    - error_message is None on success
    - is_configured is False if IMAP credentials are missing
    - include_read: if True, fetches all recent messages (not just unread)
    """
    import imaplib
    import email as email_lib
    from email.header import decode_header
    import socket

    imap_host = get_secret("imap_host", "")
    imap_port = int(get_secret("imap_port", 993))
    imap_username = get_secret("imap_username", "")
    imap_password = get_secret("imap_password", "")

    if not all([imap_host, imap_username, imap_password]):
        return [], None, False  # IMAP not configured

    mail = None
    try:
        # Set socket timeout to prevent hanging
        socket.setdefaulttimeout(30)

        # Connect to IMAP server
        mail = imaplib.IMAP4_SSL(imap_host, imap_port)
        mail.login(imap_username, imap_password)
        mail.select("INBOX", readonly=True)  # readonly to not mark as read

        # Search for emails
        if include_read:
            status, message_ids = mail.search(None, "ALL")
        else:
            status, message_ids = mail.search(None, "UNSEEN")

        if status != "OK":
            mail.logout()
            return [], "Failed to search mailbox", True

        # Get message IDs (most recent first)
        id_list = message_ids[0].split()
        if not id_list:
            mail.logout()
            return [], None, True  # No messages

        id_list = id_list[-limit:]  # Limit to most recent
        id_list.reverse()  # Most recent first

        emails: List[Dict[str, Any]] = []
        for msg_id in id_list:
            try:
                # Fetch message with flags in one request
                status, msg_data = mail.fetch(msg_id, "(FLAGS BODY.PEEK[])")
                if status != "OK" or not msg_data or not msg_data[0]:
                    continue

                # Parse flags
                is_read = False
                flags_part = msg_data[0][0] if isinstance(msg_data[0], tuple) else b""
                if isinstance(flags_part, bytes):
                    is_read = b"\\Seen" in flags_part

                # Get raw email data
                raw_email = None
                for part in msg_data:
                    if isinstance(part, tuple) and len(part) > 1:
                        raw_email = part[1]
                        break

                if not raw_email:
                    continue

                # Parse email
                msg = email_lib.message_from_bytes(raw_email)

                # Decode subject
                subject = ""
                if msg["Subject"]:
                    try:
                        decoded_subject = decode_header(msg["Subject"])
                        for part, encoding in decoded_subject:
                            if isinstance(part, bytes):
                                subject += part.decode(encoding or "utf-8", errors="ignore")
                            else:
                                subject += str(part)
                    except:
                        subject = str(msg["Subject"])

                # Get from address
                from_addr = ""
                if msg["From"]:
                    try:
                        from_decoded = decode_header(msg["From"])
                        for part, encoding in from_decoded:
                            if isinstance(part, bytes):
                                from_addr += part.decode(encoding or "utf-8", errors="ignore")
                            else:
                                from_addr += str(part)
                        # Extract just the email address
                        email_match = re.search(r'[\w\.-]+@[\w\.-]+', from_addr)
                        if email_match:
                            from_addr = email_match.group()
                    except:
                        from_addr = str(msg["From"])

                # Get date
                date_str = msg.get("Date", "")

                # Get body (simplified - just get preview)
                body_content = ""
                try:
                    if msg.is_multipart():
                        # Collect both plain text and HTML versions, use the longest one
                        plain_body = ""
                        html_body = ""
                        for part in msg.walk():
                            content_type = part.get_content_type()
                            if content_type == "text/plain":
                                payload = part.get_payload(decode=True)
                                if payload:
                                    decoded = payload.decode("utf-8", errors="ignore")
                                    if len(decoded) > len(plain_body):
                                        plain_body = decoded
                            elif content_type == "text/html":
                                payload = part.get_payload(decode=True)
                                if payload:
                                    html_content = payload.decode("utf-8", errors="ignore")
                                    # Strip HTML tags to get text content
                                    text_from_html = re.sub(r'<[^>]+>', ' ', html_content)
                                    # Clean up whitespace
                                    text_from_html = re.sub(r'\s+', ' ', text_from_html).strip()
                                    if len(text_from_html) > len(html_body):
                                        html_body = text_from_html
                        # Use the longer of the two (HTML often has complete content)
                        body_content = plain_body if len(plain_body) >= len(html_body) else html_body
                    else:
                        payload = msg.get_payload(decode=True)
                        if payload:
                            body_content = payload.decode("utf-8", errors="ignore")
                        else:
                            body_content = str(msg.get_payload())
                except:
                    body_content = "(Could not decode body)"

                # Store full body for slot detection, truncate for preview
                full_body = body_content.strip()
                preview_body = full_body
                if len(preview_body) > 500:
                    preview_body = preview_body[:500] + "..."

                emails.append({
                    "id": msg_id.decode() if isinstance(msg_id, bytes) else str(msg_id),
                    "from": from_addr,
                    "subject": subject,
                    "date": date_str,
                    "body": preview_body,  # Truncated for display
                    "full_body": full_body,  # Full body for slot detection
                    "is_read": is_read,
                })
            except Exception as e:
                # Skip problematic messages
                continue

        mail.logout()
        return emails, None, True

    except imaplib.IMAP4.error as e:
        if mail:
            try:
                mail.logout()
            except:
                pass
        log_structured(
            LogLevel.ERROR,
            f"IMAP error: {e}",
            action="imap_fetch_messages",
            error_type="imap_error",
        )
        return [], f"IMAP error: {e}", True
    except socket.timeout:
        if mail:
            try:
                mail.logout()
            except:
                pass
        return [], "IMAP connection timed out", True
    except Exception as e:
        if mail:
            try:
                mail.logout()
            except:
                pass
        log_structured(
            LogLevel.ERROR,
            f"Failed to fetch emails via IMAP: {e}",
            action="imap_fetch_messages",
            error_type="imap_error",
            exc_info=True,
        )
        return [], f"Failed to fetch emails: {e}", True
    finally:
        # Reset socket timeout
        socket.setdefaulttimeout(None)


def mark_email_read_imap(msg_id: str) -> bool:
    """
    Mark an email as read via IMAP.
    Returns True if successful, False otherwise.
    """
    import imaplib
    import socket

    imap_host = get_secret("imap_host", "")
    imap_port = int(get_secret("imap_port", 993))
    imap_username = get_secret("imap_username", "")
    imap_password = get_secret("imap_password", "")

    if not all([imap_host, imap_username, imap_password]):
        return False

    mail = None
    try:
        socket.setdefaulttimeout(10)
        mail = imaplib.IMAP4_SSL(imap_host, imap_port)
        mail.login(imap_username, imap_password)
        mail.select("INBOX")  # Not readonly - we need to modify flags

        # Mark as read by adding \Seen flag
        status, _ = mail.store(msg_id, '+FLAGS', '\\Seen')
        mail.logout()
        return status == "OK"
    except Exception as e:
        log_structured(
            LogLevel.WARNING,
            f"Failed to mark email as read: {e}",
            action="mark_email_read",
            error_type="imap_error",
        )
        if mail:
            try:
                mail.logout()
            except:
                pass
        return False
    finally:
        socket.setdefaulttimeout(None)


def fetch_unread_emails_graph(include_read: bool = False) -> Tuple[List[Dict[str, Any]], Optional[str], bool]:
    """
    Fetch emails from scheduler mailbox via Microsoft Graph API.
    Returns (emails, error_message, is_configured) tuple.
    - error_message is None on success
    - is_configured is False if Graph credentials are missing
    - include_read: if True, fetches all recent messages (not just unread)

    Uses the same Graph credentials as calendar operations.
    """
    cfg = get_graph_config()
    if not cfg:
        return [], None, False  # Graph not configured

    try:
        from graph_client import GraphClient
        client = GraphClient(cfg)
        messages = client.fetch_unread_messages(top=50, include_read=include_read)

        emails: List[Dict[str, Any]] = []
        for msg in messages:
            from_addr = ""
            from_data = msg.get("from", {})
            if from_data:
                email_addr = from_data.get("emailAddress", {})
                from_addr = email_addr.get("address", "")

            # Get body content (prefer text, fall back to HTML)
            body_content = msg.get("bodyPreview", "")
            body_data = msg.get("body", {})
            if body_data and body_data.get("content"):
                body_content = body_data.get("content", "")
                # Strip HTML tags if content type is HTML
                if body_data.get("contentType") == "html":
                    import re
                    body_content = re.sub(r'<[^>]+>', '', body_content)
                    body_content = body_content.strip()

            # Store full body for slot detection, truncate for preview display
            full_body = body_content
            preview_body = body_content
            if len(preview_body) > 500:
                preview_body = preview_body[:500] + "..."

            emails.append({
                "id": msg.get("id", ""),
                "from": from_addr,
                "subject": msg.get("subject", ""),
                "date": msg.get("receivedDateTime", ""),
                "body": preview_body,  # Truncated for display
                "full_body": full_body,  # Full body for slot detection
                "is_read": msg.get("isRead", False),
            })

        return emails, None, True  # Success, configured

    except GraphAuthError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph authentication failed: {e}",
            action="graph_fetch_messages",
            error_type="graph_auth_error",
        )
        return [], f"Graph authentication failed: {e}", True
    except GraphAPIError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph API error: {e}",
            action="graph_fetch_messages",
            error_type="graph_api_error",
            details={"status_code": e.status_code},
        )
        return [], f"Graph API error: {e}", True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to fetch emails via Graph: {e}",
            action="graph_fetch_messages",
            error_type="graph_error",
            exc_info=True,
        )
        return [], f"Failed to fetch emails: {e}", True


def _extract_slots_from_email_body(text: str) -> List[Dict[str, str]]:
    """
    Extract numbered slots from the email body.
    Parses patterns like:
      *1.* 2026-01-26 00:00–00:30
      1. 2026-01-26 00:00-00:30
      1) 2026-01-26 00:00 - 00:30
    Returns list of slot dicts with date, start, end keys.
    """
    slots = []
    # Pattern to match numbered slots with various formats
    # Handles: *1.* date time–time, 1. date time-time, 1) date time - time
    slot_pattern = re.compile(
        r'(?:\*?(\d{1,3})[\.\)]\*?\s*)'  # Slot number like *1.* or 1. or 1)
        r'(\d{4}-\d{2}-\d{2})\s+'  # Date like 2026-01-26
        r'(\d{1,2}:\d{2})\s*'  # Start time like 00:00
        r'[–\-—]\s*'  # Dash separator (various unicode dashes)
        r'(\d{1,2}:\d{2})',  # End time like 00:30
        re.MULTILINE
    )

    for match in slot_pattern.finditer(text):
        slot_num = int(match.group(1))
        date = match.group(2)
        start = match.group(3).zfill(5)  # Ensure HH:MM format
        end = match.group(4).zfill(5)

        # Insert at correct position (slot numbers are 1-indexed)
        while len(slots) < slot_num:
            slots.append(None)  # Placeholder
        slots[slot_num - 1] = {"date": date, "start": start, "end": end}

    # Remove any None placeholders
    slots = [s for s in slots if s is not None]
    return slots


def detect_slot_choice_from_text(text: str, slots: List[Dict[str, str]]) -> Optional[Dict[str, str]]:
    """
    Detect which slot the candidate chose from their reply.

    Robust approach:
    - Extract slots from quoted email if present (preferred).
    - Look for the FIRST standalone integer in the reply portion.
    - Fall back to matching full slot label or date/time.

    This is designed to work with messy email replies like:
        "6"
        "Slot 6 please"
        "I would like option 6"
        "6 thanks"
    """
    t = (text or "").strip()
    if not t:
        return None

    # Try to extract slots from the email body (quoted original email)
    email_slots = _extract_slots_from_email_body(t)

    # Use email slots if they exist and are at least as complete as session slots
    effective_slots = email_slots if len(email_slots) >= len(slots or []) else (slots or [])
    if not effective_slots:
        return None

    # Extract reply portion (before quoted email content)
    reply_text = t

    # Common reply markers that indicate quoted history begins
    for marker in [
        "\nOn ", "\r\nOn ",
        "\n>", "\r\n>",
        "Sent from",  # Outlook/Mobile signatures often start quoted section
        "From:", "Date:", "Subject:", "To:", "Cc:"
    ]:
        idx = reply_text.find(marker)
        if idx > 0:
            reply_text = reply_text[:idx]
            break

    # Normalize and split into lines
    lines = [ln.strip() for ln in reply_text.replace("\r", "\n").split("\n")]
    lines = [ln for ln in lines if ln]  # remove blanks

    # Scan line-by-line for the FIRST valid number
    # (most candidates reply with just "6" on its own line)
    for ln in lines[:20]:  # only inspect top part of message
        # Ignore obvious signature fragments
        low = ln.lower()
        if low.startswith(("sent from", "from:", "date:", "subject:", "to:", "cc:")):
            continue

        m = re.fullmatch(r"(\d{1,3})", ln)
        if m:
            slot_num = int(m.group(1))
            if 1 <= slot_num <= len(effective_slots):
                return effective_slots[slot_num - 1]

    # If not standalone, search anywhere in reply for patterns like "slot 6"
    m = re.search(r"\b(?:slot|option|choice|number|#)\s*(\d{1,3})\b", reply_text, re.IGNORECASE)
    if m:
        slot_num = int(m.group(1))
        if 1 <= slot_num <= len(effective_slots):
            return effective_slots[slot_num - 1]

    # Next best: search entire email for first valid number (fallback)
    m = re.search(r"\b(\d{1,3})\b", reply_text)
    if m:
        slot_num = int(m.group(1))
        if 1 <= slot_num <= len(effective_slots):
            return effective_slots[slot_num - 1]

    # Method: Look for full slot label in text
    t_lower = t.lower()
    for s in effective_slots:
        label = format_slot_label(s).lower()
        if label in t_lower:
            return s

    # Method: Look for date and time that match a slot
    m_date = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", reply_text)
    m_time = re.search(r"\b(\d{1,2}:\d{2})\b", reply_text)
    if m_date and m_time:
        date = m_date.group(1)
        start = m_time.group(1).zfill(5)
        for s in effective_slots:
            if s.get("date") == date and s.get("start") == start:
                return s

    return None

    # Method 1: Look for slot number at the START of the reply (before quoted text)
    # Email replies typically have the response at the top, then "On ... wrote:" and quoted text
    # Extract just the reply part (before "On " or ">")
    reply_text = t

    # Split on common reply markers
    for marker in ["\nOn ", "\n>", "\r\n>", "On Mon,", "On Tue,", "On Wed,", "On Thu,", "On Fri,", "On Sat,", "On Sun,"]:
        if marker in reply_text:
            reply_text = reply_text.split(marker)[0]
            break

    reply_text = reply_text.strip()

    # Look for slot number (1-3 digits to support up to 999 slots)
    slot_num_patterns = [
        r"^\s*(\d{1,3})\s*$",  # Just a number like "3" or "70" or "252"
        r"^\s*(\d{1,3})\s*[\n\r.,!]",  # Number at start followed by newline or punctuation
        r"^\s*(\d{1,3})\b",  # Number at the very start
        r"(?:slot|option|choice|number|#)\s*(\d{1,3})\b",  # "slot 70", "#70", etc.
        r"\b(\d{1,3})\s*(?:st|nd|rd|th)?\s*(?:slot|option|choice)",  # "70th slot"
    ]

    for pattern in slot_num_patterns:
        match = re.search(pattern, reply_text, re.IGNORECASE)
        if match:
            try:
                slot_num = int(match.group(1))
                if 1 <= slot_num <= len(effective_slots):
                    return effective_slots[slot_num - 1]  # Convert to 0-indexed
            except (ValueError, IndexError):
                pass

    # Method 2: Look for full slot label in text
    t_lower = t.lower()
    for s in effective_slots:
        label = format_slot_label(s).lower()
        if label in t_lower:
            return s

    # Method 3: Look for date and time that match a slot
    # Only match if the date+time actually corresponds to one of our slots
    m_date = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", reply_text)
    m_time = re.search(r"\b(\d{1,2}:\d{2})\b", reply_text)
    if m_date and m_time:
        date = m_date.group(1)
        start = m_time.group(1).zfill(5)
        for s in effective_slots:
            if s.get("date") == date and s.get("start") == start:
                return s

    return None


# ----------------------------
# Graph + ICS helpers
# ----------------------------
def _make_graph_client() -> Optional[GraphClient]:
    cfg = get_graph_config()
    if not cfg:
        return None
    return GraphClient(cfg)


def _build_professional_invite_body(
    *,
    time_display: str,
    role_title: str,
    duration_minutes: int,
    panel_members: Optional[List[Dict[str, str]]] = None,
    agenda: Optional[str] = None,
    candidates: Optional[List[str]] = None,  # For group invites
) -> str:
    """Build a professional HTML body for calendar invites."""
    company = get_company_config()

    # Panel section
    panel_section = ""
    if panel_members:
        panel_names = [p.get("name") or p.get("email", "") for p in panel_members]
        panel_list = "".join([f'<li style="padding: 4px 0; color: #333333;">{name}</li>' for name in panel_names])
        panel_section = f'''
        <div style="margin: 16px 0; padding: 12px 16px; background-color: #f8f9fa; border-radius: 6px; border-left: 3px solid {company.primary_color};">
            <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                Interview Panel
            </p>
            <ul style="margin: 0; padding-left: 20px; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                {panel_list}
            </ul>
        </div>
        '''

    # Candidates section (for group interviews)
    candidates_section = ""
    if candidates:
        cand_list = "".join([f'<li style="padding: 4px 0; color: #333333;">{c}</li>' for c in candidates])
        candidates_section = f'''
        <div style="margin: 16px 0; padding: 12px 16px; background-color: #fff8e6; border-radius: 6px; border-left: 3px solid #ffc107;">
            <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                Candidates ({len(candidates)})
            </p>
            <ul style="margin: 0; padding-left: 20px; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                {cand_list}
            </ul>
        </div>
        '''

    # Agenda section
    agenda_section = ""
    if agenda and agenda.strip():
        agenda_html = agenda.strip().replace(chr(10), '<br>')
        agenda_section = f'''
        <div style="margin: 16px 0;">
            <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                Agenda
            </p>
            <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #555555; line-height: 1.5;">
                {agenda_html}
            </p>
        </div>
        '''

    return f'''
    <div style="font-family: {_EMAIL_FONT_STACK}; max-width: 600px;">
        <div style="margin-bottom: 20px; padding: 16px; background-color: #e8f4fd; border-radius: 8px; border: 1px solid #b8daff;">
            <p style="margin: 0; font-size: 16px; font-weight: 600; color: {company.primary_color};">
                Interview Details
            </p>
        </div>

        <table style="width: 100%; border-collapse: collapse; margin-bottom: 16px;">
            <tr>
                <td style="padding: 8px 0; font-size: 14px; color: #666666; width: 120px; vertical-align: top;">
                    <strong>Date & Time:</strong>
                </td>
                <td style="padding: 8px 0; font-size: 14px; color: #333333;">
                    {time_display}
                </td>
            </tr>
            <tr>
                <td style="padding: 8px 0; font-size: 14px; color: #666666; vertical-align: top;">
                    <strong>Position:</strong>
                </td>
                <td style="padding: 8px 0; font-size: 14px; color: #333333;">
                    {role_title or "Interview"}
                </td>
            </tr>
            <tr>
                <td style="padding: 8px 0; font-size: 14px; color: #666666; vertical-align: top;">
                    <strong>Duration:</strong>
                </td>
                <td style="padding: 8px 0; font-size: 14px; color: #333333;">
                    {duration_minutes} minutes
                </td>
            </tr>
        </table>

        {candidates_section}
        {panel_section}
        {agenda_section}

        <div style="margin-top: 20px; padding-top: 16px; border-top: 1px solid #e9ecef;">
            <p style="margin: 0; font-size: 13px; color: #888888;">
                Please ensure you join on time. If you need to reschedule, please contact us as soon as possible.
            </p>
        </div>
    </div>
    '''


def _graph_event_payload(
    *,
    subject: str,
    body_html: str,
    start_local: datetime,
    end_local: datetime,
    time_zone: str,
    attendees: List[Tuple[str, str]],
    is_teams: bool,
    location: str,
    cc_attendees: Optional[List[Tuple[str, str]]] = None,
) -> Dict[str, Any]:
    """
    Build Graph API event payload.

    Args:
        attendees: List of (email, name) tuples for required attendees (To:)
        cc_attendees: List of (email, name) tuples for optional attendees (CC:)
    """
    # Build attendees list with required attendees
    all_attendees = [{"emailAddress": {"address": e, "name": n or e}, "type": "required"} for (e, n) in attendees]

    # Add CC attendees as optional
    if cc_attendees:
        all_attendees.extend([{"emailAddress": {"address": e, "name": n or e}, "type": "optional"} for (e, n) in cc_attendees])

    payload: Dict[str, Any] = {
        "subject": subject,
        "body": {"contentType": "HTML", "content": body_html},
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "attendees": all_attendees,
    }

    if is_teams:
        payload["isOnlineMeeting"] = True
        payload["onlineMeetingProvider"] = "teamsForBusiness"
        payload["location"] = {"displayName": "Microsoft Teams"}
    else:
        payload["location"] = {"displayName": location or "Interview"}

    return payload


def _build_ics(
    *,
    organizer_email: str,
    organizer_name: str,
    attendee_emails: List[str],
    summary: str,
    description: str,
    dtstart_utc: datetime,
    dtend_utc: datetime,
    location: str,
    url: str,
    uid_hint: str,
    display_timezone: str = "UTC",
) -> bytes:
    uid = stable_uid(uid_hint, organizer_email, ",".join(attendee_emails), dtstart_utc.isoformat())
    inv = ICSInvite(
        uid=uid,
        dtstart_utc=dtstart_utc,
        dtend_utc=dtend_utc,
        summary=summary,
        description=description,
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=attendee_emails,
        location=location,
        url=url,
        display_timezone=display_timezone,
    )
    return inv.to_ics()


# ----------------------------
# Streamlit UI - Branding Components
# ----------------------------

# NOTE: Sidebar customization and custom background colors disabled - using default Streamlit styling
# def _apply_brand_theme(company: CompanyConfig, background_color: Optional[str] = None) -> None:
#     """Apply client's brand colors to UI elements via CSS."""
#     primary = company.primary_color
#     primary_light = _lighten_color(primary, 0.9)
#     primary_dark = _darken_color(primary, 0.2)
#
#     # Background color CSS (only if custom color is set)
#     bg_css = ""
#     if background_color:
#         bg_css = f"""
# /* Custom background color */
# .stApp, [data-testid="stAppViewContainer"] {{
#     background-color: {background_color} !important;
# }}
# .stMain, [data-testid="stMain"], .main .block-container {{
#     background-color: {background_color} !important;
# }}
# """
#
#     css = f"""<style>
# {bg_css}
# /* Primary buttons */
# .stButton > button[kind="primary"], .stButton > button[data-testid="baseButton-primary"] {{
#     background-color: {primary} !important;
#     border-color: {primary} !important;
# }}
# .stButton > button[kind="primary"]:hover, .stButton > button[data-testid="baseButton-primary"]:hover {{
#     background-color: {primary_dark} !important;
#     border-color: {primary_dark} !important;
# }}
# /* All buttons hover effect */
# .stButton > button:hover {{
#     border-color: {primary} !important;
# }}
# /* Selected tabs */
# .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {{
#     border-bottom-color: {primary} !important;
#     color: {primary} !important;
# }}
# /* Tab highlight bar */
# .stTabs [data-baseweb="tab-highlight"] {{
#     background-color: {primary} !important;
# }}
# /* Links */
# a {{ color: {primary}; }}
# a:hover {{ color: {primary_dark}; }}
# /* Progress bars */
# .stProgress > div > div > div {{
#     background-color: {primary} !important;
# }}
# /* Selectbox/multiselect highlight */
# [data-baseweb="select"] [aria-selected="true"], [data-baseweb="menu"] [aria-selected="true"] {{
#     background-color: {primary_light} !important;
# }}
# /* Checkbox and radio when checked */
# .stCheckbox [data-testid="stCheckbox"] input:checked + div {{
#     background-color: {primary} !important;
#     border-color: {primary} !important;
# }}
# /* Slider */
# .stSlider [data-testid="stThumbValue"], .stSlider [data-baseweb="slider"] div[role="slider"] {{
#     background-color: {primary} !important;
# }}
# /* Sidebar accent */
# [data-testid="stSidebar"] {{
#     border-right: 3px solid {primary};
# }}
# /* Custom branded section class */
# .branded-section {{
#     border-left: 4px solid {primary};
#     padding-left: 16px;
#     margin: 16px 0;
# }}
# </style>"""
#     st.markdown(css, unsafe_allow_html=True)


def _render_header_full(company: CompanyConfig) -> None:
    """Render full header with logo and powered-by badge."""
    css = """<style>
.branded-header { display: flex; align-items: center; justify-content: space-between; padding: 1rem 0; border-bottom: 2px solid #f0f0f0; margin-bottom: 1.5rem; }
.client-branding { display: flex; align-items: center; gap: 16px; }
.client-logo { max-height: 50px; max-width: 180px; object-fit: contain; }
.app-title { font-size: 1.5rem; font-weight: 600; color: #333; margin: 0; }
.powered-by { display: flex; align-items: center; gap: 8px; font-size: 0.75rem; color: #888; }
.powerdash-logo { height: 20px; opacity: 0.7; }
</style>"""
    st.markdown(css, unsafe_allow_html=True)

    client_logo_html = ""
    client_logo_src = _get_logo_src(company.logo_url)
    if client_logo_src:
        client_logo_html = f'<img src="{client_logo_src}" class="client-logo" alt="{company.name}" />'

    powerdash_logo_path = get_secret("powerdash_logo_url", "logo.png")
    powerdash_logo_src = _get_logo_src(powerdash_logo_path)
    layout = get_layout_config()

    powered_by_html = ""
    if layout.show_powered_by and powerdash_logo_src:
        powered_by_html = f'<div class="powered-by"><span>Powered by</span><img src="{powerdash_logo_src}" class="powerdash-logo" alt="PowerDash" /></div>'

    header_html = f'<div class="branded-header"><div class="client-branding">{client_logo_html}<h1 class="app-title">{company.name} Interview Scheduler</h1></div>{powered_by_html}</div>'

    st.markdown(header_html, unsafe_allow_html=True)


def _render_header_compact(company: CompanyConfig) -> None:
    """Render compact single-line header."""
    logo_html = ""
    client_logo_src = _get_logo_src(company.logo_url)
    if client_logo_src:
        logo_html = f'<img src="{client_logo_src}" style="height: 32px; margin-right: 12px;" />'

    layout = get_layout_config()
    powered_by_text = '<span style="font-size: 0.7rem; color: #999;">Powered by PowerDash</span>' if layout.show_powered_by else ''

    html = f'<div style="display: flex; align-items: center; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid #eee; margin-bottom: 16px;"><div style="display: flex; align-items: center;">{logo_html}<span style="font-size: 1.1rem; font-weight: 600;">{company.name} Scheduler</span></div>{powered_by_text}</div>'
    st.markdown(html, unsafe_allow_html=True)


def _render_header_minimal(company: CompanyConfig) -> None:
    """Render minimal text-only header."""
    html = f'<div style="padding: 4px 0; margin-bottom: 12px;"><span style="font-size: 1rem; color: #333;">{company.name}</span><span style="font-size: 0.8rem; color: #999; margin-left: 8px;">Interview Scheduler</span></div>'
    st.markdown(html, unsafe_allow_html=True)


def _render_branded_header(company: CompanyConfig) -> None:
    """Render header based on configured style."""
    layout = get_layout_config()

    if layout.header_style == "compact":
        _render_header_compact(company)
    elif layout.header_style == "minimal":
        _render_header_minimal(company)
    else:
        _render_header_full(company)


def _render_footer() -> None:
    """Render footer with PowerDash branding and links."""
    st.markdown("---")

    css = """<style>
.app-footer { display: flex; justify-content: space-between; align-items: center; padding: 1rem 0; color: #888; font-size: 0.8rem; }
.footer-left { display: flex; align-items: center; gap: 8px; }
.footer-logo { height: 36px; opacity: 0.6; }
.footer-links a { color: #888; text-decoration: none; margin-left: 16px; }
.footer-links a:hover { color: #555; }
</style>"""
    st.markdown(css, unsafe_allow_html=True)

    powerdash_logo_path = get_secret("powerdash_logo_url", "logo.png")
    powerdash_logo_src = _get_logo_src(powerdash_logo_path)
    current_year = datetime.now().year

    logo_html = f'<img src="{powerdash_logo_src}" class="footer-logo" alt="PowerDash" />' if powerdash_logo_src else ''
    footer_html = f'<div class="app-footer"><div class="footer-left">{logo_html}<span>&copy; {current_year} PowerDash HR. All rights reserved.</span></div><div class="footer-links"><a href="https://powerdashhr.com/support" target="_blank">Support</a><a href="https://powerdashhr.com/privacy" target="_blank">Privacy</a></div></div>'

    st.markdown(footer_html, unsafe_allow_html=True)


def _save_current_branding() -> None:
    """Save current branding settings from session state to persistent storage."""
    settings = {
        "company_name": st.session_state.get("custom_company_name"),
        "logo_data": st.session_state.get("custom_logo_data"),
        "primary_color": st.session_state.get("custom_primary_color"),
        "background_color": st.session_state.get("custom_background_color"),
    }
    # Only save if at least one setting is customized
    if any(v is not None for v in settings.values()):
        _save_branding_settings(settings)
    else:
        # Remove the file if all settings are default
        path = _get_branding_settings_path()
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass


# def _render_branding_sidebar() -> None:
#     """Render sidebar with branding customization settings."""
#     with st.sidebar:
#         st.markdown("### Settings")
#
#         # Company name customization
#         default_name = get_secret("company_name", "PowerDash HR")
#         current_name = st.session_state.get("custom_company_name") or default_name
#
#         new_name = st.text_input(
#             "Company Name",
#             value=current_name,
#             key="branding_name_input",
#             help="Customize the company name displayed in the header"
#         )
#
#         if new_name != current_name:
#             if new_name and new_name != default_name:
#                 st.session_state["custom_company_name"] = new_name
#             elif new_name == default_name:
#                 st.session_state["custom_company_name"] = None
#             _save_current_branding()
#             st.rerun()
#
#         st.markdown("---")
#
#         # Logo upload
#         st.markdown("**Company Logo**")
#
#         # Show current logo if set
#         current_logo = st.session_state.get("custom_logo_data")
#         if current_logo:
#             st.image(current_logo, width=150)
#             if st.button("Remove Logo", key="remove_logo_btn"):
#                 st.session_state["custom_logo_data"] = None
#                 _save_current_branding()
#                 st.rerun()
#         else:
#             default_logo_path = get_secret("company_logo_url")
#             if default_logo_path:
#                 logo_src = _get_logo_src(default_logo_path)
#                 if logo_src:
#                     st.image(logo_src, width=150)
#                     st.caption("Default logo from settings")
#
#         uploaded_logo = st.file_uploader(
#             "Upload New Logo",
#             type=["png", "jpg", "jpeg", "gif", "svg"],
#             key="logo_uploader",
#             help="Upload a company logo (PNG, JPG, GIF, or SVG)"
#         )
#
#         if uploaded_logo is not None:
#             # Convert to base64 data URL
#             data = uploaded_logo.read()
#             ext = os.path.splitext(uploaded_logo.name)[1].lower()
#             mime_types = {
#                 '.png': 'image/png',
#                 '.jpg': 'image/jpeg',
#                 '.jpeg': 'image/jpeg',
#                 '.gif': 'image/gif',
#                 '.svg': 'image/svg+xml',
#             }
#             mime_type = mime_types.get(ext, 'image/png')
#             b64 = base64.b64encode(data).decode('utf-8')
#             data_url = f"data:{mime_type};base64,{b64}"
#
#             st.session_state["custom_logo_data"] = data_url
#             _save_current_branding()
#             st.rerun()
#
#         st.markdown("---")
#
#         # Brand color customization
#         st.markdown("**Brand Color**")
#
#         default_color = get_secret("company_primary_color", "#0066CC")
#         current_color = st.session_state.get("custom_primary_color") or default_color
#
#         # Show color preview with computed variants
#         col1, col2 = st.columns([1, 2])
#         with col1:
#             new_color = st.color_picker(
#                 "Primary",
#                 value=current_color,
#                 key="brand_color_picker",
#                 help="Main brand color for buttons, links, and accents"
#             )
#         with col2:
#             light_color = _lighten_color(current_color, 0.9)
#             dark_color = _darken_color(current_color, 0.2)
#             st.markdown(f'<div style="display:flex;gap:4px;margin-top:26px;"><div style="width:24px;height:24px;background:{current_color};border-radius:4px;" title="Primary"></div><div style="width:24px;height:24px;background:{light_color};border-radius:4px;" title="Light"></div><div style="width:24px;height:24px;background:{dark_color};border-radius:4px;" title="Dark"></div></div>', unsafe_allow_html=True)
#             st.caption("Primary · Light · Dark")
#
#         if new_color != current_color:
#             if new_color != default_color:
#                 st.session_state["custom_primary_color"] = new_color
#             else:
#                 st.session_state["custom_primary_color"] = None
#             _save_current_branding()
#             st.rerun()
#
#         # Background color
#         st.markdown("**Background Color**")
#         current_bg = st.session_state.get("custom_background_color")
#
#         col1, col2 = st.columns([1, 2])
#         with col1:
#             # Default to white if not set
#             new_bg = st.color_picker(
#                 "Background",
#                 value=current_bg or "#FFFFFF",
#                 key="bg_color_picker",
#                 help="Page background color"
#             )
#         with col2:
#             if current_bg:
#                 st.markdown(f'<div style="margin-top:26px;padding:8px;background:{current_bg};border:1px solid #ddd;border-radius:4px;font-size:11px;color:#666;">Custom</div>', unsafe_allow_html=True)
#             else:
#                 st.markdown('<div style="margin-top:26px;padding:8px;background:#f0f0f0;border-radius:4px;font-size:11px;color:#888;">Default</div>', unsafe_allow_html=True)
#
#         # Only save if changed from current (and not white which is default)
#         if current_bg and new_bg != current_bg:
#             if new_bg.upper() == "#FFFFFF":
#                 st.session_state["custom_background_color"] = None
#             else:
#                 st.session_state["custom_background_color"] = new_bg
#             _save_current_branding()
#             st.rerun()
#         elif not current_bg and new_bg.upper() != "#FFFFFF":
#             st.session_state["custom_background_color"] = new_bg
#             _save_current_branding()
#             st.rerun()
#
#         st.markdown("---")
#
#         # Reset to defaults button
#         if st.button("Reset to Defaults", key="reset_branding_btn"):
#             st.session_state["custom_company_name"] = None
#             st.session_state["custom_logo_data"] = None
#             st.session_state["custom_primary_color"] = None
#             st.session_state["custom_background_color"] = None
#             _save_current_branding()
#             st.rerun()
#
#         # PowerDash branding at bottom
#         st.markdown("---")
#         powerdash_logo_src = _get_logo_src(get_secret("powerdash_logo_url", "logo.png"))
#         if powerdash_logo_src:
#             st.image(powerdash_logo_src, width=100)
#         st.caption("Powered by PowerDash HR")


# ----------------------------
# Streamlit UI - Main App
# ----------------------------

def main() -> None:
    # Page config must come first - use secrets for initial title
    base_name = get_secret("company_name", "PowerDash HR")
    st.set_page_config(
        page_title=f"{base_name} Interview Scheduler",
        page_icon=get_secret("company_favicon_url", "🗓️"),
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    ensure_session_state()

    # NOTE: Sidebar customization and custom background colors disabled - using default Streamlit styling
    # # Apply brand theme CSS immediately so all UI elements use brand colors
    # company = get_company_config()
    # background_color = st.session_state.get("custom_background_color")
    # _apply_brand_theme(company, background_color)

    # # Render branding settings sidebar
    # _render_branding_sidebar()

    # # Refresh company config in case sidebar changed it
    # company = get_company_config()
    # layout = get_layout_config()
    # background_color = st.session_state.get("custom_background_color")

    # Use default company config without sidebar customization
    company = get_company_config()
    layout = get_layout_config()

    audit = AuditLog(get_audit_log_path())
    _render_branded_header(company)

    tab_new, tab_inbox, tab_invites, tab_audit = st.tabs([
        "📝 New Request",
        "📥 Inbox",
        "📅 Interviews",
        "📜 Audit Log",
    ])

    # ========= TAB: New Scheduling Request =========
    with tab_new:
        st.subheader("New Scheduling Request")

        col_left, col_center, col_right = st.columns([1.2, 1.5, 1.2], gap="large")

        with col_left:
            st.markdown("#### Hiring Manager & Recruiter")
            role_title = st.text_input("Role Title", key="role_title")
            hiring_manager_name = st.text_input("Hiring Manager Name", key="hm_name")
            hiring_manager_email = st.text_input("Hiring Manager Email (required)", key="hm_email")
            recruiter_name = st.text_input("Recruiter Name", key="rec_name")
            recruiter_email = st.text_input("Recruiter Email (optional attendee)", key="rec_email")
            scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")
            st.text_input("Recruiter / Scheduling Mailbox Email", value=str(scheduler_mailbox), disabled=True)

        with col_center:
            st.markdown("#### Interviewer Availability")

            # Ensure at least one interviewer exists
            if not st.session_state.get("panel_interviewers"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"] = [{
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                }]

            interviewers = st.session_state["panel_interviewers"]

            # Render each interviewer
            for idx, interviewer in enumerate(interviewers):
                with st.container(border=True):
                    cols = st.columns([3, 3, 1])
                    with cols[0]:
                        name = st.text_input(
                            "Name *",
                            value=interviewer.get("name", ""),
                            key=f"interviewer_name_{interviewer['id']}",
                            placeholder="e.g., John Smith"
                        )
                        interviewers[idx]["name"] = name
                    with cols[1]:
                        email = st.text_input(
                            "Email *",
                            value=interviewer.get("email", ""),
                            key=f"interviewer_email_{interviewer['id']}",
                            placeholder="john@company.com"
                        )
                        interviewers[idx]["email"] = email
                    with cols[2]:
                        st.write("")  # Spacing
                        # Remove button (disabled if only 1 interviewer)
                        if len(interviewers) > 1:
                            if st.button("✕", key=f"remove_{interviewer['id']}", help="Remove interviewer"):
                                st.session_state["panel_interviewers"] = [
                                    i for i in interviewers if i["id"] != interviewer["id"]
                                ]
                                _save_persisted_slots()
                                st.rerun()

                    # Interviewer timezone selector
                    current_tz = interviewer.get("timezone", st.session_state["selected_timezone"])
                    tz_options = _common_timezones()
                    tz_idx = tz_options.index(current_tz) if current_tz in tz_options else 0
                    interviewer_tz = st.selectbox(
                        "Timezone",
                        options=tz_options,
                        index=tz_idx,
                        key=f"interviewer_tz_{interviewer['id']}",
                        help="The timezone of this interviewer's calendar"
                    )
                    interviewers[idx]["timezone"] = interviewer_tz

                    # File uploader
                    uploaded = st.file_uploader(
                        f"Calendar ({interviewer.get('name') or f'Interviewer {idx+1}'})",
                        type=["pdf", "png", "jpg", "jpeg", "docx"],
                        key=f"file_{interviewer['id']}",
                    )
                    interviewers[idx]["file"] = uploaded

                    # Show slot count with breakdown
                    slot_count = len(interviewer.get("slots", []))
                    manual_count = len([s for s in interviewer.get("slots", []) if s.get("source") == "manual"])
                    uploaded_count = slot_count - manual_count
                    if slot_count > 0:
                        if manual_count > 0 and uploaded_count > 0:
                            st.caption(f"{slot_count} slot(s) ({manual_count} manual, {uploaded_count} uploaded)")
                        elif manual_count > 0:
                            st.caption(f"{slot_count} manual slot(s)")
                        else:
                            st.caption(f"{slot_count} uploaded slot(s)")

                    # Per-interviewer parse button
                    if interviewer.get("file") or interviewer.get("slots"):
                        if st.button(
                            f"Parse {interviewer.get('name') or f'Interviewer {idx+1}'}",
                            key=f"parse_single_{interviewer['id']}",
                            type="primary",
                        ):
                            if not interviewer.get("name", "").strip():
                                st.error("Name is required before parsing.")
                            elif not interviewer.get("email", "").strip():
                                st.error("Email is required before parsing.")
                            else:
                                _parse_single_interviewer_availability(idx)

                    # Manual slot entry expander
                    with st.expander("Manual Slot Entry", expanded=False):
                        _render_manual_slot_form(idx, interviewer["id"])
                        _render_edit_slot_form(idx, interviewer["id"])
                        _render_interviewer_slots(idx, interviewer["id"])

            st.session_state["panel_interviewers"] = interviewers

            # Add interviewer button
            if st.button("+ Add Interviewer", key="add_interviewer_btn"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"].append({
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                })
                _save_persisted_slots()
                st.rerun()

            st.markdown("---")

            st.session_state["duration_minutes"] = st.number_input(
                "Interview duration (minutes)", min_value=15, max_value=240, step=15, value=int(st.session_state["duration_minutes"])
            )
            # Ensure selected_timezone is valid before widget renders
            if st.session_state["selected_timezone"] not in _common_timezones():
                st.session_state["selected_timezone"] = get_default_timezone()
            tz_name = st.selectbox(
                "Display timezone",
                options=_common_timezones(),
                key="selected_timezone",
            )

            # Real-time clock showing current time in selected timezone vs system timezone
            from timezone_utils import from_utc
            now_utc = datetime.now(timezone.utc)
            now_system = datetime.now().astimezone()  # System local time
            system_tz_name = now_system.strftime("%Z")  # e.g., "PST", "GMT"

            try:
                now_selected = from_utc(now_utc, tz_name)
                selected_time = now_selected.strftime("%I:%M %p %Z")
                system_time = now_system.strftime("%I:%M %p %Z")

                st.caption(f"**{tz_name}**: {selected_time} | **Your system ({system_tz_name})**: {system_time}")
            except:
                pass

            parse_btn = st.button("Parse All Availability", type="primary")

            if parse_btn:
                # Validate all interviewers have name and email
                missing = []
                for iv in st.session_state.get("panel_interviewers", []):
                    if not iv.get("name", "").strip() or not iv.get("email", "").strip():
                        label = iv.get("name", "").strip() or f"Interviewer {iv['id']}"
                        missing.append(label)
                if missing:
                    st.error(f"Name and email are required for: {', '.join(missing)}")
                else:
                    _parse_all_panel_availability()

            st.markdown("#### Available Time Slots")

            intersections = st.session_state.get("computed_intersections", [])
            panel_interviewers = st.session_state.get("panel_interviewers", [])
            interviewer_count = len([i for i in panel_interviewers if i.get("slots")])

            if st.session_state["slots"]:
                # Filter mode selector (only show if multiple interviewers)
                if interviewer_count > 1:
                    from slot_intersection import filter_slots_by_availability

                    filter_col1, filter_col2 = st.columns([2, 1])
                    with filter_col1:
                        filter_options = [
                            ("all_available", f"All {interviewer_count} must be available"),
                            ("any_n", "At least N are available"),
                            ("show_all", "Show all slots"),
                        ]
                        filter_mode = st.selectbox(
                            "Show slots where:",
                            options=filter_options,
                            format_func=lambda x: x[1],
                            key="slot_filter_mode_select"
                        )
                        st.session_state["slot_filter_mode"] = filter_mode[0]

                    with filter_col2:
                        if filter_mode[0] == "any_n":
                            min_n = st.number_input(
                                "Minimum N",
                                min_value=1,
                                max_value=interviewer_count,
                                value=max(1, interviewer_count - 1),
                                key="slot_filter_min_n_input"
                            )
                            st.session_state["slot_filter_min_n"] = min_n

                    # Apply filter
                    filtered_slots = filter_slots_by_availability(
                        intersections,
                        st.session_state.get("slot_filter_mode", "all_available"),
                        st.session_state.get("slot_filter_min_n", 1),
                        interviewer_count
                    )
                else:
                    filtered_slots = st.session_state["slots"]

                # Store filtered slots for email generation to use
                st.session_state["filtered_slots_for_email"] = filtered_slots

                if not filtered_slots:
                    st.warning("No slots match the current filter. Try relaxing the availability requirement.")
                    selected_slot = None
                else:
                    st.info("Select a slot to create an invite, or generate a candidate email.")

                    # Editable slots list
                    st.markdown("---")
                    _render_parsed_slots_list(filtered_slots)
                    _render_parsed_slot_edit_form()
                    _render_add_parsed_slot_form()
                    _render_parser_debug_panel()
                    st.markdown("---")

                    # Build slot labels with availability info
                    from slot_intersection import format_slot_label_with_availability

                    def get_slot_label(slot):
                        if interviewer_count > 1:
                            return format_slot_label_with_availability(slot, interviewer_count)
                        return format_slot_label(slot)

                    slot_labels = [get_slot_label(s) for s in filtered_slots]
                    selected_label = st.selectbox("Select slot", options=slot_labels, key="selected_slot_label")
                    selected_slot = filtered_slots[slot_labels.index(selected_label)]

                    # Show availability indicator for panel interviews
                    if interviewer_count > 1 and selected_slot:
                        avail = selected_slot.get("available_count", interviewer_count)
                        total = selected_slot.get("total_interviewers", interviewer_count)
                        available_names = selected_slot.get("available_names", [])

                        if avail == total:
                            st.success(f"All {total} interviewers available")
                        elif avail >= total * 0.75:
                            missing = [
                                i.get("name") or i.get("email")
                                for i in panel_interviewers
                                if i["id"] not in selected_slot.get("available_interviewers", [])
                                and i.get("slots")
                            ]
                            st.info(f"{avail}/{total} available. Missing: {', '.join(missing) if missing else 'None'}")
                        else:
                            st.warning(f"Only {avail}/{total} interviewers available: {', '.join(available_names)}")

                    # Real-time timezone conversion preview
                    if selected_slot:
                        from timezone_utils import safe_zoneinfo, from_utc, format_time_for_display
                        try:
                            # Parse the slot time as display timezone
                            slot_dt_naive = datetime.strptime(
                                f"{selected_slot['date']}T{selected_slot['start']}:00",
                                "%Y-%m-%dT%H:%M:%S"
                            )
                            zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                            slot_dt_local = slot_dt_naive.replace(tzinfo=zi)

                            # Convert to UTC for reference
                            from timezone_utils import to_utc
                            slot_utc = to_utc(slot_dt_local)

                            # Show conversion to common timezones
                            st.markdown("**Time Conversion Preview:**")
                            preview_tzs = ["UTC", "America/New_York", "America/Los_Angeles", "Europe/London", "Asia/Tokyo"]
                            # Add display timezone if not in list
                            if tz_name not in preview_tzs:
                                preview_tzs.insert(0, tz_name)

                            conversion_items = []
                            for preview_tz in preview_tzs:
                                try:
                                    converted = from_utc(slot_utc, preview_tz)
                                    time_str = converted.strftime("%a %b %d, %I:%M %p %Z")
                                    # Highlight the display timezone
                                    if preview_tz == tz_name:
                                        conversion_items.append(f"**{preview_tz}**: {time_str} *(selected)*")
                                    else:
                                        conversion_items.append(f"{preview_tz}: {time_str}")
                                except Exception:
                                    pass

                            st.caption(" | ".join(conversion_items[:4]))  # Show top 4
                        except (ValueError, TypeError):
                            pass  # Skip preview on invalid date

                    # DST Warning Check
                    if selected_slot:
                        from timezone_utils import is_near_dst_transition
                        try:
                            slot_date = datetime.strptime(selected_slot["date"], "%Y-%m-%d").date()
                            slot_dt = datetime.combine(slot_date, datetime.min.time())

                            # Check display timezone for DST transition
                            is_near, trans_date, trans_type = is_near_dst_transition(slot_dt, tz_name, days_threshold=7)
                            if is_near and trans_date:
                                direction = "spring forward" if trans_type == "spring_forward" else "fall back"
                                st.warning(
                                    f"DST Alert: Clocks {direction} on {trans_date.strftime('%B %d, %Y')} "
                                    f"in {tz_name}. Please verify the scheduled time."
                                )
                        except (ValueError, TypeError):
                            pass  # Skip DST check on invalid date
            else:
                st.info("No slots extracted yet. Upload availability and click Parse All Availability.")
                selected_slot = None

        with col_right:
            st.markdown("#### Candidates")
            st.caption("Enter one or more emails separated by semicolons. Format: email or Name <email>")

            candidate_input = st.text_area(
                "Candidate Email(s) (required)",
                key="multi_cand_input",
                height=80,
                placeholder="john@example.com; Jane Doe <jane@example.com>; bob@example.com"
            )

            # Parse and validate candidates
            candidate_results = parse_candidate_emails(candidate_input)
            valid_candidates = [r for r in candidate_results if r.is_valid]
            invalid_count = len(candidate_results) - len(valid_candidates)

            # Display validation results
            if candidate_results:
                if valid_candidates and not invalid_count:
                    st.success(f"All {len(valid_candidates)} candidate(s) validated")
                elif valid_candidates and invalid_count:
                    st.warning(f"{len(valid_candidates)} valid, {invalid_count} invalid candidate(s)")
                elif invalid_count:
                    st.error(f"All {invalid_count} candidate(s) have validation errors")

                # Show validation details in expander
                if len(candidate_results) > 1 or invalid_count > 0:
                    with st.expander("Validation Details", expanded=bool(invalid_count)):
                        for r in candidate_results:
                            if r.is_valid:
                                display_name = f"{r.name} ({r.email})" if r.name else r.email
                                st.markdown(f":white_check_mark: {display_name}")
                            else:
                                st.markdown(f":x: {r.original} - {r.error}")

            # Option to proceed with valid only when there are errors
            proceed_with_valid = True
            if invalid_count > 0 and len(valid_candidates) > 0:
                proceed_with_valid = st.checkbox(
                    f"Proceed with {len(valid_candidates)} valid candidate(s) only",
                    value=True,
                    key="proceed_with_valid"
                )
                if not proceed_with_valid:
                    st.info("Fix invalid candidates before proceeding, or check the box above to skip them.")

            # Scheduling mode selection (only show if multiple valid candidates)
            scheduling_mode = "individual"
            if len(valid_candidates) > 1:
                st.markdown("##### Scheduling Mode")
                scheduling_mode = st.radio(
                    "How to schedule these candidates:",
                    options=["individual", "group"],
                    format_func=lambda x: {
                        "individual": f"Individual Interviews - {len(valid_candidates)} separate invites",
                        "group": "Group Interview - All candidates in one meeting"
                    }[x],
                    key="scheduling_mode",
                    horizontal=True
                )
                if scheduling_mode == "individual":
                    st.caption("Each candidate will receive their own calendar invite.")
                else:
                    st.caption("All candidates will be invited to a single shared meeting.")

            # For backward compatibility, set candidate_email from first valid candidate
            candidate_email = valid_candidates[0].email if valid_candidates else ""
            candidate_name = valid_candidates[0].name if valid_candidates else ""

            # Candidate timezone - pre-populate with inferred timezone from calendar
            inferred_tz = extract_common_timezone(st.session_state.get("slots", []))
            if inferred_tz:
                # Update session state if inference found a timezone
                st.session_state["candidate_timezone"] = inferred_tz

            candidate_tz_default = st.session_state.get("candidate_timezone", get_default_timezone())
            candidate_tz_idx = _common_timezones().index(candidate_tz_default) if candidate_tz_default in _common_timezones() else 0

            candidate_timezone = st.selectbox(
                "Candidate Timezone",
                options=_common_timezones(),
                index=candidate_tz_idx,
                key="candidate_timezone_select",
                help="Times in the invitation will be shown in this timezone"
            )

            if inferred_tz and inferred_tz == candidate_timezone:
                st.caption("Auto-detected from calendar screenshot")

            # Show candidate's view of the selected time
            if selected_slot:
                from timezone_utils import safe_zoneinfo, to_utc, from_utc, format_datetime_for_display
                try:
                    slot_dt_naive = datetime.strptime(
                        f"{selected_slot['date']}T{selected_slot['start']}:00",
                        "%Y-%m-%dT%H:%M:%S"
                    )
                    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                    slot_dt_local = slot_dt_naive.replace(tzinfo=zi)
                    slot_utc = to_utc(slot_dt_local)

                    # Show what candidate will see
                    candidate_view = format_datetime_for_display(slot_utc, candidate_timezone)
                    st.success(f"Candidate will see: **{candidate_view}**")
                except (ValueError, TypeError):
                    pass

            st.markdown("#### Invite details")
            is_teams = st.selectbox("Interview type", options=["Teams", "Non-Teams"], index=0, key="interview_type") == "Teams"
            subject = st.text_input("Subject/title", value=f"Interview: {role_title}" if role_title else "Interview", key="subject")
            agenda = st.text_area("Description/agenda", value="Interview discussion.", key="agenda")
            location = st.text_input("Location (non-Teams)", value="", key="location")

            include_recruiter = st.checkbox("Include recruiter as attendee", value=True, key="include_recruiter")

            st.markdown("----")
            st.markdown("#### Email Branding")

            # Logo URL input
            current_logo = st.session_state.get("email_logo_url") or get_secret("company_logo_url", "")
            logo_url = st.text_input(
                "Logo URL (for email header)",
                value=current_logo or "",
                placeholder="https://example.com/logo.png",
                key="email_logo_url_input",
                help="Enter a publicly accessible URL to your company logo. Leave blank for no logo."
            )
            if logo_url != st.session_state.get("email_logo_url"):
                st.session_state["email_logo_url"] = logo_url if logo_url else None

            # Template management
            st.markdown("##### Email Templates")
            templates = _load_email_templates()
            template_names = list(templates.keys())

            col_load, col_save = st.columns(2)

            with col_load:
                if template_names:
                    selected_template = st.selectbox(
                        "Load template",
                        options=[""] + template_names,
                        key="template_select",
                        format_func=lambda x: "Select a template..." if x == "" else x
                    )
                    if selected_template and st.button("Load", key="load_template_btn"):
                        tpl = templates[selected_template]
                        st.session_state["email_logo_url"] = tpl.get("logo_url")
                        if tpl.get("company_name"):
                            st.session_state["custom_company_name"] = tpl.get("company_name")
                        if tpl.get("primary_color"):
                            st.session_state["custom_primary_color"] = tpl.get("primary_color")
                        st.success(f"Loaded template: {selected_template}")
                        st.rerun()
                else:
                    st.caption("No saved templates yet")

            with col_save:
                new_template_name = st.text_input(
                    "Save as template",
                    placeholder="Template name",
                    key="new_template_name"
                )
                if st.button("Save", key="save_template_btn"):
                    if new_template_name:
                        template_data = {
                            "logo_url": st.session_state.get("email_logo_url"),
                            "company_name": st.session_state.get("custom_company_name"),
                            "primary_color": st.session_state.get("custom_primary_color"),
                        }
                        if _save_email_template(new_template_name, template_data):
                            st.success(f"Template '{new_template_name}' saved!")
                            st.rerun()
                    else:
                        st.warning("Enter a template name to save")

            # Delete template option
            if template_names:
                with st.expander("Delete template"):
                    del_template = st.selectbox(
                        "Select template to delete",
                        options=template_names,
                        key="delete_template_select"
                    )
                    if st.button("Delete", key="delete_template_btn", type="secondary"):
                        if _delete_email_template(del_template):
                            st.success(f"Deleted template: {del_template}")
                            st.rerun()

            st.markdown("----")
            st.markdown("#### Actions")

            # Interviewer selection for email
            panel_interviewers = st.session_state.get("panel_interviewers", [])
            interviewers_with_slots = [i for i in panel_interviewers if i.get("slots")]

            selected_interviewers = []  # List of selected interviewer dicts

            if len(interviewers_with_slots) > 1:
                st.markdown("**Select Interviewer(s) for Interview:**")
                interviewer_options = [
                    {
                        "id": i["id"],
                        "name": i.get("name") or i.get("email") or f"Interviewer {i['id']}",
                        "email": i.get("email", ""),
                        "slots": i.get("slots", []),
                    }
                    for i in interviewers_with_slots
                ]

                # Multi-select for interviewers
                option_names = [f"{opt['name']} ({len(opt['slots'])} slots)" for opt in interviewer_options]
                selected_names = st.multiselect(
                    "Select interviewer(s) to include",
                    options=option_names,
                    default=[option_names[0]] if option_names else [],
                    key="email_interviewer_multiselect",
                    help="Select one or more interviewers. Email will show slots where ALL selected interviewers are available."
                )

                # Map selected names back to interviewer data
                for name in selected_names:
                    idx = option_names.index(name)
                    selected_interviewers.append(interviewer_options[idx])

                if len(selected_interviewers) == 1:
                    st.info(f"📋 Will send **{selected_interviewers[0]['name']}**'s available slots to candidate")
                elif len(selected_interviewers) > 1:
                    names = ", ".join([i['name'] for i in selected_interviewers])
                    st.info(f"📋 Will send slots where **all {len(selected_interviewers)} interviewers** are available to candidate ({names})")
                else:
                    st.warning("Please select at least one interviewer")

            elif len(interviewers_with_slots) == 1:
                # Single interviewer - auto-select
                selected_interviewers = [{
                    "id": interviewers_with_slots[0]["id"],
                    "name": interviewers_with_slots[0].get("name") or interviewers_with_slots[0].get("email") or "Interviewer",
                    "email": interviewers_with_slots[0].get("email", ""),
                    "slots": interviewers_with_slots[0].get("slots", []),
                }]

            # Generate branded email to candidate
            if st.button("Generate Candidate Scheduling Email"):
                company = get_company_config()

                # Use the current filtered/edited slots from session state
                # This ensures edits/deletions made in the UI are reflected in the email
                # and respects the current availability filter mode
                source_slots = st.session_state.get("filtered_slots_for_email", st.session_state.get("slots", []))

                if not source_slots:
                    st.error("No slots available. Please add availability first.")
                else:
                    # Split slots by duration for email (ensures discrete meeting slots)
                    email_slots = []
                    duration = st.session_state.get("duration_minutes", 30)
                    for slot in source_slots:
                        split_result = split_slot_by_duration(slot, duration)
                        if split_result:
                            for ss in split_result:
                                # Preserve metadata
                                for key in slot:
                                    if key not in ss:
                                        ss[key] = slot[key]
                            email_slots.extend(split_result)
                        else:
                            email_slots.append(slot)

                    # Filter out slots that have already been scheduled
                    scheduled_interviews = audit.get_active_interviews()
                    slots_before_filter = len(email_slots)
                    email_slots = filter_out_scheduled_slots(
                        email_slots,
                        scheduled_interviews,
                        source_timezone=tz_name,
                    )
                    slots_removed = slots_before_filter - len(email_slots)
                    if slots_removed > 0:
                        st.info(f"Excluded {slots_removed} slot(s) that already have interviews scheduled.")

                    if not email_slots:
                        st.error("No available slots remaining after excluding already-scheduled interviews.")
                    else:
                        html_body = build_branded_email_html(
                            candidate_name=candidate_name,
                            role_title=role_title or "Position",
                            slots=email_slots,
                            company=company,
                            source_timezone=tz_name,
                            target_timezone=candidate_timezone,
                        )
                        plain_body = build_branded_email_plain(
                            candidate_name=candidate_name,
                            role_title=role_title or "Position",
                            slots=email_slots,
                            company=company,
                            source_timezone=tz_name,
                            target_timezone=candidate_timezone,
                        )
                        st.session_state["candidate_email_html"] = html_body
                        st.session_state["candidate_email_plain"] = plain_body
                        st.session_state["candidate_email_generated_at"] = datetime.now().strftime("%H:%M:%S")

                        if len(selected_interviewers) == 1:
                            interviewer_note = f" (for {selected_interviewers[0]['name']})"
                        elif len(selected_interviewers) > 1:
                            names = ", ".join([i['name'] for i in selected_interviewers])
                            interviewer_note = f" (panel: {names})"
                        else:
                            interviewer_note = ""
                        st.success(f"Generated branded HTML email template with {len(email_slots)} slot(s){interviewer_note}")

            if st.session_state.get("candidate_email_html"):
                gen_time = st.session_state.get("candidate_email_generated_at", "unknown")
                st.caption(f"📧 Branded HTML email template (generated at {gen_time})")

                # Preview mode toggle
                preview_mode = st.radio(
                    "Preview mode",
                    options=["Rendered", "HTML Source", "Plain Text"],
                    horizontal=True,
                    key="email_preview_mode"
                )

                if preview_mode == "Rendered":
                    st.markdown("**Email Preview (Rendered):**")
                    import streamlit.components.v1 as components
                    components.html(st.session_state["candidate_email_html"], height=500, scrolling=True)
                elif preview_mode == "HTML Source":
                    st.code(st.session_state["candidate_email_html"], language="html")
                else:
                    st.text_area("Email preview (Plain Text)", st.session_state["candidate_email_plain"], height=300)

                # Slots email goes ONLY to candidate - Hiring Manager and Recruiter
                # will receive the calendar invite when the interview is scheduled
                st.info(f"📧 Slots email will be sent to: **{candidate_email}** only (Interviewers, Hiring Manager & Recruiter will receive the calendar invite when scheduled)")

                company = get_company_config()
                if st.button("Send Email"):
                    # Validate recipient before attempting to send
                    if not candidate_email:
                        st.error("Candidate email is required. Please enter a valid email address above.")
                    else:
                        html_body = st.session_state["candidate_email_html"]
                        plain_body = st.session_state.get("candidate_email_plain")

                        # Debug: verify HTML structure
                        is_valid_html = html_body.strip().startswith("<!DOCTYPE html>")
                        if not is_valid_html:
                            st.warning("Warning: Email body doesn't appear to be valid HTML. Please regenerate the template.")

                        ok = send_email_graph(
                            subject=f"Interview Opportunity at {company.name}: {role_title}",
                            body=html_body,
                            to_emails=[candidate_email],
                            content_type="HTML",
                            plain_text_body=plain_body,
                        )
                        audit.log(
                            "graph_sent_scheduling_email" if ok else "graph_send_failed",
                            actor=recruiter_email or "",
                            candidate_email=candidate_email or "",
                            hiring_manager_email=hiring_manager_email or "",
                            recruiter_email=recruiter_email or "",
                            role_title=role_title or "",
                            payload={
                                "subject": f"Interview Opportunity at {company.name}: {role_title}",
                            },
                            status="success" if ok else "failed",
                            error_message="" if ok else "Graph email send failed",
                        )
                        if ok:
                            st.success("Email sent.")
                        else:
                            st.error("Email send failed (see message above).")

            # Create Graph event
            # Collect panel interviewers from session state
            panel_interviewers_for_invite = [
                {"name": i.get("name", ""), "email": i.get("email", "")}
                for i in st.session_state.get("panel_interviewers", [])
                if i.get("email")  # Only include interviewers with valid emails
            ]

            # Determine if we have enough info to create invite
            has_interviewers = bool(panel_interviewers_for_invite) or bool(hiring_manager_email)
            has_valid_candidates = len(valid_candidates) > 0 and (proceed_with_valid or not invalid_count)
            create_disabled = not (selected_slot and has_interviewers and has_valid_candidates)

            # # Button text reflects number of candidates
            # button_text = "Create & Send Interview Invite"
            # if len(valid_candidates) > 1:
            #     button_text = f"Create & Send {len(valid_candidates)} Interview Invite(s)"
            #
            # # Validate Before Send button
            # col_validate, col_send = st.columns([1, 2])
            # with col_validate:
            #     if st.button("Validate Before Send", disabled=create_disabled, help="Preview who will receive invites"):
            #         report = _validate_invite_flow(
            #             selected_slot=selected_slot,
            #             tz_name=tz_name,
            #             candidate_timezone=candidate_timezone,
            #             duration_minutes=int(st.session_state["duration_minutes"]),
            #             candidates=valid_candidates,
            #             hiring_manager=(hiring_manager_email, hiring_manager_name),
            #             recruiter=(recruiter_email, recruiter_name),
            #             include_recruiter=include_recruiter,
            #             panel_interviewers=panel_interviewers_for_invite if panel_interviewers_for_invite else None,
            #             is_teams=is_teams,
            #             role_title=role_title,
            #         )
            #
            #         # Display validation report
            #         if report.is_valid:
            #             st.success(report.summary)
            #         else:
            #             st.error(report.summary)
            #
            #         # Show intended recipients
            #         if report.intended_recipients:
            #             st.markdown("**Invites will be sent to:**")
            #             for recipient in report.intended_recipients:
            #                 st.markdown(f"- {recipient}")
            #
            #         # Show errors
            #         if report.errors:
            #             st.markdown("**Errors (must fix):**")
            #             for err in report.errors:
            #                 st.error(err)
            #
            #         # Show warnings
            #         if report.warnings:
            #             st.markdown("**Warnings:**")
            #             for w in report.warnings:
            #                 st.warning(w)

            # with col_send:
            #     if st.button(button_text, disabled=create_disabled):
            #         with st.spinner(f"Scheduling {len(valid_candidates)} interview(s)..."):
            #             results = _handle_multi_candidate_invite(
            #                 audit=audit,
            #                 selected_slot=selected_slot,
            #                 tz_name=tz_name,
            #                 candidate_timezone=candidate_timezone,
            #                 duration_minutes=int(st.session_state["duration_minutes"]),
            #                 role_title=role_title,
            #                 subject=subject,
            #                 agenda=agenda,
            #                 location=location,
            #                 is_teams=is_teams,
            #                 candidates=valid_candidates,
            #                 hiring_manager=(hiring_manager_email, hiring_manager_name),
            #                 recruiter=(recruiter_email, recruiter_name),
            #                 include_recruiter=include_recruiter,
            #                 panel_interviewers=panel_interviewers_for_invite if panel_interviewers_for_invite else None,
            #                 scheduling_mode=scheduling_mode,
            #             )
            #
            #         # Display batch results
            #         _render_batch_results(results)

            # ICS fallback download button (available after generation)
            if st.session_state.get("last_invite_ics_bytes"):
                st.download_button(
                    "Download .ics (Add to calendar)",
                    data=st.session_state["last_invite_ics_bytes"],
                    file_name="powerdash_interview_invite.ics",
                    mime="text/calendar",
                )
                audit.log(
                    "ics_downloaded",
                    actor=recruiter_email or "",
                    candidate_email=candidate_email or "",
                    hiring_manager_email=hiring_manager_email or "",
                    recruiter_email=recruiter_email or "",
                    role_title=role_title or "",
                    event_id=st.session_state.get("last_graph_event_id", ""),
                    payload={"uid": st.session_state.get("last_invite_uid", "")},
                    status="success",
                )

                # Optional email ICS via Graph
                if st.button("Email .ics (optional)"):
                    ok = send_email_graph(
                        subject=subject,
                        body=agenda,
                        to_emails=[candidate_email, hiring_manager_email] + ([recruiter_email] if include_recruiter and recruiter_email else []),
                        attachment={
                            "data": st.session_state["last_invite_ics_bytes"],
                            "maintype": "text",
                            "subtype": "calendar",
                            "filename": "invite.ics",
                        },
                    )
                    audit.log(
                        "graph_sent_ics" if ok else "graph_send_failed",
                        actor=recruiter_email or "",
                        candidate_email=candidate_email or "",
                        hiring_manager_email=hiring_manager_email or "",
                        recruiter_email=recruiter_email or "",
                        role_title=role_title or "",
                        event_id=st.session_state.get("last_graph_event_id", ""),
                        payload={"uid": st.session_state.get("last_invite_uid", "")},
                        status="success" if ok else "failed",
                        error_message="" if ok else "Graph email send failed",
                    )
                    st.success("ICS emailed.") if ok else st.error("Failed to email ICS.")

    # ========= TAB: Scheduler Inbox =========
    with tab_inbox:
        # Header with refresh button
        header_col1, header_col2 = st.columns([3, 1])
        with header_col1:
            st.subheader("Scheduler Inbox")
        with header_col2:
            if st.button("Refresh Inbox", key="refresh_inbox", type="primary"):
                st.rerun()

        # Filter options
        col_filter1, col_filter2, col_filter3 = st.columns(3)
        with col_filter1:
            hide_bounces = st.checkbox("Hide bounce/undeliverable messages", value=True)
        with col_filter2:
            include_read = st.checkbox("Include already-read messages", value=False)
        with col_filter3:
            auto_send_invites = st.checkbox("Auto-send invites when slot detected", value=True, help="Automatically send calendar invites when a slot choice is detected from candidate replies")

        # Initialize session state for tracking auto-processed emails
        if "auto_processed_emails" not in st.session_state:
            st.session_state.auto_processed_emails = set()

        # Try IMAP first (Gmail), fall back to Graph API (Office 365)
        emails, error, is_configured = fetch_emails_imap(include_read=include_read)
        email_source = "Gmail (IMAP)"

        if not is_configured:
            # Fall back to Graph API
            emails, error, is_configured = fetch_unread_emails_graph(include_read=include_read)
            email_source = "Office 365 (Graph API)"

        # Get current time for "last fetched" display
        fetch_time = datetime.now().strftime("%I:%M:%S %p")

        if not is_configured:
            st.warning("Email inbox not configured. Add IMAP settings (imap_host, imap_username, imap_password) for Gmail, or Graph API settings for Office 365.")
        elif error:
            st.caption(f"Reading emails via {email_source} | Last fetched: {fetch_time}")
            st.error(f"Failed to fetch emails: {error}")
        elif not emails:
            st.caption(f"Reading emails via {email_source} | Last fetched: {fetch_time}")
            st.success("✓ Connected to mailbox. No unread emails found.")
        else:
            st.caption(f"Reading emails via {email_source} | Last fetched: {fetch_time}")
            # Filter out bounce messages if requested
            if hide_bounces:
                filtered_emails = [
                    e for e in emails
                    if not (
                        e.get("subject", "").lower().startswith("undeliverable") or
                        "postmaster" in e.get("from", "").lower() or
                        "mailer-daemon" in e.get("from", "").lower() or
                        "microsoftexchange" in e.get("from", "").lower()
                    )
                ]
                bounce_count = len(emails) - len(filtered_emails)
                if bounce_count > 0:
                    st.info(f"Hiding {bounce_count} bounce/undeliverable message(s)")
                emails = filtered_emails

            if not emails:
                st.success("✓ No candidate replies found (after filtering).")
            else:
                st.write(f"Found {len(emails)} email(s).")

            # Helper function to send invite for a detected slot
            def _send_invite_for_email(email_data: Dict[str, Any], detected_slot: Dict[str, str]) -> bool:
                """Send calendar invite for a detected slot choice. Returns True on success."""
                cand_email = email_data.get("from", "")

                # Get current form values from session state (keys match form input keys)
                hm_email = st.session_state.get("hm_email", "")
                hm_name = st.session_state.get("hm_name", "")
                rec_email = st.session_state.get("rec_email", "")
                rec_name = st.session_state.get("rec_name", "")
                role_title = st.session_state.get("role_title", "")
                duration = int(st.session_state.get("duration_minutes", 60))
                tz_name = st.session_state.get("tz_name", "UTC")
                candidate_tz = st.session_state.get("candidate_timezone", tz_name)
                is_teams = st.session_state.get("is_teams", True)
                subject = st.session_state.get("subject", "")
                agenda = st.session_state.get("agenda", "")
                location = st.session_state.get("location", "")
                include_recruiter = st.session_state.get("include_recruiter", True)
                panel_interviewers = st.session_state.get("panel_interviewers", [])

                # Validate we have minimum required info
                if not hm_email and not panel_interviewers:
                    st.error(f"Cannot auto-send to {cand_email}: Please set hiring manager email or panel interviewers in the New Scheduling Request tab first.")
                    return False

                # Parse candidate name from email if possible
                candidate_name = _ensure_candidate_name("", cand_email)

                # Create the invite
                result = _create_individual_invite(
                    audit=audit,
                    selected_slot=detected_slot,
                    tz_name=tz_name,
                    candidate_timezone=candidate_tz,
                    duration_minutes=duration,
                    role_title=role_title,
                    subject=subject,
                    agenda=agenda,
                    location=location,
                    is_teams=is_teams,
                    candidate=(cand_email, candidate_name),
                    hiring_manager=(hm_email, hm_name),
                    recruiter=(rec_email, rec_name),
                    include_recruiter=include_recruiter,
                    panel_interviewers=[
                        {"name": p.get("name", ""), "email": p.get("email", "")}
                        for p in panel_interviewers if p.get("email")
                    ] if panel_interviewers else None,
                )

                if result.success:
                    # Mark email as read after successful invite
                    email_id = email_data.get("id", "")
                    if email_id:
                        client = _make_graph_client()
                        if client:
                            client.mark_message_read(email_id)
                        mark_message_processed(email_id, action="invite_sent")
                    return True
                else:
                    st.error(f"Failed to send invite to {cand_email}: {result.error}")
                    return False

            # AUTO-SEND PROCESSING: Process emails with detected slots before rendering UI
            if auto_send_invites and emails:
                auto_send_results = []
                for email_item in emails:
                    email_id = email_item.get("id", "")
                    # Skip if already processed or already read
                    if email_id in st.session_state.auto_processed_emails:
                        continue
                    # Persistent idempotency (across reruns)
                    if has_processed_message(email_id, action="invite_sent"):
                        st.session_state.auto_processed_emails.add(email_id)
                        continue
                        continue
                    if email_item.get("is_read"):
                        continue

                    # Check for slot choice
                    full_body = email_item.get("full_body", email_item.get("body", ""))
                    detected_choice = detect_slot_choice_from_text(full_body, st.session_state.get("slots", []))

                    if detected_choice:
                        cand_email = email_item.get("from", "")
                        # Safety: ignore messages that originate from the scheduler mailbox itself
                        if cand_email and cand_email.strip().lower() == (SCHEDULER_MAILBOX or "").strip().lower():
                            st.session_state.auto_processed_emails.add(email_id)
                            continue
                        slot_display = f"{detected_choice.get('date')} {detected_choice.get('start')}-{detected_choice.get('end')}"

                        # Try to send the invite
                        with st.spinner(f"Auto-sending invite to {cand_email} for slot {slot_display}..."):
                            success = _send_invite_for_email(email_item, detected_choice)

                        if success:
                            auto_send_results.append((cand_email, slot_display, True))
                            # Mark as processed to avoid duplicate sends
                            st.session_state.auto_processed_emails.add(email_id)
                        else:
                            auto_send_results.append((cand_email, slot_display, False))

                # Show auto-send results
                if auto_send_results:
                    for cand, slot, success in auto_send_results:
                        if success:
                            st.success(f"Auto-sent invite to {cand} for slot {slot}")
                        # Errors already shown by _send_invite_for_email

            # Render email list UI
            for i, e in enumerate(emails, start=1):
                email_id = e.get("id", "")
                was_auto_processed = email_id in st.session_state.auto_processed_emails
                read_status = "[SENT]" if was_auto_processed else ("[READ]" if e.get("is_read") else "[NEW]")

                with st.expander(f"{i}. {read_status} {e['subject'] or '(no subject)'} — {e['from']}"):
                    st.write(e.get("date", ""))
                    preview_body = e.get("body", "")
                    full_body = e.get("full_body", preview_body)  # Fall back to preview if full not available
                    st.text_area("Body Preview", preview_body, height=160, key=f"inbox_body_{i}")

                    # Show full body info for debugging
                    slots_in_session = len(st.session_state.get("slots", []))
                    email_slots = _extract_slots_from_email_body(full_body)
                    slots_in_email = len(email_slots)
                    st.caption(f"Full body: {len(full_body):,} chars | Session slots: {slots_in_session} | Email slots: {slots_in_email}")

                    candidate_email = e.get("from", "")

                    # Try slot detection - works even if session slots are empty (extracts from email)
                    choice = detect_slot_choice_from_text(full_body, st.session_state.get("slots", []))
                    if choice:
                        st.success(f"Detected slot choice: {choice.get('date')} {choice.get('start')}-{choice.get('end')}")

                        # Check if already auto-processed
                        if was_auto_processed:
                            st.info(f"Invite already auto-sent to: {candidate_email}")
                        elif auto_send_invites:
                            # Already processed above, but email might be read or had an error
                            if e.get("is_read"):
                                st.info(f"Email already read - skipped auto-send. Use manual button if needed.")
                            # Show manual button as fallback
                            if st.button(f"Re-send Invite", key=f"resend_slot_{i}"):
                                with st.spinner("Sending calendar invite..."):
                                    if _send_invite_for_email(e, choice):
                                        st.success(f"Calendar invite sent to {candidate_email}!")
                                        st.session_state.auto_processed_emails.add(email_id)
                                        st.rerun()
                        else:
                            # Manual mode - show confirm button
                            col1, col2 = st.columns([2, 1])
                            with col1:
                                st.info(f"Ready to send invite to: {candidate_email}")
                            with col2:
                                if st.button(f"Confirm & Send Invite", key=f"confirm_slot_{i}"):
                                    with st.spinner("Sending calendar invite..."):
                                        if _send_invite_for_email(e, choice):
                                            st.success(f"Calendar invite sent to {candidate_email}!")
                                            st.rerun()
                    else:
                        st.info("No slot choice detected from this email.")

    # ========= TAB: Calendar Invites =========
    with tab_invites:
        st.subheader("Interview Management")
        st.caption("Manage scheduled interviews: reschedule, cancel, or view history.")

        def _format_candidates_display(interview_row: Dict[str, Any]) -> str:
            """Format candidate display for table, handling multi-candidate interviews."""

            candidates_json = interview_row.get("candidates_json")
            if candidates_json:
                try:
                    candidates = json.loads(candidates_json)
                    if len(candidates) > 2:
                        return f"{candidates[0]['email']} +{len(candidates)-1} more"
                    return "; ".join(c['email'] for c in candidates)
                except (json.JSONDecodeError, KeyError, TypeError):
                    pass
            return interview_row.get("candidate_email", "")

        def _format_interview_type(interview_row: Dict[str, Any]) -> str:
            """Format interview type indicator."""
            is_group = interview_row.get("is_group_interview")
            is_panel = interview_row.get("is_panel_interview")
            if is_group:
                return "Group"
            elif is_panel:
                return "Panel"
            return "Individual"

        def _get_status_badge(status: str) -> str:
            """Return status with emoji badge."""
            status_badges = {
                "pending": "Pending",
                "confirmed": "Confirmed",
                "rescheduled": "Rescheduled",
                "cancelled": "Cancelled",
                "completed": "Completed",
                "no_show": "No Show",
                "created": "Created",
                "scheduled": "Scheduled",
            }
            return status_badges.get(status.lower() if status else "", status or "Unknown")

        # --- Audit Log Rendering Functions ---

        def _render_audit_timeline(entries: List[Dict[str, Any]]):
            """Render audit log as timeline view with status indicators."""
            if not entries:
                st.info("No audit entries to display.")
                return

            # Action color mapping
            action_colors = {
                "graph_create_event": "#28a745",      # Green - success
                "graph_create_group_event": "#28a745",
                "interview_rescheduled": "#ffc107",   # Yellow - change
                "graph_reschedule_event": "#ffc107",
                "interview_cancelled": "#dc3545",     # Red - cancellation
                "graph_cancel_event": "#dc3545",
                "notification_sent": "#17a2b8",       # Blue - notification
                "candidate_notification_sent": "#17a2b8",
                "email_sent": "#17a2b8",
                "graph_sent_scheduling_email": "#17a2b8",
            }
            default_color = "#6c757d"  # Gray

            current_date = None

            for entry in entries:
                # Date separator
                timestamp = entry.get("timestamp", "")
                entry_date = timestamp[:12] if timestamp else ""

                if entry_date and entry_date != current_date:
                    current_date = entry_date
                    st.markdown(f"### {current_date}")

                # Get color for this action
                action_code = entry.get("action_code", "")
                color = action_colors.get(action_code, default_color)

                # Status icon
                status = entry.get("status", "")
                if status == "success":
                    status_icon = "✅"
                elif status == "failed":
                    status_icon = "❌"
                else:
                    status_icon = "ℹ️"

                # Entry card
                with st.container():
                    col_status, col_content = st.columns([0.5, 11.5])

                    with col_status:
                        st.markdown(f"### {status_icon}")

                    with col_content:
                        st.markdown(f"**{entry.get('action_display', '')}**")
                        st.markdown(entry.get("summary", ""))

                        # Meta line
                        time_part = timestamp[13:] if len(timestamp) > 13 else ""
                        meta = f"_{time_part} by {entry.get('actor_display', 'System')}_"
                        st.caption(meta)

                        # Details if present
                        if entry.get("details"):
                            st.caption(entry["details"])

                        # Expandable raw data
                        with st.expander("View details"):
                            st.json(entry.get("raw", {}))

                    st.markdown("---")

        def _render_audit_table(entries: List[Dict[str, Any]]):
            """Render audit log as formatted table view."""
            if not entries:
                st.info("No audit entries to display.")
                return

            # Status icon mapping
            def get_status_icon(status: str) -> str:
                if status == "success":
                    return "✅"
                elif status == "failed":
                    return "❌"
                return "ℹ️"

            table_data = [
                {
                    "Time": e.get("timestamp", ""),
                    "Status": get_status_icon(e.get("status", "")),
                    "Action": e.get("action_display", ""),
                    "Summary": e.get("summary", ""),
                    "Actor": e.get("actor_display", ""),
                }
                for e in entries
            ]

            st.dataframe(
                table_data,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Status": st.column_config.TextColumn(width="small"),
                    "Summary": st.column_config.TextColumn(width="large"),
                }
            )

            # Expandable row details
            with st.expander("View entry details"):
                idx = st.number_input(
                    "Entry index (0 = most recent)",
                    min_value=0,
                    max_value=max(0, len(entries) - 1),
                    value=0,
                    key="audit_table_row_idx"
                )
                selected = entries[int(idx)]
                st.markdown(f"**{selected.get('action_display', '')}**")
                st.markdown(selected.get("summary", ""))
                if selected.get("details"):
                    st.caption(selected["details"])
                st.json(selected.get("raw", {}))

        def _render_audit_raw(entries: List[Dict[str, Any]]):
            """Render audit log as raw dataframe (original view)."""
            if not entries:
                st.info("No audit entries to display.")
                return

            st.dataframe(
                [
                    {
                        "timestamp_utc": r.get("timestamp_utc", ""),
                        "action": r.get("action", ""),
                        "status": r.get("status", ""),
                        "candidate": r.get("candidate_email", ""),
                        "hiring_manager": r.get("hiring_manager_email", ""),
                        "event_id": r.get("event_id", ""),
                        "error": (r.get("error_message", "")[:80] + "…")
                            if r.get("error_message") and len(r.get("error_message", "")) > 80
                            else (r.get("error_message") or ""),
                    }
                    for r in entries
                ],
                use_container_width=True,
                hide_index=True,
            )

            with st.expander("Show raw payload for a row"):
                idx = st.number_input(
                    "Row index (0 = most recent)",
                    min_value=0,
                    max_value=max(0, len(entries) - 1),
                    value=0,
                    key="audit_raw_row_idx"
                )
                st.json(entries[int(idx)])

        def _render_interview_history(audit_instance: AuditLog, event_id: str):
            """Display complete history for selected interview."""
            st.markdown("#### Interview History")

            history = audit_instance.get_interview_history(event_id)

            if not history:
                st.info("No history entries for this interview.")
                return

            # Action color mapping
            action_colors = {
                "graph_create_event": "#28a745",     # Green
                "interview_rescheduled": "#ffc107", # Yellow
                "interview_cancelled": "#dc3545",   # Red
                "notification_sent": "#17a2b8",      # Blue
            }

            for entry in history:
                action = entry.get("action", "")
                timestamp = entry.get("timestamp_utc", "")[:16]
                actor = entry.get("actor", "System") or "System"
                status = entry.get("status", "")

                color = action_colors.get(action, "#6c757d")

                col_time, col_action, col_status = st.columns([2, 4, 2])
                with col_time:
                    st.caption(timestamp)
                with col_action:
                    st.markdown(f"<span style='color: {color}; font-weight: 600;'>{action}</span>", unsafe_allow_html=True)
                with col_status:
                    st.caption(f"{status} | {actor[:20]}")

                # Show details in expander
                payload_json = entry.get("payload_json")
                if payload_json:
                    with st.expander("Details", expanded=False):
                        try:
                            st.json(json.loads(payload_json))
                        except json.JSONDecodeError:
                            st.text(payload_json)

                st.markdown("---")

        # Filter controls
        col_filter1, col_filter2 = st.columns(2)
        with col_filter1:
            status_filter = st.selectbox(
                "Filter by status",
                options=["All", "Pending", "Confirmed", "Rescheduled", "Cancelled", "Scheduled", "Created"],
                key="invites_status_filter"
            )
        with col_filter2:
            search_term = st.text_input("Search candidate/role", key="invites_search", placeholder="Search...")

        # Get interviews with optional filter
        filter_value = None if status_filter == "All" else status_filter.lower()
        interviews = audit.list_interviews(limit=200, status_filter=filter_value)

        # Apply search filter
        if search_term:
            search_lower = search_term.lower()
            interviews = [
                r for r in interviews
                if search_lower in (r.get("candidate_email", "") or "").lower()
                or search_lower in (r.get("role_title", "") or "").lower()
            ]

        if not interviews:
            st.info("No interviews match the current filters. Create an invite from the first tab.")
        else:
            # Show compact table
            st.dataframe(
                [
                    {
                        "created": r["created_utc"][:10] if r.get("created_utc") else "",
                        "role": r.get("role_title", ""),
                        "type": _format_interview_type(r),
                        "candidate(s)": _format_candidates_display(r),
                        "start_utc": r.get("start_utc", "")[:16] if r.get("start_utc") else "",
                        "status": _get_status_badge(r.get("last_status", "")),
                        "event_id": (r.get("graph_event_id", "") or "")[:12] + "..." if r.get("graph_event_id") else "",
                    }
                    for r in interviews
                ],
                use_container_width=True,
                hide_index=True,
            )

            # Export controls section
            with st.expander("Export Data", expanded=False):
                col_tz, col_fields, col_status = st.columns(3)

                with col_tz:
                    export_tz = st.selectbox(
                        "Timezone",
                        options=_common_timezones(),
                        index=_tz_index(get_default_timezone()),
                        key="export_interviews_tz"
                    )

                with col_fields:
                    include_all_fields = st.checkbox(
                        "Include all fields",
                        value=False,
                        help="Include extended fields like Teams links, Event IDs",
                        key="export_all_fields"
                    )

                with col_status:
                    export_status_filter = st.multiselect(
                        "Filter by status",
                        options=["Pending", "Confirmed", "Rescheduled", "Cancelled", "Completed"],
                        default=["Pending", "Confirmed", "Rescheduled"],
                        key="export_status_filter"
                    )

                col_date_range, col_custom = st.columns(2)

                with col_date_range:
                    date_range = st.selectbox(
                        "Date range",
                        options=["All time", "Today", "This week", "This month", "Last 30 days", "Custom"],
                        key="export_date_range"
                    )

                date_from = None
                date_to = None
                if date_range == "Custom":
                    with col_custom:
                        col_from, col_to = st.columns(2)
                        with col_from:
                            date_from = st.date_input("From", key="export_date_from")
                        with col_to:
                            date_to = st.date_input("To", key="export_date_to")

                # Filter and generate export
                filtered_interviews = filter_interviews_for_export(
                    interviews,
                    status_filter=export_status_filter if export_status_filter else None,
                    date_range=date_range,
                    date_from=date_from,
                    date_to=date_to,
                )

                col_info, col_download = st.columns([3, 1])

                with col_info:
                    st.caption(f"{len(filtered_interviews)} interview(s) match the filters")

                with col_download:
                    if filtered_interviews:
                        csv_bytes = export_interviews_csv(
                            filtered_interviews,
                            display_timezone=export_tz,
                            include_all_fields=include_all_fields,
                        )
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        st.download_button(
                            label="Download CSV",
                            data=csv_bytes,
                            file_name=f"interviews_export_{timestamp}.csv",
                            mime="text/csv",
                            key="download_interviews_csv"
                        )
                    else:
                        st.button("Download CSV", disabled=True, key="download_disabled")

            st.markdown("----")

            # Check for active confirmation dialogs
            cancelling_id = st.session_state.get("cancelling_interview_id")
            rescheduling_id = st.session_state.get("rescheduling_interview_id")
            viewing_history_id = st.session_state.get("viewing_interview_history")

            # Interview History View
            if viewing_history_id:
                st.markdown("### Interview History")
                if st.button("Back to Interview List", key="back_from_history"):
                    st.session_state["viewing_interview_history"] = None
                    st.rerun()
                _render_interview_history(audit, viewing_history_id)

            # Cancellation Confirmation Dialog
            elif cancelling_id:
                cancel_row = next((r for r in interviews if r.get("graph_event_id") == cancelling_id), None)
                if cancel_row:
                    st.markdown("### Confirm Cancellation")
                    st.warning(f"""
                    **You are about to cancel this interview:**
                    - **Role:** {cancel_row.get('role_title', '')}
                    - **Candidate:** {cancel_row.get('candidate_email', '')}
                    - **Time:** {cancel_row.get('start_utc', '')}

                    This action cannot be undone. The candidate will receive a cancellation notice.
                    """)

                    cancel_reason = st.selectbox(
                        "Cancellation reason",
                        options=[
                            "Candidate requested",
                            "Position filled",
                            "Position closed",
                            "Interviewer unavailable",
                            "Scheduling conflict",
                            "Other"
                        ],
                        key="cancel_reason_select"
                    )

                    custom_reason = ""
                    if cancel_reason == "Other":
                        custom_reason = st.text_input("Please specify reason", key="cancel_custom_reason")

                    notify_candidate = st.checkbox("Send cancellation email to candidate", value=True, key="cancel_notify")
                    candidate_message = ""
                    if notify_candidate:
                        candidate_message = st.text_area(
                            "Message to candidate (optional)",
                            placeholder="We apologize for any inconvenience...",
                            key="cancel_message"
                        )

                    col_confirm, col_back = st.columns(2)
                    with col_confirm:
                        if st.button("Confirm Cancellation", type="primary", key="confirm_cancel_btn"):
                            final_reason = custom_reason if cancel_reason == "Other" else cancel_reason
                            _handle_cancel(
                                audit=audit,
                                event_id=cancelling_id,
                                context_row=cancel_row,
                                reason=final_reason,
                                notify_candidate=notify_candidate,
                                candidate_message=candidate_message,
                            )
                            st.session_state["cancelling_interview_id"] = None
                            st.rerun()
                    with col_back:
                        if st.button("Go Back", key="cancel_back_btn"):
                            st.session_state["cancelling_interview_id"] = None
                            st.rerun()
                else:
                    st.session_state["cancelling_interview_id"] = None
                    st.rerun()

            # Reschedule Confirmation Dialog
            elif rescheduling_id:
                resched_row = next((r for r in interviews if r.get("graph_event_id") == rescheduling_id), None)
                if resched_row:
                    st.markdown("### Reschedule Interview")

                    # Display current info
                    st.info(f"""
                    **Rescheduling interview:**
                    - **Role:** {resched_row.get('role_title', '')}
                    - **Candidate:** {resched_row.get('candidate_email', '')}
                    - **Current Time:** {resched_row.get('start_utc', '')}
                    """)

                    display_tz = st.selectbox(
                        "Timezone",
                        options=_common_timezones(),
                        index=_tz_index(resched_row.get("display_timezone")),
                        key="resched_tz"
                    )

                    try:
                        start_local = from_utc(datetime.fromisoformat(resched_row["start_utc"]), display_tz)
                    except Exception:
                        start_local = None

                    col_date, col_time = st.columns(2)
                    with col_date:
                        new_date = st.date_input(
                            "New date",
                            value=start_local.date() if start_local else datetime.now().date(),
                            key="resched_date"
                        )
                    with col_time:
                        new_time = st.time_input(
                            "New time",
                            value=start_local.time().replace(second=0, microsecond=0) if start_local else datetime.now().time().replace(second=0, microsecond=0),
                            key="resched_time"
                        )

                    new_duration = st.number_input(
                        "Duration (minutes)",
                        min_value=15,
                        max_value=240,
                        step=15,
                        value=int(resched_row.get("duration_minutes") or 30),
                        key="resched_duration"
                    )

                    reschedule_reason = st.selectbox(
                        "Reason for reschedule",
                        options=[
                            "Candidate requested",
                            "Interviewer unavailable",
                            "Scheduling conflict",
                            "Time zone adjustment",
                            "Other"
                        ],
                        key="reschedule_reason_select"
                    )

                    notify_candidate = st.checkbox(
                        "Send update notification to candidate",
                        value=True,
                        key="reschedule_notify"
                    )

                    col_confirm, col_back = st.columns(2)
                    with col_confirm:
                        if st.button("Confirm Reschedule", type="primary", key="confirm_resched_btn"):
                            _handle_reschedule(
                                audit=audit,
                                event_id=rescheduling_id,
                                new_date=new_date,
                                new_time=new_time,
                                duration_minutes=int(new_duration),
                                tz_name=display_tz,
                                context_row=resched_row,
                                reason=reschedule_reason,
                                notify_candidate=notify_candidate,
                            )
                            st.session_state["rescheduling_interview_id"] = None
                            st.rerun()
                    with col_back:
                        if st.button("Go Back", key="resched_back_btn"):
                            st.session_state["rescheduling_interview_id"] = None
                            st.rerun()
                else:
                    st.session_state["rescheduling_interview_id"] = None
                    st.rerun()

            # Default view - Interview management
            else:
                st.markdown("#### Manage Interview")

                # Only show non-cancelled interviews for management
                manageable = [r for r in interviews if r.get("last_status", "").lower() != "cancelled"]
                if not manageable:
                    st.info("No active interviews to manage. All interviews are cancelled or none exist.")
                else:
                    event_ids = [r["graph_event_id"] for r in manageable if r.get("graph_event_id")]
                    if event_ids:
                        selected_event_id = st.selectbox(
                            "Select interview",
                            options=event_ids,
                            format_func=lambda x: next(
                                (f"{r.get('role_title', '')} - {r.get('candidate_email', '')} ({r.get('start_utc', '')[:10]})"
                                 for r in manageable if r.get("graph_event_id") == x),
                                x
                            ),
                            key="manage_event_select"
                        )
                        selected_row = next((r for r in manageable if r.get("graph_event_id") == selected_event_id), None)

                        if selected_row:
                            # Show interview details
                            st.markdown(f"""
                            **Selected Interview:**
                            - **Role:** {selected_row.get('role_title', '')}
                            - **Candidate:** {selected_row.get('candidate_email', '')}
                            - **Start:** {selected_row.get('start_utc', '')}
                            - **Status:** {_get_status_badge(selected_row.get('last_status', ''))}
                            """)

                            # Action buttons
                            col_resched, col_cancel, col_history = st.columns(3)
                            with col_resched:
                                if st.button("Reschedule", type="primary", key="btn_reschedule"):
                                    st.session_state["rescheduling_interview_id"] = selected_event_id
                                    st.rerun()
                            with col_cancel:
                                if st.button("Cancel", type="secondary", key="btn_cancel"):
                                    st.session_state["cancelling_interview_id"] = selected_event_id
                                    st.rerun()
                            with col_history:
                                if st.button("View History", key="btn_history"):
                                    st.session_state["viewing_interview_history"] = selected_event_id
                                    st.rerun()

    # ========= TAB: Audit Log =========
    with tab_audit:
        st.subheader("Audit Log")
        st.caption("Complete history of all scheduling actions.")

        # Controls row
        col_view, col_limit = st.columns([3, 1])

        with col_view:
            view_mode = st.radio(
                "View",
                options=["Table", "Timeline", "Raw"],
                horizontal=True,
                key="audit_view_mode"
            )

        with col_limit:
            entry_limit = st.selectbox(
                "Entries",
                options=[100, 300, 500],
                index=1,
                key="audit_entry_limit"
            )

        # Filter controls
        col_action, col_status, col_search = st.columns([2, 1, 2])

        with col_action:
            action_options = ["All"] + sorted(set(AUDIT_ACTION_DESCRIPTIONS.values()))
            action_filter = st.selectbox(
                "Filter by action",
                options=action_options,
                key="audit_action_filter"
            )

        with col_status:
            status_filter = st.selectbox(
                "Status",
                options=["All", "Success", "Failed"],
                key="audit_status_filter"
            )

        with col_search:
            audit_search = st.text_input(
                "Search",
                placeholder="candidate, role...",
                key="audit_search"
            )

        # Fetch entries
        raw_entries = audit.list_recent_audit(limit=entry_limit)

        if not raw_entries:
            st.info("No audit entries yet.")
        else:
            # Apply filters
            filtered_entries = filter_audit_entries(
                raw_entries,
                action_filter=action_filter,
                status_filter=status_filter,
                search_term=audit_search,
            )

            if not filtered_entries:
                st.warning("No entries match the current filters.")
            else:
                # Format entries for display
                formatted_entries = [format_audit_entry_human(e) for e in filtered_entries]

                # Export button
                col_count, col_export = st.columns([3, 1])

                with col_count:
                    st.caption(f"Showing {len(formatted_entries)} of {len(raw_entries)} entries")

                with col_export:
                    csv_bytes = export_audit_log_csv(formatted_entries)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.download_button(
                        label="Export CSV",
                        data=csv_bytes,
                        file_name=f"audit_log_{timestamp}.csv",
                        mime="text/csv",
                        key="download_audit_csv"
                    )

                st.markdown("---")

                # Render based on view mode
                if view_mode == "Timeline":
                    _render_audit_timeline(formatted_entries)
                elif view_mode == "Table":
                    _render_audit_table(formatted_entries)
                else:
                    _render_audit_raw(filtered_entries)

    # Render footer if enabled
    if layout.show_footer:
        _render_footer()


# ----------------------------
# Internal UI handlers
# ----------------------------
def _parse_availability_upload(upload, interviewer_timezone: Optional[str] = None, display_timezone: Optional[str] = None) -> List[Dict[str, str]]:
    data = upload.read()
    name = (upload.name or "").lower()
    slots: List[Dict[str, str]] = []

    if name.endswith(".pdf"):
        imgs = pdf_to_images(data, max_pages=3)
        for img in imgs:
            slots.extend(parse_slots_from_image(img, interviewer_timezone, display_timezone))

    elif name.endswith(".docx"):
        # Strategy: Extract text + tables, then also check embedded images

        # 1. Parse text content (paragraphs + tables)
        text = docx_to_text(data)
        if text:
            slots.extend(parse_slots_from_text(text))

        # 2. Extract and parse embedded images (optional enhancement)
        embedded_images = docx_extract_images(data, max_images=3)
        for img in embedded_images:
            slots.extend(parse_slots_from_image(img, interviewer_timezone, display_timezone))

    else:
        # Assume image file (png, jpg, jpeg)
        img = Image.open(io.BytesIO(data)).convert("RGB")
        slots.extend(parse_slots_from_image(img, interviewer_timezone, display_timezone))

    # De-duplicate slots by (date, start, end) tuple
    uniq = {(s["date"], s["start"], s["end"]): s for s in slots}
    return list(uniq.values())

def _parse_single_interviewer_availability(interviewer_idx: int) -> None:
    """Parse availability for a single interviewer and recompute intersection."""
    from slot_intersection import (
        normalize_slots_to_utc,
        merge_adjacent_slots,
        compute_intersection,
    )

    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.error("Invalid interviewer index")
        return

    tz_name = st.session_state["selected_timezone"]
    min_duration = st.session_state["duration_minutes"]
    interviewer = interviewers[interviewer_idx]
    interviewer_name = interviewer.get("name") or f"Interviewer {interviewer.get('id', '?')}"

    # Parse this interviewer's file
    interviewer_tz = interviewer.get("timezone", tz_name)
    existing_manual_slots = [s for s in interviewer.get("slots", []) if s.get("source") == "manual"]

    try:
        if interviewer.get("file"):
            interviewer["file"].seek(0)
            uploaded_slots = _parse_availability_upload(interviewer["file"], interviewer_tz, tz_name)
            for s in uploaded_slots:
                s["source"] = "uploaded"
            interviewer["slots"] = _merge_slots(existing_manual_slots, uploaded_slots)
            uploaded_count = len(uploaded_slots)
            st.success(f"Parsed {uploaded_count} slot(s) for {interviewer_name}.")
        elif existing_manual_slots:
            interviewer["slots"] = existing_manual_slots
            st.info(f"No file uploaded for {interviewer_name}. Using {len(existing_manual_slots)} manual slot(s).")
        else:
            st.warning(f"No file or manual slots for {interviewer_name}.")
            return
    except Exception as e:
        st.error(f"{interviewer_name}: {e}")
        return

    st.session_state["panel_interviewers"][interviewer_idx] = interviewer

    # Recompute intersection across all interviewers
    all_interviewer_slots: Dict[int, List] = {}
    interviewer_names_map: Dict[int, str] = {}

    for iv in interviewers:
        if iv.get("slots"):
            name = iv.get("name") or iv.get("email") or f"Interviewer {iv['id']}"
            interviewer_names_map[iv["id"]] = name
            normalized = normalize_slots_to_utc(iv["slots"], tz_name)
            merged = merge_adjacent_slots(normalized)
            all_interviewer_slots[iv["id"]] = merged

    if all_interviewer_slots:
        intersections = compute_intersection(
            all_interviewer_slots,
            min_duration_minutes=min_duration,
            display_timezone=tz_name,
            interviewer_names=interviewer_names_map,
        )

        # Split continuous windows into discrete meeting-sized slots
        split_intersections = []
        for slot in intersections:
            split_slots = split_slot_by_duration(slot, min_duration)
            if split_slots:
                # Preserve metadata from original slot in each split slot
                for ss in split_slots:
                    for key in slot:
                        if key not in ss:
                            ss[key] = slot[key]
                split_intersections.extend(split_slots)
            else:
                # Keep original if splitting returns empty (edge case)
                split_intersections.append(slot)

        st.session_state["computed_intersections"] = split_intersections
        st.session_state["slots"] = split_intersections
        _save_persisted_slots()


def _parse_all_panel_availability() -> None:
    """Parse availability for all interviewers and compute intersection.

    Handles both uploaded files and manually-entered slots.
    """
    from slot_intersection import (
        normalize_slots_to_utc,
        merge_adjacent_slots,
        compute_intersection,
    )

    interviewers = st.session_state.get("panel_interviewers", [])
    tz_name = st.session_state["selected_timezone"]
    min_duration = st.session_state["duration_minutes"]

    all_interviewer_slots: Dict[int, List] = {}
    interviewer_names: Dict[int, str] = {}
    parse_errors = []
    total_uploaded = 0
    total_manual = 0

    for interviewer in interviewers:
        # Get existing manual slots (preserve them)
        existing_manual_slots = [s for s in interviewer.get("slots", []) if s.get("source") == "manual"]
        total_manual += len(existing_manual_slots)
        interviewer_tz = interviewer.get("timezone", tz_name)

        try:
            if interviewer.get("file"):
                # Reset file position before reading
                interviewer["file"].seek(0)
                # Parse the uploaded file
                uploaded_slots = _parse_availability_upload(interviewer["file"], interviewer_tz, tz_name)
                # Mark uploaded slots with source
                for s in uploaded_slots:
                    s["source"] = "uploaded"
                total_uploaded += len(uploaded_slots)
                # Merge manual + uploaded, preferring manual for duplicates
                interviewer["slots"] = _merge_slots(existing_manual_slots, uploaded_slots)
            elif existing_manual_slots:
                # No file but has manual slots - keep them
                interviewer["slots"] = existing_manual_slots

            # Include interviewer if they have any slots
            if interviewer.get("slots"):
                # Build interviewer name for display
                name = interviewer.get("name") or interviewer.get("email") or f"Interviewer {interviewer['id']}"
                interviewer_names[interviewer["id"]] = name

                # Normalize to UTC for intersection
                normalized = normalize_slots_to_utc(interviewer["slots"], tz_name)
                merged = merge_adjacent_slots(normalized)
                all_interviewer_slots[interviewer["id"]] = merged

        except Exception as e:
            interviewer_name = interviewer.get("name") or f"Interviewer {interviewer.get('id', '?')}"
            parse_errors.append(f"{interviewer_name}: {e}")

    if parse_errors:
        for err in parse_errors:
            st.error(err)

    # Compute intersection
    if all_interviewer_slots:
        intersections = compute_intersection(
            all_interviewer_slots,
            min_duration_minutes=min_duration,
            display_timezone=tz_name,
            interviewer_names=interviewer_names,
        )

        # Split continuous windows into discrete meeting-sized slots
        split_intersections = []
        for slot in intersections:
            split_slots = split_slot_by_duration(slot, min_duration)
            if split_slots:
                # Preserve metadata from original slot in each split slot
                for ss in split_slots:
                    for key in slot:
                        if key not in ss:
                            ss[key] = slot[key]
                split_intersections.extend(split_slots)
            else:
                # Keep original if splitting returns empty (edge case)
                split_intersections.append(slot)

        st.session_state["computed_intersections"] = split_intersections

        # Also update legacy "slots" for backward compatibility
        st.session_state["slots"] = split_intersections
        _save_persisted_slots()

        num_interviewers = len(all_interviewer_slots)
        total_slots = total_uploaded + total_manual

        if num_interviewers == 1:
            if total_manual > 0 and total_uploaded > 0:
                st.success(f"Processed {total_slots} slot(s) ({total_manual} manual, {total_uploaded} uploaded).")
            elif total_manual > 0:
                st.success(f"Processed {total_manual} manual slot(s).")
            else:
                st.success(f"Extracted {total_uploaded} slot(s) from uploaded file.")
        else:
            full_overlap = sum(1 for s in intersections if s.get("is_full_overlap", False))
            source_info = []
            if total_manual > 0:
                source_info.append(f"{total_manual} manual")
            if total_uploaded > 0:
                source_info.append(f"{total_uploaded} uploaded")
            source_str = f" ({', '.join(source_info)})" if source_info else ""
            st.success(
                f"Processed {total_slots} total slots{source_str} from {num_interviewers} interviewers. "
                f"Found {len(intersections)} intersection slot(s) ({full_overlap} with all available)."
            )
    else:
        st.warning("No availability found. Please upload calendars or add slots manually.")


def _zoneinfo(tz_name: str):
    """Get ZoneInfo with validation. Falls back to UTC if invalid."""
    zi, was_valid = safe_zoneinfo(tz_name, fallback="UTC")
    if not was_valid:
        st.warning(f"Invalid timezone '{tz_name}', using UTC")
    return zi


def _common_timezones() -> List[str]:
    # Keep concise; you can expand later.
    return [
        "UTC",
        "Europe/London",
        "Europe/Dublin",
        "Europe/Paris",
        "Europe/Rome",
        "Europe/Berlin",
        "America/New_York",
        "America/Chicago",
        "America/Denver",
        "America/Los_Angeles",
        "America/Toronto",
        "America/Sao_Paulo",
        "Asia/Dubai",
        "Asia/Kolkata",
        "Asia/Singapore",
        "Asia/Tokyo",
        "Australia/Sydney",
    ]


def _tz_index(tz_name: str | None) -> int:
    tzs = _common_timezones()
    if tz_name and tz_name in tzs:
        return tzs.index(tz_name)
    return tzs.index(get_default_timezone()) if get_default_timezone() in tzs else 0


def _render_batch_results(results: List[SchedulingResult]) -> None:
    """Display results of batch scheduling operation."""
    if not results:
        return

    successful = [r for r in results if r.success]
    failed = [r for r in results if not r.success]

    # Summary
    st.markdown("### Scheduling Results")

    if successful and not failed:
        st.success(f"Successfully scheduled all {len(successful)} interview(s)")
    elif successful and failed:
        st.warning(f"Scheduled {len(successful)} interview(s), {len(failed)} failed")
    else:
        st.error(f"Failed to schedule all {len(failed)} interview(s)")

    # Success details
    if successful:
        with st.expander(f"Successful ({len(successful)})", expanded=True):
            for r in successful:
                display = f"{r.candidate_name} ({r.candidate_email})" if r.candidate_name else r.candidate_email
                st.markdown(f":white_check_mark: **{display}**")

                # Show recipients
                if r.recipients:
                    st.caption(f"Invites sent to: {', '.join(r.recipients)}")

                # Show warnings
                if r.warnings:
                    for w in r.warnings:
                        st.warning(w, icon="warning")

                if r.teams_url:
                    st.link_button("Open Teams Link", r.teams_url)

    # Failure details
    if failed:
        with st.expander(f"Failed ({len(failed)})", expanded=True):
            for r in failed:
                display = f"{r.candidate_name} ({r.candidate_email})" if r.candidate_name else r.candidate_email
                st.markdown(f":x: **{display}**")
                st.caption(f"Error: {r.error}")


def _handle_multi_candidate_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidates: List[CandidateValidationResult],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
    scheduling_mode: str = "individual",
) -> List[SchedulingResult]:
    """
    Handle creating invites for multiple candidates.
    Returns list of SchedulingResult for UI display.
    """
    results: List[SchedulingResult] = []

    if scheduling_mode == "group" and len(candidates) > 1:
        # Group interview: single invite with all candidates
        result = _create_group_invite(
            audit=audit,
            selected_slot=selected_slot,
            tz_name=tz_name,
            candidate_timezone=candidate_timezone,
            duration_minutes=duration_minutes,
            role_title=role_title,
            subject=subject,
            agenda=agenda,
            location=location,
            is_teams=is_teams,
            candidates=candidates,
            hiring_manager=hiring_manager,
            recruiter=recruiter,
            include_recruiter=include_recruiter,
            panel_interviewers=panel_interviewers,
        )
        results.append(result)
    else:
        # Individual interviews: one invite per candidate
        for candidate in candidates:
            if not candidate.is_valid:
                results.append(SchedulingResult(
                    candidate_email=candidate.original,
                    candidate_name=candidate.name,
                    success=False,
                    event_id=None,
                    teams_url=None,
                    error=candidate.error or "Invalid candidate"
                ))
                continue

            result = _create_individual_invite(
                audit=audit,
                selected_slot=selected_slot,
                tz_name=tz_name,
                candidate_timezone=candidate_timezone,
                duration_minutes=duration_minutes,
                role_title=role_title,
                subject=subject,
                agenda=agenda,
                location=location,
                is_teams=is_teams,
                candidate=(candidate.email, candidate.name),
                hiring_manager=hiring_manager,
                recruiter=recruiter,
                include_recruiter=include_recruiter,
                panel_interviewers=panel_interviewers,
            )
            results.append(result)

    return results


@dataclass
class ValidationReport:
    """Result of validating invite parameters without sending."""
    is_valid: bool
    intended_recipients: List[str]
    errors: List[str]
    warnings: List[str]
    summary: str


def _validate_invite_flow(
    *,
    selected_slot: Optional[Dict[str, str]],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    candidates: List[CandidateValidationResult],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
    is_teams: bool = False,
) -> ValidationReport:
    """
    Dry-run validation of invite parameters without sending.
    Returns a validation report with all intended recipients, errors, and warnings.
    """
    errors: List[str] = []
    warnings: List[str] = []
    intended_recipients: List[str] = []

    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # Validate role title
    if not role_title or not role_title.strip():
        warnings.append("Role/position title is empty - subject will use generic 'Interview' text")

    # Validate timezone
    if not is_valid_timezone(tz_name):
        warnings.append(f"Display timezone '{tz_name}' is invalid, will use UTC")

    if not is_valid_timezone(candidate_timezone):
        warnings.append(f"Candidate timezone '{candidate_timezone}' is invalid, will use display timezone")

    # Validate candidates
    valid_candidates = [c for c in candidates if c.is_valid and c.email]
    invalid_candidates = [c for c in candidates if not c.is_valid]

    if not valid_candidates:
        errors.append("No valid candidate emails provided")
    else:
        for c in valid_candidates:
            intended_recipients.append(c.email)

    for c in invalid_candidates:
        warnings.append(f"Invalid candidate email skipped: {c.original} - {c.error}")

    # Validate hiring manager
    try:
        if panel_interviewers:
            hm_email = validate_email_optional(hm_email_raw, "Hiring manager email")
        else:
            hm_email = validate_email(hm_email_raw, "Hiring manager email")
        if hm_email:
            intended_recipients.append(hm_email)
    except ValidationError as e:
        errors.append(str(e))
        hm_email = None

    # Validate recruiter
    rec_email = None
    if include_recruiter and rec_email_raw:
        try:
            rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
            if rec_email:
                intended_recipients.append(rec_email)
        except ValidationError as e:
            warnings.append(f"Invalid recruiter email: {e}")

    # Validate panel interviewers
    if panel_interviewers:
        seen_emails = {(c.email or "").lower() for c in valid_candidates}
        if hm_email:
            seen_emails.add(hm_email.lower())

        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            pi_name = pi.get("name", "")
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    intended_recipients.append(validated_email)
                    seen_emails.add(validated_email.lower())
                except ValidationError as ve:
                    warnings.append(f"Invalid panel email '{pi_email}': {ve.message}")

    # Validate slot
    if not selected_slot:
        errors.append("No time slot selected")
    else:
        try:
            validate_slot(selected_slot)
            # Try to parse the datetime
            datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
        except ValidationError as e:
            errors.append(f"Invalid slot: {e.message}")
        except ValueError as e:
            errors.append(f"Invalid date/time format: {e}")

    # Check Graph API configuration
    client = _make_graph_client()
    if not client:
        warnings.append("Graph API not configured - invites will fail unless configured")

    # Check Teams meeting capability
    if is_teams and not client:
        warnings.append("Teams meeting requested but Graph API is not available")

    # Build summary
    is_valid = len(errors) == 0
    if is_valid:
        summary = f"Ready to send invites to {len(intended_recipients)} recipient(s)"
    else:
        summary = f"Cannot send: {len(errors)} error(s) found"

    return ValidationReport(
        is_valid=is_valid,
        intended_recipients=intended_recipients,
        errors=errors,
        warnings=warnings,
        summary=summary,
    )


def _create_individual_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidate: Tuple[str, str],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> SchedulingResult:
    """
    Create a single individual interview invite.
    Returns SchedulingResult with success/failure status.
    """
    candidate_email_raw, candidate_name = candidate
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    if not is_valid_timezone(tz_name):
        tz_name = "UTC"
    if not is_valid_timezone(candidate_timezone):
        candidate_timezone = tz_name

    try:
        candidate_email = validate_email(candidate_email_raw, "Candidate email")
        # Hiring manager email is only required if no panel interviewers
        if panel_interviewers:
            hm_email = validate_email_optional(hm_email_raw, "Hiring manager email")
        else:
            hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=candidate_email_raw,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e),  # Include field name in error message
            warnings=None,
            recipients=None,
        )

    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid slot: {e.message}",
            warnings=None,
            recipients=None,
        )

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid date/time format: {e}",
            warnings=None,
            recipients=None,
        )

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Build attendees list - candidate and hiring manager are required (To:)
    attendees: List[Tuple[str, str]] = [(candidate_email, candidate_name)]
    seen_emails: set = {candidate_email.lower()}
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []
    skipped_panel_members: List[str] = []
    cc_attendees: List[Tuple[str, str]] = []  # Interviewers go in CC

    # ALWAYS add hiring manager as required attendee (if valid and not duplicate)
    if hm_email and hm_email.lower() not in seen_emails:
        attendees.append((hm_email, hm_name))
        seen_emails.add(hm_email.lower())

    # Add panel interviewers to CC (optional attendees) with deduplication
    if panel_interviewers:
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    cc_attendees.append((validated_email, pi.get("name", "")))  # CC instead of required
                    seen_emails.add(validated_email.lower())
                except ValidationError as ve:
                    skipped_panel_members.append(f"{pi_email}: {ve.message}")

    # Optionally add recruiter to CC
    if include_recruiter and rec_email and rec_email.lower() not in seen_emails:
        cc_attendees.append((rec_email, rec_name))  # CC instead of required
        seen_emails.add(rec_email.lower())

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Ensure we have a display name for subject
    effective_name = _ensure_candidate_name(candidate_name, candidate_email)
    has_role = role_title and role_title.strip()

    effective_subject = subject
    if is_panel:
        if not subject.startswith("Panel Interview"):
            if has_role:
                effective_subject = f"Panel Interview: {role_title.strip()} - {effective_name}"
            else:
                effective_subject = f"Panel Interview with {effective_name}"
    elif not subject:
        if has_role:
            effective_subject = f"Interview: {role_title.strip()} - {effective_name}"
        else:
            effective_subject = f"Interview with {effective_name}"

    # Generate ICS fallback
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|{candidate_email}|{hm_email}", organizer_email, start_utc.isoformat())

    client = _make_graph_client()
    if not client:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error="Graph API not configured",
            warnings=None,
            recipients=None,
        )

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    # Build professional HTML body using helper
    body_html = _build_professional_invite_body(
        time_display=candidate_time_display,
        role_title=role_title,
        duration_minutes=duration_minutes,
        panel_members=validated_panel if is_panel else None,
        agenda=agenda,
    )

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,
        attendees=attendees,
        cc_attendees=cc_attendees,  # Interviewers in CC
        is_teams=is_teams,
        location=location,
    )

    # Log payload for debugging Teams issues
    log_structured(
        LogLevel.INFO,
        f"Creating individual event with is_teams={is_teams}",
        action="create_individual_event",
        details={
            "is_teams": is_teams,
            "isOnlineMeeting": payload.get("isOnlineMeeting"),
            "onlineMeetingProvider": payload.get("onlineMeetingProvider"),
            "attendee_count": len(attendees),
        },
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        warnings: List[str] = []

        # Add warnings for skipped panel members
        if skipped_panel_members:
            warnings.append(f"Skipped invalid panel emails: {', '.join(skipped_panel_members)}")

        if is_teams:
            online_meeting = created.get("onlineMeeting")
            if online_meeting:
                teams_url = online_meeting.get("joinUrl") or ""

            # Try alternative locations for Teams URL
            if not teams_url:
                # Check webLink field
                web_link = created.get("webLink", "")
                if web_link and "teams.microsoft.com" in web_link:
                    teams_url = web_link

            # If Teams URL not in initial response, wait and fetch the event again
            # (Graph API sometimes needs a moment to generate the Teams meeting)
            if not teams_url and event_id:
                import time
                time.sleep(2)  # Wait 2 seconds for Teams meeting to be provisioned
                try:
                    refreshed = client.get_event(event_id)
                    online_meeting = refreshed.get("onlineMeeting")
                    if online_meeting:
                        teams_url = online_meeting.get("joinUrl") or ""
                    # Also check webLink in refreshed response
                    if not teams_url:
                        web_link = refreshed.get("webLink", "")
                        if web_link and "teams.microsoft.com" in web_link:
                            teams_url = web_link
                except Exception:
                    pass  # Ignore errors on retry, we'll show warning below

            if not teams_url:
                # Log what we got from Graph for debugging
                log_structured(
                    LogLevel.WARNING,
                    "Teams URL not found in Graph response",
                    action="create_event_teams",
                    details={
                        "has_onlineMeeting": created.get("onlineMeeting") is not None,
                        "onlineMeeting_keys": list(created.get("onlineMeeting", {}).keys()) if created.get("onlineMeeting") else [],
                        "has_webLink": bool(created.get("webLink")),
                        "isOnlineMeeting": created.get("isOnlineMeeting"),
                    },
                )
                warnings.append("Teams meeting was requested but no join URL was returned")

        # Validate Graph response contains expected attendees
        response_attendees = created.get("attendees", [])
        response_emails = {
            (a.get("emailAddress", {}).get("address") or "").lower()
            for a in response_attendees
        }
        expected_emails = {email.lower() for email, _ in attendees}
        missing = expected_emails - response_emails
        if missing:
            warnings.append(f"Some attendees may not be on invite: {', '.join(missing)}")

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            panel_json = json.dumps(validated_panel)

        audit.upsert_interview(
            role_title=role_title,
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
        )

        audit.log(
            "graph_create_event",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload=payload,
            status="success",
        )

        # Send meeting invitation email explicitly (Graph app-only auth doesn't auto-send invites)
        all_recipient_emails = [a[0] for a in attendees]
        cc_recipient_emails = [a[0] for a in cc_attendees] if cc_attendees else []
        try:
            client.send_meeting_invite(
                subject=effective_subject,
                html_body=body_html,
                ics_bytes=ics_bytes,
                to_recipients=all_recipient_emails,
                cc_recipients=cc_recipient_emails if cc_recipient_emails else None,
                organizer_name=organizer_name,
                organizer_email=organizer_email,
            )
            log_structured(

                LogLevel.INFO,
                f"Sent meeting invitation email to {len(all_recipient_emails)} recipients",
                action="send_invite_email",
                details={
                    "to": all_recipient_emails,
                    "cc": cc_recipient_emails,
                    "subject": effective_subject,
                },
            )
        except Exception as mail_err:
            log_structured(
                LogLevel.WARNING,
                f"Failed to send invitation email: {mail_err}",
                action="send_invite_email_failed",
                details={"error": str(mail_err)},
            )
            warnings.append(f"Calendar event created but email send failed: {mail_err}")

        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=True,
            event_id=event_id,
            teams_url=teams_url,
            error=None,
            warnings=warnings if warnings else None,
            recipients=[a[0] for a in attendees] + [a[0] for a in cc_attendees],
        )

    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        audit.log(
            "graph_create_failed",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details},
            status="failed",
            error_message=str(e),
        )
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e),
            warnings=None,
            recipients=None,
        )


def _create_group_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidates: List[CandidateValidationResult],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> SchedulingResult:
    """
    Create a single group interview invite with all candidates.
    Returns SchedulingResult with success/failure status.
    """
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    if not is_valid_timezone(tz_name):
        tz_name = "UTC"
    if not is_valid_timezone(candidate_timezone):
        candidate_timezone = tz_name

    try:
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e),  # Include field name in error message
            warnings=None,
            recipients=None,
        )

    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid slot: {e.message}",
            warnings=None,
            recipients=None,
        )

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid date/time format: {e}",
            warnings=None,
            recipients=None,
        )

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Build attendees list - candidates and hiring manager are required (To:)
    valid_candidates = [c for c in candidates if c.is_valid and c.email]
    attendees: List[Tuple[str, str]] = [(c.email, c.name) for c in valid_candidates]
    seen_emails: set = {(c.email or "").lower() for c in valid_candidates}
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []
    skipped_panel_members: List[str] = []
    cc_attendees: List[Tuple[str, str]] = []  # Interviewers go in CC

    # ALWAYS add hiring manager as required attendee (if valid and not duplicate)
    if hm_email and hm_email.lower() not in seen_emails:
        attendees.append((hm_email, hm_name))
        seen_emails.add(hm_email.lower())

    # Add panel interviewers to CC (optional attendees) with deduplication
    if panel_interviewers:
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    cc_attendees.append((validated_email, pi.get("name", "")))  # CC instead of required
                    seen_emails.add(validated_email.lower())
                except ValidationError as ve:
                    skipped_panel_members.append(f"{pi_email}: {ve.message}")

    # Optionally add recruiter to CC
    if include_recruiter and rec_email and rec_email.lower() not in seen_emails:
        cc_attendees.append((rec_email, rec_name))  # CC instead of required
        seen_emails.add(rec_email.lower())

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Use group interview subject
    has_role = role_title and role_title.strip()
    if has_role:
        effective_subject = f"Group Interview: {role_title.strip()}"
    else:
        effective_subject = "Group Interview"
    if subject and not subject.startswith("Interview:"):
        effective_subject = subject

    # Build candidates JSON for storage
    candidates_json = json.dumps([
        {"email": c.email, "name": c.name} for c in valid_candidates
    ])

    # Generate ICS fallback
    primary_candidate = valid_candidates[0] if valid_candidates else None
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|group|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|group|{hm_email}", organizer_email, start_utc.isoformat())

    client = _make_graph_client()
    if not client:
        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error="Graph API not configured",
            warnings=None,
            recipients=None,
        )

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    # Build professional HTML body using helper
    candidate_names = [c.name if c.name else c.email for c in valid_candidates]
    body_html = _build_professional_invite_body(
        time_display=candidate_time_display,
        role_title=role_title,
        duration_minutes=duration_minutes,
        panel_members=validated_panel if is_panel else None,
        agenda=agenda,
        candidates=candidate_names,
    )

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,
        attendees=attendees,
        cc_attendees=cc_attendees,  # Interviewers in CC
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        warnings: List[str] = []

        # Add warnings for skipped panel members
        if skipped_panel_members:
            warnings.append(f"Skipped invalid panel emails: {', '.join(skipped_panel_members)}")

        if is_teams:
            online_meeting = created.get("onlineMeeting")
            if online_meeting:
                teams_url = online_meeting.get("joinUrl") or ""

            # Try alternative locations for Teams URL
            if not teams_url:
                # Check webLink field
                web_link = created.get("webLink", "")
                if web_link and "teams.microsoft.com" in web_link:
                    teams_url = web_link

            # If Teams URL not in initial response, wait and fetch the event again
            # (Graph API sometimes needs a moment to generate the Teams meeting)
            if not teams_url and event_id:
                import time
                time.sleep(2)  # Wait 2 seconds for Teams meeting to be provisioned
                try:
                    refreshed = client.get_event(event_id)
                    online_meeting = refreshed.get("onlineMeeting")
                    if online_meeting:
                        teams_url = online_meeting.get("joinUrl") or ""
                    # Also check webLink in refreshed response
                    if not teams_url:
                        web_link = refreshed.get("webLink", "")
                        if web_link and "teams.microsoft.com" in web_link:
                            teams_url = web_link
                except Exception:
                    pass  # Ignore errors on retry, we'll show warning below

            if not teams_url:
                # Log what we got from Graph for debugging
                log_structured(
                    LogLevel.WARNING,
                    "Teams URL not found in Graph response",
                    action="create_event_teams",
                    details={
                        "has_onlineMeeting": created.get("onlineMeeting") is not None,
                        "onlineMeeting_keys": list(created.get("onlineMeeting", {}).keys()) if created.get("onlineMeeting") else [],
                        "has_webLink": bool(created.get("webLink")),
                        "isOnlineMeeting": created.get("isOnlineMeeting"),
                    },
                )
                warnings.append("Teams meeting was requested but no join URL was returned")

        # Validate Graph response contains expected attendees
        response_attendees = created.get("attendees", [])
        response_emails = {
            (a.get("emailAddress", {}).get("address") or "").lower()
            for a in response_attendees
        }
        expected_emails = {email.lower() for email, _ in attendees}
        missing = expected_emails - response_emails
        if missing:
            warnings.append(f"Some attendees may not be on invite: {', '.join(missing)}")

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            panel_json = json.dumps(validated_panel)

        # Store with primary candidate email for backward compatibility
        audit.upsert_interview(
            role_title=role_title,
            candidate_email=primary_candidate.email if primary_candidate else "",
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
            candidates_json=candidates_json,
            is_group_interview=True,
        )

        audit.log(
            "graph_create_group_event",
            actor=rec_email or "",
            candidate_email=", ".join(c.email for c in valid_candidates),
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload={"candidates_count": len(valid_candidates)},
            status="success",
        )

        # Send meeting invitation email explicitly (Graph app-only auth doesn't auto-send invites)
        all_recipient_emails = [a[0] for a in attendees]
        cc_recipient_emails = [a[0] for a in cc_attendees] if cc_attendees else []
        try:
            client.send_meeting_invite(
                subject=effective_subject,
                html_body=body_html,
                ics_bytes=ics_bytes,
                to_recipients=all_recipient_emails,
                cc_recipients=cc_recipient_emails if cc_recipient_emails else None,
                organizer_name=organizer_name,
                organizer_email=organizer_email,
            )
            log_structured(

                LogLevel.INFO,
                f"Sent group meeting invitation email to {len(all_recipient_emails)} recipients",
                action="send_group_invite_email",
                details={
                    "to": all_recipient_emails,
                    "cc": cc_recipient_emails,
                    "subject": effective_subject,
                },
            )
        except Exception as mail_err:
            log_structured(
                LogLevel.WARNING,
                f"Failed to send group invitation email: {mail_err}",
                action="send_group_invite_email_failed",
                details={"error": str(mail_err)},
            )
            warnings.append(f"Calendar event created but email send failed: {mail_err}")

        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name=f"Group ({len(valid_candidates)} candidates)",
            success=True,
            event_id=event_id,
            teams_url=teams_url,
            error=None,
            warnings=warnings if warnings else None,
            recipients=[a[0] for a in attendees] + [a[0] for a in cc_attendees],
        )

    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        audit.log(
            "graph_create_group_failed",
            actor=rec_email or "",
            candidate_email=", ".join(c.email for c in valid_candidates),
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details, "candidates_count": len(valid_candidates)},
            status="failed",
            error_message=str(e),
        )
        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e),
            warnings=None,
            recipients=None,
        )


def _handle_create_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidate: Tuple[str, str],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> None:
    candidate_email_raw, candidate_name = candidate
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    # Validate timezones
    if not is_valid_timezone(tz_name):
        st.warning(f"Invalid display timezone '{tz_name}', using UTC")
        tz_name = "UTC"

    if not is_valid_timezone(candidate_timezone):
        st.warning(f"Invalid candidate timezone '{candidate_timezone}', using display timezone")
        candidate_timezone = tz_name

    # Validate emails
    try:
        candidate_email = validate_email(candidate_email_raw, "Candidate email")
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        st.error(f"Validation error: {e.message}")
        return

    # Validate slot format
    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        st.error(f"Invalid time slot: {e.message}")
        return

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        st.error(f"Selected slot has invalid date/time format: {e}")
        return

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)

    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # === IDEMPOTENCY CHECK ===
    existing = audit.interview_exists(
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        role_title=role_title,
        start_utc=iso_utc(start_utc),
    )
    if existing:
        st.warning(
            f"An interview already exists for this candidate at this time. "
            f"Event ID: {existing.get('graph_event_id', 'N/A')}"
        )
        # Use a unique key based on the slot to avoid Streamlit duplicate key errors
        checkbox_key = f"force_dup_{selected_slot['date']}_{selected_slot['start']}"
        if not st.checkbox("Create duplicate anyway?", key=checkbox_key):
            return

    attendees: List[Tuple[str, str]] = [(candidate_email, candidate_name)]
    cc_attendees: List[Tuple[str, str]] = []  # Interviewers go in CC

    # Build attendees from panel interviewers if provided, otherwise use hiring manager
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []

    if panel_interviewers:
        seen_emails = {candidate_email.lower()}  # Avoid duplicating candidate
        # Add hiring manager as required attendee
        if hm_email and hm_email.lower() not in seen_emails:
            attendees.append((hm_email, hm_name))
            seen_emails.add(hm_email.lower())
        # Add panel interviewers to CC
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    cc_attendees.append((validated_email, pi.get("name", "")))  # CC instead of required
                    seen_emails.add(validated_email.lower())
                except ValidationError:
                    pass  # Skip invalid emails
    else:
        # Fall back to single hiring manager (backward compatibility)
        attendees.append((hm_email, hm_name))

    if include_recruiter and rec_email:
        cc_attendees.append((rec_email, rec_name))  # Recruiter in CC

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Update subject for panel interviews
    effective_subject = subject
    if is_panel:
        if not subject.startswith("Panel Interview"):
            effective_subject = f"Panel Interview: {role_title} - {candidate_name}"

    # Always generate ICS (so we have a fallback even if Graph works)
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|{candidate_email}|{hm_email}", organizer_email, start_utc.isoformat())
    audit.log(
        "ics_generated",
        actor=rec_email or "",
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        recruiter_email=rec_email or "",
        role_title=role_title,
        payload={"uid": st.session_state["last_invite_uid"]},
        status="success",
    )

    client = _make_graph_client()
    if not client:
        st.warning("Graph is not configured. Using .ics fallback only.")
        return

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    # Build professional HTML body using helper
    body_html = _build_professional_invite_body(
        time_display=candidate_time_display,
        role_title=role_title,
        duration_minutes=duration_minutes,
        panel_members=validated_panel if is_panel else None,
        agenda=agenda,
    )

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,  # Use candidate timezone for calendar event
        attendees=attendees,
        cc_attendees=cc_attendees,  # Interviewers in CC
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        if is_teams:
            teams_url = (created.get("onlineMeeting") or {}).get("joinUrl") or ""
        st.session_state["last_graph_event_id"] = event_id
        st.session_state["last_teams_join_url"] = teams_url

        # Re-generate ICS including Teams URL if present (better fallback)
        if teams_url:
            st.session_state["last_invite_ics_bytes"] = _build_ics(
                organizer_email=organizer_email,
                organizer_name=organizer_name,
                attendee_emails=[a[0] for a in attendees],
                summary=effective_subject,
                description=agenda,
                dtstart_utc=start_utc,
                dtend_utc=end_utc,
                location="Microsoft Teams",
                url=teams_url,
                uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
                display_timezone=candidate_timezone,
            )

        audit.log(
            "graph_create_event",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload=payload,
            status="success",
        )

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            import json as _json
            panel_json = _json.dumps(validated_panel)

        audit.upsert_interview(
            role_title=role_title,
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
        )

        st.success("Invite created and sent via Microsoft Graph.")
        if teams_url:
            st.link_button("Open Teams meeting link", teams_url)
    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        st.error("Graph scheduling failed. .ics fallback is available for download.")
        if details:
            st.json(details)
        audit.log(
            "graph_create_failed",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details},
            status="failed",
            error_message=str(e),
        )


def _extract_candidate_name_from_context(context_row: Dict[str, Any]) -> str:
    """
    Extract candidate name from interview context row.

    Tries candidates_json first (for group interviews), then falls back to
    parsing from candidate_email if no name is found.
    """
    # Try to get name from candidates_json (for group/multi-candidate interviews)
    candidates_json = context_row.get("candidates_json")
    if candidates_json:
        try:
            candidates = json.loads(candidates_json)
            if candidates and isinstance(candidates, list) and len(candidates) > 0:
                first_candidate = candidates[0]
                if isinstance(first_candidate, dict) and first_candidate.get("name"):
                    return first_candidate["name"]
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    # Fall back to empty string - email will be used as fallback in templates
    return ""


def _format_interview_time_for_candidate(
    utc_time_str: str,
    candidate_timezone: Optional[str],
    display_timezone: Optional[str] = None,
) -> str:
    """
    Format interview time in candidate's timezone for notifications.

    Args:
        utc_time_str: ISO format UTC time string
        candidate_timezone: Candidate's preferred timezone
        display_timezone: Fallback display timezone

    Returns:
        Human-readable formatted time string
    """
    from timezone_utils import from_utc, is_valid_timezone

    # Determine which timezone to use
    tz_to_use = None
    if candidate_timezone and is_valid_timezone(candidate_timezone):
        tz_to_use = candidate_timezone
    elif display_timezone and is_valid_timezone(display_timezone):
        tz_to_use = display_timezone

    try:
        # Parse the UTC time
        dt_utc = datetime.fromisoformat(utc_time_str.replace("+00:00", "").replace("Z", ""))
        dt_utc = dt_utc.replace(tzinfo=timezone.utc)

        if tz_to_use:
            # Convert to candidate/display timezone
            dt_local = from_utc(dt_utc, tz_to_use)
            tz_abbrev = dt_local.strftime("%Z") or tz_to_use
            return dt_local.strftime(f"%A, %B %d, %Y at %I:%M %p {tz_abbrev}")
        else:
            # Fall back to UTC display
            return dt_utc.strftime("%A, %B %d, %Y at %I:%M %p UTC")
    except Exception:
        return utc_time_str


def _send_cancellation_email(
    client: GraphClient,
    candidate_email: str,
    candidate_name: str,
    role_title: str,
    interview_time: str,
    reason: str,
    custom_message: str,
    company: CompanyConfig,
) -> bool:
    """Send cancellation notification email to candidate."""
    try:
        html_body = build_cancellation_email_html(
            candidate_name=candidate_name,
            role_title=role_title,
            interview_time=interview_time,
            reason=reason,
            custom_message=custom_message if custom_message else None,
            company=company,
        )
        client.send_mail(
            subject=f"Interview Cancelled: {role_title} at {company.name}",
            body=html_body,
            to_recipients=[candidate_email],
            content_type="HTML",
        )
        log_structured(
            LogLevel.INFO,
            f"Cancellation email sent to {candidate_email}",
            action="cancellation_email_sent",
        )
        return True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to send cancellation email: {e}",
            action="cancellation_email_failed",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def _send_reschedule_email(
    client: GraphClient,
    candidate_email: str,
    candidate_name: str,
    role_title: str,
    old_time: str,
    new_time: str,
    teams_url: Optional[str],
    company: CompanyConfig,
) -> bool:
    """Send reschedule notification email to candidate."""
    try:
        html_body = build_reschedule_email_html(
            candidate_name=candidate_name,
            role_title=role_title,
            old_time=old_time,
            new_time=new_time,
            teams_url=teams_url,
            company=company,
        )
        client.send_mail(
            subject=f"Interview Rescheduled: {role_title} at {company.name}",
            body=html_body,
            to_recipients=[candidate_email],
            content_type="HTML",
        )
        log_structured(
            LogLevel.INFO,
            f"Reschedule email sent to {candidate_email}",
            action="reschedule_email_sent",
        )
        return True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to send reschedule email: {e}",
            action="reschedule_email_failed",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def _handle_reschedule(
    *,
    audit: AuditLog,
    event_id: str,
    new_date,
    new_time,
    duration_minutes: int,
    tz_name: str,
    context_row: Dict[str, Any],
    reason: str = "",
    notify_candidate: bool = True,
) -> None:
    """
    Handle interview reschedule with status tracking and notifications.

    Args:
        audit: AuditLog instance
        event_id: Graph event ID
        new_date: New interview date
        new_time: New interview time
        duration_minutes: Interview duration
        tz_name: Display timezone
        context_row: Interview data from database
        reason: Reason for reschedule
        notify_candidate: Whether to send email notification
    """
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    # Calculate new times
    start_local = datetime.combine(new_date, new_time).replace(tzinfo=_zoneinfo(tz_name))
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Store old time for notification - use candidate timezone for better UX
    old_time_str = context_row.get("start_utc", "")
    candidate_tz = context_row.get("candidate_timezone") or context_row.get("display_timezone") or tz_name
    old_time_formatted = _format_interview_time_for_candidate(old_time_str, candidate_tz, tz_name)

    # Format new time in candidate's timezone
    new_time_formatted = start_local.strftime("%A, %B %d, %Y at %I:%M %p") + f" ({tz_name})"

    patch = {
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
    }

    try:
        # Patch the Graph event
        client.patch_event(event_id, patch, send_updates="all")

        # Update interview status in database
        audit.update_interview_status(
            event_id=event_id,
            new_status=InterviewStatus.RESCHEDULED,
            reason=reason,
            updated_by=context_row.get("recruiter_email"),
        )

        # Increment ICS sequence for proper calendar client update
        audit.increment_ics_sequence(event_id)

        # Log success
        audit.log(
            "interview_rescheduled",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload={
                "old_start": old_time_str,
                "new_start": start_utc.isoformat(),
                "new_end": end_utc.isoformat(),
                "reason": reason,
                "notification_sent": notify_candidate,
            },
            status="success",
        )

        # Send notification email
        notification_sent = False
        if notify_candidate:
            company = get_company_config()
            candidate_email = context_row.get("candidate_email", "")
            if candidate_email:
                # Extract candidate name from context (candidates_json or fallback)
                candidate_name = _extract_candidate_name_from_context(context_row)
                notification_sent = _send_reschedule_email(
                    client=client,
                    candidate_email=candidate_email,
                    candidate_name=candidate_name,
                    role_title=context_row.get("role_title", ""),
                    old_time=old_time_formatted,
                    new_time=new_time_formatted,
                    teams_url=context_row.get("teams_join_url"),
                    company=company,
                )
                if notification_sent:
                    audit.log(
                        "notification_sent",
                        actor=context_row.get("recruiter_email", "") or "",
                        candidate_email=candidate_email,
                        event_id=event_id,
                        payload={"type": "reschedule"},
                        status="success",
                    )

        if notify_candidate and notification_sent:
            st.success("Event rescheduled. Candidate notified via email.")
        else:
            st.success("Event rescheduled. Attendees should receive updated invites.")

    except GraphAPIError as e:
        st.error("Reschedule failed.")
        st.json(e.response_json)
        audit.log(
            "interview_reschedule_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


def _handle_cancel(
    *,
    audit: AuditLog,
    event_id: str,
    context_row: Dict[str, Any],
    reason: str = "",
    notify_candidate: bool = True,
    candidate_message: str = "",
) -> None:
    """
    Handle interview cancellation with status tracking and notifications.

    Args:
        audit: AuditLog instance
        event_id: Graph event ID
        context_row: Interview data from database
        reason: Cancellation reason
        notify_candidate: Whether to send email notification
        candidate_message: Optional custom message for candidate
    """
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    # Format interview time for notification - use candidate timezone for better UX
    interview_time_str = context_row.get("start_utc", "")
    candidate_tz = context_row.get("candidate_timezone") or context_row.get("display_timezone")
    interview_time_formatted = _format_interview_time_for_candidate(interview_time_str, candidate_tz)

    try:
        # Delete the calendar event
        client.delete_event(event_id)

        # Update interview status in database
        audit.update_interview_status(
            event_id=event_id,
            new_status=InterviewStatus.CANCELLED,
            reason=reason,
            updated_by=context_row.get("recruiter_email"),
        )

        # Log success
        audit.log(
            "interview_cancelled",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload={
                "reason": reason,
                "notification_sent": notify_candidate,
            },
            status="success",
        )

        # Send notification email
        notification_sent = False
        if notify_candidate:
            company = get_company_config()
            candidate_email = context_row.get("candidate_email", "")
            if candidate_email:
                # Extract candidate name from context (candidates_json or fallback)
                candidate_name = _extract_candidate_name_from_context(context_row)
                notification_sent = _send_cancellation_email(
                    client=client,
                    candidate_email=candidate_email,
                    candidate_name=candidate_name,
                    role_title=context_row.get("role_title", ""),
                    interview_time=interview_time_formatted,
                    reason=reason,
                    custom_message=candidate_message,
                    company=company,
                )
                if notification_sent:
                    audit.log(
                        "notification_sent",
                        actor=context_row.get("recruiter_email", "") or "",
                        candidate_email=candidate_email,
                        event_id=event_id,
                        payload={"type": "cancellation", "reason": reason},
                        status="success",
                    )

        if notify_candidate and notification_sent:
            st.success("Interview cancelled. Candidate notified via email.")
        else:
            st.success("Interview cancelled. Attendees should receive cancellation notices.")

    except GraphAPIError as e:
        st.error("Cancel failed.")
        st.json(e.response_json)
        audit.log(
            "interview_cancel_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


if __name__ == "__main__":
    main()
